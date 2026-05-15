"""
Unit tests for GLOW MCP Server endpoints using FastAPI TestClient.
"""
import io
from pathlib import Path
from fastapi.testclient import TestClient
import pytest



import sys
from pathlib import Path
# Add project root and mcp_server to sys.path
project_root = Path(__file__).parent.parent.parent.resolve()
sys.path.insert(0, str(project_root))
sys.path.insert(0, str(project_root / "mcp_server"))
from main import app

client = TestClient(app)


def test_health():
    resp = client.get("/health")
    assert resp.status_code == 200
    assert resp.json()["status"] == "ok"


def test_audit_markdown(tmp_path):
    md_file = tmp_path / "test.md"
    md_file.write_text("# Heading\n\nSome text.")
    with md_file.open("rb") as f:
        resp = client.post("/audit", files={"file": ("test.md", f, "text/markdown")}, data={"format": "markdown"})
    assert resp.status_code == 200
    assert "score" in resp.text or "result" in resp.text


def test_audit_docx(tmp_path):
    # Minimal DOCX file (could use python-docx to generate, or skip if not available)
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")
    docx_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_heading("Test Heading", 0)
    doc.add_paragraph("Some text.")
    doc.save(docx_file)
    with docx_file.open("rb") as f:
        resp = client.post("/audit", files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}, data={"format": "docx"})
    assert resp.status_code == 200
    assert "score" in resp.text or "result" in resp.text

# Additional tests for /fix, /convert, /report can be added similarly.
