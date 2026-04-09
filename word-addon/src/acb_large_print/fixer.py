"""Fix Word documents to comply with ACB Large Print Guidelines."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from . import constants as C
from .auditor import AuditResult, audit_document


def _apply_acb_font(font_obj, spec: C.FontDef) -> None:
    """Apply ACB font settings to a python-docx font object."""
    font_obj.name = spec.name
    font_obj.size = Pt(spec.size_pt)
    font_obj.bold = spec.bold
    font_obj.italic = False
    font_obj.all_caps = False
    font_obj.color.rgb = RGBColor(0, 0, 0)


def _fix_styles(doc: Document) -> int:
    """Fix all named styles to match ACB specs. Returns count of fixes."""
    fixes = 0
    for style_name, style_def in C.ACB_STYLES.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue

        font = style.font
        pf = style.paragraph_format
        changed = False

        # Font
        if font.name != style_def.font.name:
            font.name = style_def.font.name
            changed = True
        if font.size is None or abs(font.size.pt - style_def.font.size_pt) > 0.5:
            font.size = Pt(style_def.font.size_pt)
            changed = True
        if font.bold != style_def.font.bold:
            font.bold = style_def.font.bold
            changed = True
        if font.italic:
            font.italic = False
            changed = True
        if font.all_caps:
            font.all_caps = False
            changed = True

        # Paragraph format
        if pf.alignment != WD_ALIGN_PARAGRAPH.LEFT:
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            changed = True
        if pf.line_spacing != style_def.para.line_spacing:
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = style_def.para.line_spacing
            changed = True
        if pf.widow_control is not True:
            pf.widow_control = True
            changed = True

        # Remove theme colors from headings
        if style_name.startswith("Heading"):
            rPr = style.element.find(qn("w:rPr"))
            if rPr is not None:
                color_elem = rPr.find(qn("w:color"))
                if color_elem is not None:
                    theme = color_elem.get(qn("w:themeColor"))
                    if theme:
                        color_elem.attrib.pop(qn("w:themeColor"), None)
                        color_elem.attrib.pop(qn("w:themeShade"), None)
                        color_elem.set(qn("w:val"), "000000")
                        changed = True

        if changed:
            fixes += 1

    return fixes


def _fix_page_setup(doc: Document, bound: bool = False) -> int:
    """Fix page margins and orientation. Returns count of fixes."""
    fixes = 0
    tolerance = Inches(0.05)
    for section in doc.sections:
        left_target = C.MARGIN_LEFT_IN + (C.BINDING_EXTRA_IN if bound else 0)

        if section.top_margin is None or abs(section.top_margin - Inches(C.MARGIN_TOP_IN)) > tolerance:
            section.top_margin = Inches(C.MARGIN_TOP_IN)
            fixes += 1
        if section.bottom_margin is None or abs(section.bottom_margin - Inches(C.MARGIN_BOTTOM_IN)) > tolerance:
            section.bottom_margin = Inches(C.MARGIN_BOTTOM_IN)
            fixes += 1
        if section.left_margin is None or abs(section.left_margin - Inches(left_target)) > tolerance:
            section.left_margin = Inches(left_target)
            fixes += 1
        if section.right_margin is None or abs(section.right_margin - Inches(C.MARGIN_RIGHT_IN)) > tolerance:
            section.right_margin = Inches(C.MARGIN_RIGHT_IN)
            fixes += 1

    return fixes


def _fix_paragraph_formatting(doc: Document) -> int:
    """Fix direct formatting overrides in paragraph content. Returns fix count."""
    fixes = 0

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        is_heading = bool(re.match(r"^Heading \d+$", style_name))

        # Fix alignment
        if para.paragraph_format.alignment is not None:
            if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                fixes += 1

        for run in para.runs:
            # Remove italic everywhere
            if run.font.italic:
                run.font.italic = False
                # Convert italic text to underline (ACB emphasis)
                run.font.underline = True
                fixes += 1

            # Fix bold in body text: convert to underline
            if run.font.bold and not is_heading and run.text.strip():
                if len(run.text.strip()) > 5:
                    run.font.bold = False
                    run.font.underline = True
                    fixes += 1

            # Fix font family
            if run.font.name and run.font.name != C.FONT_FAMILY:
                run.font.name = C.FONT_FAMILY
                fixes += 1

            # Fix font size below minimum
            if run.font.size and run.font.size.pt < C.MIN_SIZE_PT - 0.5:
                run.font.size = Pt(C.MIN_SIZE_PT)
                fixes += 1

            # Remove all caps
            if run.font.all_caps:
                run.font.all_caps = False
                fixes += 1

    return fixes


def _fix_hyphenation(doc: Document) -> int:
    """Disable automatic hyphenation. Returns 1 if fixed, 0 if already correct."""
    settings = doc.settings.element
    auto_hyph = settings.find(qn("w:autoHyphenation"))
    if auto_hyph is not None:
        val = auto_hyph.get(qn("w:val"), "true")
        if val.lower() not in ("false", "0", "off"):
            auto_hyph.set(qn("w:val"), "false")
            return 1
        return 0
    # Add element explicitly set to false
    elem = OxmlElement("w:autoHyphenation")
    elem.set(qn("w:val"), "false")
    settings.append(elem)
    return 1


def _fix_page_numbers(doc: Document) -> int:
    """Add page numbers to footer if missing. Returns 1 if added."""
    has_page_numbers = False
    for section in doc.sections:
        footer = section.footer
        if footer is None:
            continue
        footer_xml = footer._element.xml
        if "PAGE" in footer_xml or "w:fldChar" in footer_xml:
            has_page_numbers = True
            break

    if has_page_numbers:
        return 0

    # Add page numbers to first section footer (others will link)
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = paragraph.add_run()
    run.font.name = C.FONT_FAMILY
    run.font.size = Pt(C.FOOTER_SIZE_PT)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    return 1


def _fix_document_language(doc: Document, lang: str = "en-US") -> int:
    """Set document language if missing. Returns 1 if fixed."""
    styles_elem = doc.styles.element
    rpr = styles_elem.find(qn("w:docDefaults/w:rPrDefault/w:rPr"))
    if rpr is not None:
        lang_elem = rpr.find(qn("w:lang"))
        if lang_elem is None:
            lang_elem = OxmlElement("w:lang")
            rpr.append(lang_elem)
        existing = lang_elem.get(qn("w:val"))
        if existing and existing.strip():
            return 0
        lang_elem.set(qn("w:val"), lang)
        lang_elem.set(qn("w:eastAsia"), lang)
        lang_elem.set(qn("w:bidi"), lang)
        return 1
    return 0


def fix_document(
    file_path: str | Path,
    output_path: str | Path | None = None,
    *,
    bound: bool = False,
) -> tuple[Path, int, AuditResult]:
    """Fix a Word document for ACB compliance.

    Args:
        file_path: Input .docx file.
        output_path: Where to save fixed document. Defaults to overwriting.
        bound: Add binding margin if True.

    Returns:
        Tuple of (output_path, total_fixes, post_fix_audit_result).
    """
    file_path = Path(file_path)
    if output_path is None:
        output_path = file_path
    output_path = Path(output_path)

    doc = Document(str(file_path))

    total_fixes = 0
    total_fixes += _fix_styles(doc)
    total_fixes += _fix_page_setup(doc, bound=bound)
    total_fixes += _fix_paragraph_formatting(doc)
    total_fixes += _fix_hyphenation(doc)
    total_fixes += _fix_page_numbers(doc)
    total_fixes += _fix_document_language(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    # Run post-fix audit to show remaining issues
    post_audit = audit_document(output_path)

    return output_path, total_fixes, post_audit
