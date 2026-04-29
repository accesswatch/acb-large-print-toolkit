"""Fix route -- upload a document and download a remediated copy."""

from pathlib import Path

from flask import Blueprint, current_app, render_template, request, send_file
from werkzeug.utils import secure_filename

from ..rules import (
    build_rule_policy,
    filter_findings,
    get_all_rule_ids,
    get_profile_label,
    get_rule_ids_by_category,
    get_rule_ids_by_format,
    get_rule_ids_by_profile,
    get_rule_ids_by_severity,
    get_rules_by_format,
)
from flask import url_for

from ..upload import (
    UploadError,
    cleanup_token,
    get_temp_dir,
    validate_upload,
)
from ..gating import ai_gate, GatingError, RETRY_AFTER_SECONDS
from ..customization_warning import detect_fix_customizations, generate_customization_warning

fix_bp = Blueprint("fix", __name__)


def _busy_response(operation: str, back_url: str):
    """Render the 503 busy page for gating rejections."""
    from flask import render_template as _rt, make_response
    resp = make_response(
        _rt(
            "busy.html",
            operation=operation,
            retry_seconds=RETRY_AFTER_SECONDS,
            back_url=back_url,
        ),
        503,
    )
    resp.headers["Retry-After"] = str(RETRY_AFTER_SECONDS)
    return resp


_SEVERITY_DEDUCTIONS = {
    "critical": 15,
    "high": 10,
    "medium": 5,
    "low": 2,
}


def _severity_key(severity: object) -> str:
    """Normalize severity from enum/string to lowercase key."""
    value = getattr(severity, "value", severity)
    return str(value).strip().lower()


def _build_penalty_breakdown(findings: list) -> dict[str, int]:
    """Summarize weighted score penalty and counts by severity."""
    critical = high = medium = low = 0
    for finding in findings:
        sev = _severity_key(getattr(finding, "severity", ""))
        if sev == "critical":
            critical += 1
        elif sev == "high":
            high += 1
        elif sev == "medium":
            medium += 1
        elif sev == "low":
            low += 1

    weighted_penalty = (
        critical * _SEVERITY_DEDUCTIONS["critical"]
        + high * _SEVERITY_DEDUCTIONS["high"]
        + medium * _SEVERITY_DEDUCTIONS["medium"]
        + low * _SEVERITY_DEDUCTIONS["low"]
    )

    return {
        "critical": critical,
        "high": high,
        "medium": medium,
        "low": low,
        "total": len(findings),
        "weighted_penalty": weighted_penalty,
    }


def _fix_by_extension(
    saved_path: Path,
    output_path: Path,
    *,
    bound: bool = False,
    list_indent_in: float = 0.0,
    list_hanging_in: float = 0.0,
    list_level_indents: dict[int, float] | None = None,
    para_indent_in: float = 0.0,
    first_line_indent_in: float = 0.0,
    preserve_heading_alignment: bool = False,
    detect_headings: bool = False,
    use_ai: bool = False,
    heading_threshold: int | None = None,
    confirmed_headings: list | None = None,
    heading_accuracy: str = "balanced",
):
    """Dispatch to the correct fixer based on file extension.

    Returns (output_path, total_fixes, fix_records, post_audit, warnings).
    """
    ext = saved_path.suffix.lower()
    if ext == ".xlsx":
        from acb_large_print.xlsx_auditor import audit_workbook

        post_audit = audit_workbook(saved_path)
        return (
            saved_path,
            0,
            [],
            post_audit,
            [
                "Excel workbooks cannot be auto-fixed yet. "
                "Review the audit findings and fix them manually in Excel."
            ],
        )
    elif ext == ".pptx":
        from acb_large_print.pptx_auditor import audit_presentation

        post_audit = audit_presentation(saved_path)
        return (
            saved_path,
            0,
            [],
            post_audit,
            [
                "PowerPoint presentations cannot be auto-fixed yet. "
                "Review the audit findings and fix them manually in PowerPoint."
            ],
        )
    elif ext == ".md":
        from acb_large_print.md_auditor import audit_markdown

        post_audit = audit_markdown(saved_path)
        return (
            saved_path,
            0,
            [],
            post_audit,
            [
                "Markdown auto-fix is coming soon. "
                "Review the audit findings and fix them in your text editor."
            ],
        )
    elif ext == ".pdf":
        from acb_large_print.pdf_auditor import audit_pdf

        post_audit = audit_pdf(saved_path)
        return (
            saved_path,
            0,
            [],
            post_audit,
            [
                "PDF files cannot be auto-fixed. "
                "Use Adobe Acrobat Pro or re-export from the source application."
            ],
        )
    elif ext == ".epub":
        from acb_large_print.epub_auditor import audit_epub

        post_audit = audit_epub(saved_path)
        return (
            saved_path,
            0,
            [],
            post_audit,
            [
                "ePub files cannot be auto-fixed yet. "
                "Review the audit findings and fix them in your ePub editor."
            ],
        )
    else:
        from acb_large_print.fixer import fix_document

        ai_provider = None
        if detect_headings and use_ai:
            try:
                from acb_large_print.ai_provider import get_provider

                ai_provider = get_provider()
            except Exception:
                pass  # Fall back to heuristic-only
        result = fix_document(
            saved_path,
            output_path,
            bound=bound,
            list_indent_in=list_indent_in,
            list_hanging_in=list_hanging_in,
            list_level_indents=list_level_indents,
            para_indent_in=para_indent_in,
            first_line_indent_in=first_line_indent_in,
            preserve_heading_alignment=preserve_heading_alignment,
            detect_headings=detect_headings,
            ai_provider=ai_provider,
            heading_threshold=heading_threshold,
            confirmed_headings=confirmed_headings,
            heading_accuracy_level=heading_accuracy,
        )
        # Tag whether AI was actually used (ai_provider set and invoked)
        return result[:5] + ({"ai_used": ai_provider is not None},) if len(result) == 5 else result


def _audit_by_extension(
    saved_path: Path,
    *,
    list_indent_in: float | None = None,
    list_level_indents: dict[int, float] | None = None,
    para_indent_in: float | None = None,
    first_line_indent_in: float | None = None,
):
    """Dispatch to the correct auditor based on file extension."""
    ext = saved_path.suffix.lower()
    if ext == ".xlsx":
        from acb_large_print.xlsx_auditor import audit_workbook

        return audit_workbook(saved_path)
    elif ext == ".pptx":
        from acb_large_print.pptx_auditor import audit_presentation

        return audit_presentation(saved_path)
    elif ext == ".md":
        from acb_large_print.md_auditor import audit_markdown

        return audit_markdown(saved_path)
    elif ext == ".pdf":
        from acb_large_print.pdf_auditor import audit_pdf

        return audit_pdf(saved_path)
    elif ext == ".epub":
        from acb_large_print.epub_auditor import audit_epub

        return audit_epub(saved_path)
    else:
        from acb_large_print.auditor import audit_document

        return audit_document(
            saved_path,
            list_indent_in=list_indent_in,
            list_level_indents=list_level_indents,
            para_indent_in=para_indent_in,
            first_line_indent_in=first_line_indent_in,
        )


def _format_from_path(saved_path: Path) -> str:
    return saved_path.suffix.lower().lstrip(".")


def _parse_allowed_heading_levels(form) -> list[int]:
    """Parse allowed heading levels from form data with safe defaults."""
    raw_levels = form.getlist("allowed_heading_levels")
    parsed_levels: list[int] = []
    for raw in raw_levels:
        try:
            level = int(raw)
        except (TypeError, ValueError):
            continue
        if 1 <= level <= 6:
            parsed_levels.append(level)

    if not parsed_levels:
        return [1, 2, 3, 4, 5, 6]

    return sorted(set(parsed_levels))


def _closest_allowed_level(level: int, allowed_levels: list[int]) -> int:
    """Choose the nearest allowed heading level for out-of-range suggestions."""
    if not allowed_levels:
        return level
    if level in allowed_levels:
        return level
    return min(allowed_levels, key=lambda allowed: abs(allowed - level))


def _estimate_pre_fix_body_font_pt(saved_path: Path) -> float | None:
    """Estimate typical body font size (pt) for Word docs using run-level sizes."""
    if saved_path.suffix.lower() != ".docx":
        return None

    try:
        from docx import Document

        sizes: list[float] = []
        doc = Document(str(saved_path))

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            style_name = (para.style.name if para.style is not None else "").lower()
            if "heading" in style_name:
                continue

            for run in para.runs:
                if not run.text.strip():
                    continue
                if run.font.size is not None:
                    sizes.append(float(run.font.size.pt))

        if not sizes:
            return None

        sizes.sort()
        return round(sizes[len(sizes) // 2], 1)
    except Exception:
        return None


@fix_bp.route("/", methods=["GET"])
def fix_form():
    from ..ai_features import ai_heading_fix_enabled

    return render_template(
        "fix_form.html",
        ai_available=ai_heading_fix_enabled(),
        rules_by_format=get_rules_by_format(),
    )


def _parse_form_options(form):
    """Extract fix options from the submitted form (used by both submit and confirm)."""
    from ..ai_features import ai_heading_fix_enabled

    bound = form.get("bound") == "on"
    flush_lists = form.get("flush_lists") == "on"
    use_list_levels = form.get("use_list_levels") == "on"
    mode = form.get("mode", "full")
    preserve_heading_alignment = form.get("preserve_heading_alignment") == "on"
    # Legacy boolean suppression fields kept for backward-compat; policy absorbs them.
    suppress_link_text = form.get("suppress_link_text") == "on"
    suppress_missing_alt_text = form.get("suppress_missing_alt_text") == "on"
    suppress_faux_heading = form.get("suppress_faux_heading") == "on"

    if flush_lists:
        list_indent_in = 0.0
        list_hanging_in = 0.0
    else:
        try:
            list_indent_in = float(form.get("list_indent", "0.5"))
        except (ValueError, TypeError):
            list_indent_in = 0.5
        try:
            list_hanging_in = float(form.get("list_hanging", "0.25"))
        except (ValueError, TypeError):
            list_hanging_in = 0.25
        list_indent_in = max(0.0, min(2.0, list_indent_in))
        list_hanging_in = max(0.0, min(2.0, list_hanging_in))

    list_level_indents = None
    if not flush_lists and use_list_levels:
        list_level_indents = {}
        defaults = {1: 0.25, 2: 0.50, 3: 0.75}
        for level in (1, 2, 3):
            key = f"list_indent_level_{level}"
            try:
                value = float(form.get(key, str(defaults[level])))
            except (ValueError, TypeError):
                value = defaults[level]
            list_level_indents[level - 1] = max(0.0, min(2.0, value))

    flush_paragraphs = form.get("flush_paragraphs") == "on"
    if flush_paragraphs:
        para_indent_in = 0.0
        first_line_indent_in = 0.0
    else:
        try:
            para_indent_in = float(form.get("para_indent", "0.0"))
        except (ValueError, TypeError):
            para_indent_in = 0.0
        try:
            first_line_indent_in = float(form.get("first_line_indent", "0.0"))
        except (ValueError, TypeError):
            first_line_indent_in = 0.0
        para_indent_in = max(0.0, min(2.0, para_indent_in))
        first_line_indent_in = max(0.0, min(2.0, first_line_indent_in))

    detect_headings = form.get("detect_headings") == "on"
    use_ai = form.get("use_ai") == "on" and ai_heading_fix_enabled()
    try:
        heading_threshold = int(form.get("heading_threshold", "50"))
        heading_threshold = max(0, min(100, heading_threshold))
    except (ValueError, TypeError):
        heading_threshold = 50
    
    heading_accuracy = form.get("heading_accuracy", "balanced")
    if heading_accuracy not in ("conservative", "balanced", "thorough"):
        heading_accuracy = "balanced"

    allowed_heading_levels = _parse_allowed_heading_levels(form)

    return {
        "bound": bound,
        "mode": mode,
        "list_indent_in": list_indent_in,
        "list_hanging_in": list_hanging_in,
        "list_level_indents": list_level_indents,
        "para_indent_in": para_indent_in,
        "first_line_indent_in": first_line_indent_in,
        "preserve_heading_alignment": preserve_heading_alignment,
        "detect_headings": detect_headings,
        "use_ai": use_ai,
        "suppress_link_text": suppress_link_text,
        "suppress_missing_alt_text": suppress_missing_alt_text,
        "suppress_faux_heading": suppress_faux_heading,
        "heading_threshold": heading_threshold,
        "heading_accuracy": heading_accuracy,
        "allowed_heading_levels": allowed_heading_levels,
        "rule_policy": build_rule_policy(form),
    }


def _is_heading_alignment_finding(finding) -> bool:
    """Return True when an ACB-ALIGNMENT finding refers to heading content."""
    if getattr(finding, "rule_id", "") != "ACB-ALIGNMENT":
        return False
    msg = str(getattr(finding, "message", "")).lower()
    loc = str(getattr(finding, "location", "")).lower()
    return "heading" in msg or "style: heading" in loc


def _run_fix_and_render(
    saved_path,
    token,
    opts,
    *,
    confirmed_headings=None,
    confirmed_headings_received: int | None = None,
):
    """Run the fixer and render fix_result.html. Shared by submit and confirm."""
    doc_format = _format_from_path(saved_path)
    pre_audit = _audit_by_extension(
        saved_path,
        list_indent_in=opts["list_indent_in"],
        list_level_indents=opts["list_level_indents"],
        para_indent_in=opts["para_indent_in"],
        first_line_indent_in=opts["first_line_indent_in"],
    )
    pre_body_font_pt = _estimate_pre_fix_body_font_pt(saved_path)

    ext = saved_path.suffix.lower()
    output_name = saved_path.stem + "-fixed" + ext
    output_path = saved_path.parent / output_name

    # Gate AI heading-detection sessions
    if opts.get("use_ai"):
        try:
            ctx = ai_gate()
            ctx.__enter__()
            _ai_ctx = ctx
        except GatingError:
            return _busy_response(
                "AI heading detection",
                back_url=url_for("fix.fix_form"),
            )
    else:
        _ai_ctx = None

    try:
        fix_result_tuple = _fix_by_extension(
            saved_path,
            output_path,
            bound=opts["bound"],
            list_indent_in=opts["list_indent_in"],
            list_hanging_in=opts["list_hanging_in"],
            list_level_indents=opts["list_level_indents"],
            para_indent_in=opts["para_indent_in"],
            first_line_indent_in=opts["first_line_indent_in"],
            preserve_heading_alignment=opts["preserve_heading_alignment"],
            detect_headings=opts["detect_headings"],
            use_ai=opts["use_ai"],
            heading_threshold=opts["heading_threshold"],
            confirmed_headings=confirmed_headings,
            heading_accuracy=opts["heading_accuracy"],
        )
    finally:
        if _ai_ctx is not None:
            _ai_ctx.__exit__(None, None, None)
    # Unpack; _fix_by_extension may return 5 or 6 elements (with ai_meta)
    if len(fix_result_tuple) == 6:
        _, total_fixes, fix_records, post_audit, warnings, ai_meta = fix_result_tuple
    else:
        _, total_fixes, fix_records, post_audit, warnings = fix_result_tuple
        ai_meta = {}
    ai_used = ai_meta.get("ai_used", False)

    warnings = list(warnings)

    # Suppress findings for rules intentionally bypassed by user settings.
    suppressed_rules: list[str] = []
    _rule_policy = opts.get("rule_policy")
    suppress_rule_ids: set[str] = set(_rule_policy.suppressed) if _rule_policy is not None else set()

    # Special case: when heading detection is off, always suppress faux-heading
    # findings since the fixer did not attempt to convert them.
    if not opts["detect_headings"]:
        suppress_rule_ids.add("ACB-FAUX-HEADING")

    if suppress_rule_ids:
        suppressed_rules.extend(sorted(suppress_rule_ids))

    filtered_findings = []
    heading_alignment_suppressed = False
    for finding in post_audit.findings:
        if finding.rule_id in suppress_rule_ids:
            continue
        if (
            opts["preserve_heading_alignment"]
            and _is_heading_alignment_finding(finding)
        ):
            heading_alignment_suppressed = True
            continue
        filtered_findings.append(finding)

    post_audit.findings = filtered_findings
    if heading_alignment_suppressed:
        suppressed_rules.append("ACB-ALIGNMENT (headings)")

    if pre_body_font_pt is not None and pre_body_font_pt < 17.5:
        warnings.insert(
            0,
            "This document appears to use body text around "
            f"{pre_body_font_pt:g}pt. Fixing to the ACB minimum of 18pt can increase "
            "page count, especially in long newsletters. Review margins and layout "
            "after download if page growth affects print production.",
        )

    applied_heading_fixes = sum(
        1 for rec in fix_records if rec.rule_id == "ACB-FAUX-HEADING"
    )

    # Compute document availability expiry for UI display
    from ..upload import get_upload_expiry
    import os as _os
    _max_age = int(_os.environ.get("UPLOAD_MAX_AGE_HOURS", "1"))
    doc_expiry = get_upload_expiry(token, max_age_hours=_max_age)

    if confirmed_headings_received is not None:
        current_app.logger.info(
            "Heading review apply summary | token=%s | file=%s | received=%d | applied=%d",
            token,
            saved_path.name,
            confirmed_headings_received,
            applied_heading_fixes,
        )

    categories = request.form.getlist("category") or ["acb", "msac"]
    standards_profile = request.form.get("standards_profile", "acb_2025")
    profile_label = get_profile_label(standards_profile)
    category_rule_ids = get_rule_ids_by_category(*categories)
    format_rule_ids = get_rule_ids_by_format(doc_format)
    profile_rule_ids = get_rule_ids_by_profile(standards_profile)
    mode = opts["mode"]

    if mode == "essentials":
        selected = get_rule_ids_by_severity("Critical", "High")
        mode_label = "Essentials Fix -- Critical and High focus"
    elif mode == "custom":
        selected = set(request.form.getlist("rule"))
        if not selected:
            selected = get_all_rule_ids()
        mode_label = f"Custom Fix -- {len(selected)} rules selected"
    else:
        selected = get_all_rule_ids()
        mode_label = "Full Fix -- all rules"

    selected = selected & category_rule_ids & format_rule_ids & profile_rule_ids
    post_findings = filter_findings(post_audit.findings, selected)
    pre_breakdown = _build_penalty_breakdown(pre_audit.findings)
    post_breakdown = _build_penalty_breakdown(post_audit.findings)
    score_delta = post_audit.score - pre_audit.score

    # Detect and warn about customizations from ACB defaults
    has_customizations, customization_reasons = detect_fix_customizations(opts)
    customization_warning = ""
    if has_customizations:
        customization_warning = generate_customization_warning(customization_reasons)

    return render_template(
        "fix_result.html",
        pre_score=pre_audit.score,
        pre_grade=pre_audit.grade,
        pre_finding_ids=",".join(f.rule_id for f in pre_audit.findings),
        post_score=post_audit.score,
        post_grade=post_audit.grade,
        score_delta=score_delta,
        pre_breakdown=pre_breakdown,
        post_breakdown=post_breakdown,
        confirmed_headings_received=confirmed_headings_received,
        applied_heading_fixes=applied_heading_fixes,
        total_fixes=total_fixes,
        fix_records=fix_records,
        remaining=len(post_audit.findings),
        post_findings=post_findings,
        warnings=warnings,
        suppressed_rules=suppressed_rules,
        mode_label=mode_label,
        profile_label=profile_label,
        download_name=output_name,
        token=token,
        ai_used=ai_used,
        doc_expiry=doc_expiry,
        customization_warning=customization_warning,
    )


@fix_bp.route("/", methods=["POST"])
def fix_submit():
    token = None
    try:
        from ..tool_usage import record as _record_usage
        _record_usage("fix")
        token, saved_path = validate_upload(request.files.get("document"))
        opts = _parse_form_options(request.form)
        ext = saved_path.suffix.lower()

        # Interactive heading review: detect candidates first for .docx
        if opts["detect_headings"] and ext == ".docx":
            from docx import Document as _Document
            from acb_large_print.heading_detector import detect_headings as _detect

            ai_provider = None
            if opts["use_ai"]:
                try:
                    from acb_large_print.ai_provider import get_provider

                    ai_provider = get_provider()
                except Exception:
                    pass

            doc = _Document(str(saved_path))
            candidates = _detect(
                doc,
                ai_provider=ai_provider,
                threshold=opts["heading_threshold"],
            )

            allowed_levels = opts["allowed_heading_levels"]
            for candidate in candidates:
                candidate.suggested_level = _closest_allowed_level(
                    candidate.suggested_level,
                    allowed_levels,
                )

            if candidates:
                # Build hidden form fields to preserve all original options
                form_fields = {}
                for key in request.form:
                    values = request.form.getlist(key)
                    if len(values) == 1:
                        form_fields[key] = values[0]
                    else:
                        form_fields[key] = values
                from ..upload import get_upload_expiry
                import os as _os
                _max_age = int(_os.environ.get("UPLOAD_MAX_AGE_HOURS", "1"))
                doc_expiry = get_upload_expiry(token, max_age_hours=_max_age)
                return render_template(
                    "fix_review_headings.html",
                    candidates=candidates,
                    allowed_heading_levels=allowed_levels,
                    token=token,
                    filename=saved_path.name,
                    form_fields=form_fields,
                    ai_used=ai_provider is not None,
                    doc_expiry=doc_expiry,
                )

        return _run_fix_and_render(saved_path, token, opts)

    except UploadError as e:
        from ..ai_features import ai_heading_fix_enabled

        if token:
            cleanup_token(token)

        return (
            render_template(
                "fix_form.html",
                error=str(e),
                ai_available=ai_heading_fix_enabled(),
                rules_by_format=get_rules_by_format(),
            ),
            400,
        )
    except Exception:
        from ..ai_features import ai_heading_fix_enabled

        if token:
            cleanup_token(token)

        return (
            render_template(
                "fix_form.html",
                error="An error occurred while fixing the document. "
                "Please ensure it is a valid Office file and try again.",
                ai_available=ai_heading_fix_enabled(),
                rules_by_format=get_rules_by_format(),
            ),
            500,
        )
    # Note: do NOT clean up here -- the temp dir must survive for the download


@fix_bp.route("/confirm", methods=["POST"])
def fix_confirm():
    """Apply fixes with user-confirmed heading selections."""
    token = request.form.get("token", "")
    temp_dir = get_temp_dir(token)
    if temp_dir is None:
        from ..ai_features import ai_heading_fix_enabled

        return (
            render_template(
                "fix_form.html",
                error="Your session has expired or the uploaded file is no longer available. "
                      "Documents are kept for up to 1 hour after upload. "
                      "Please upload and fix the document again.",
                ai_available=ai_heading_fix_enabled(),
                rules_by_format=get_rules_by_format(),
            ),
            400,
        )

    # Find the uploaded file in the temp dir
    saved_path = None
    for f in temp_dir.iterdir():
        if f.is_file() and f.suffix.lower() in {
            ".docx",
            ".xlsx",
            ".pptx",
            ".md",
            ".pdf",
            ".epub",
        }:
            saved_path = f
            break
    if saved_path is None:
        cleanup_token(token)
        from ..ai_features import ai_heading_fix_enabled

        return (
            render_template(
                "fix_form.html",
                error="Uploaded file not found. Please upload and fix the document again.",
                ai_available=ai_heading_fix_enabled(),
                rules_by_format=get_rules_by_format(),
            ),
            400,
        )

    try:
        opts = _parse_form_options(request.form)
        allowed_levels = set(opts["allowed_heading_levels"])

        # Parse confirmed headings from the review form
        confirmed_headings = []
        count = int(request.form.get("candidate_count", "0"))
        for i in range(count):
            level = request.form.get(f"level_{i}")
            if level == "skip":
                continue
            try:
                level_int = int(level)
            except (ValueError, TypeError):
                continue
            if not 1 <= level_int <= 6:
                continue
            if level_int not in allowed_levels:
                continue
            idx = int(request.form.get(f"idx_{i}", "-1"))
            text = request.form.get(f"text_{i}", "")
            if idx >= 0 and text:
                confirmed_headings.append((idx, level_int, text))

        return _run_fix_and_render(
            saved_path,
            token,
            opts,
            confirmed_headings=confirmed_headings if confirmed_headings else None,
            confirmed_headings_received=len(confirmed_headings),
        )

    except Exception:
        cleanup_token(token)
        from ..ai_features import ai_heading_fix_enabled

        return (
            render_template(
                "fix_form.html",
                error="An error occurred while fixing the document. "
                "Please try again.",
                ai_available=ai_heading_fix_enabled(),
            ),
            500,
        )


@fix_bp.route("/download", methods=["POST"])
def fix_download():
    """Serve the fixed file and clean up."""
    token = request.form.get("token", "")
    raw_name = request.form.get("download_name", "fixed.docx")
    download_name = secure_filename(raw_name) or "fixed.docx"

    temp_dir = get_temp_dir(token)
    if temp_dir is None:
        return (
            render_template(
                "error.html",
                title="Download Failed",
                message="The fixed file is no longer available. Please run the fix again.",
            ),
            404,
        )

    file_path = temp_dir / download_name
    # Ensure resolved path stays inside the temp directory (prevent traversal)
    try:
        file_path.resolve().relative_to(temp_dir.resolve())
    except ValueError:
        cleanup_token(token)
        return (
            render_template(
                "error.html",
                title="Download Failed",
                message="Invalid file name.",
            ),
            400,
        )

    if not file_path.exists():
        cleanup_token(token)
        return (
            render_template(
                "error.html",
                title="Download Failed",
                message="The fixed file is no longer available. Please run the fix again.",
            ),
            404,
        )

    response = send_file(
        str(file_path),
        as_attachment=True,
        download_name=download_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    @response.call_on_close
    def _cleanup():
        cleanup_token(token)

    return response


def _find_fixable_file(temp_dir):
    """Return the first fixable file in a temp dir, or None."""
    from ..upload import ALLOWED_EXTENSIONS
    for f in sorted(temp_dir.iterdir()):
        if f.is_file() and f.suffix.lower() in ALLOWED_EXTENSIONS:
            return f
    return None


@fix_bp.route("/from-audit/<token>", methods=["GET"])
def fix_from_audit_form(token: str):
    """Show the fix form pre-filled with a file from a completed audit session."""
    from ..ai_features import ai_heading_fix_enabled
    from flask import redirect
    temp_dir = get_temp_dir(token)
    if temp_dir is None:
        return redirect(url_for("fix.fix_form", notice="session_expired"))
    prefill_file = _find_fixable_file(temp_dir)
    if prefill_file is None:
        return redirect(url_for("fix.fix_form"))
    return render_template(
        "fix_form.html",
        ai_available=ai_heading_fix_enabled(),
        rules_by_format=get_rules_by_format(),
        prefill_token=token,
        prefill_filename=prefill_file.name,
    )


@fix_bp.route("/from-audit/<token>", methods=["POST"])
def fix_from_audit_submit(token: str):
    """Run fix using an existing file from a completed audit session (no re-upload)."""
    from ..ai_features import ai_heading_fix_enabled
    temp_dir = get_temp_dir(token)
    if temp_dir is None:
        return (
            render_template(
                "fix_form.html",
                error="Your session has expired. Please upload your document again.",
                ai_available=ai_heading_fix_enabled(),
                rules_by_format=get_rules_by_format(),
            ),
            400,
        )

    saved_path = _find_fixable_file(temp_dir)
    if saved_path is None:
        return (
            render_template(
                "fix_form.html",
                error="Original document not found. Please upload again.",
                ai_available=ai_heading_fix_enabled(),
                rules_by_format=get_rules_by_format(),
            ),
            400,
        )

    try:
        opts = _parse_form_options(request.form)
        ext = saved_path.suffix.lower()

        # Interactive heading review: same path as fix_submit
        if opts["detect_headings"] and ext == ".docx":
            from docx import Document as _Document
            from acb_large_print.heading_detector import detect_headings as _detect

            ai_provider = None
            if opts["use_ai"]:
                try:
                    from acb_large_print.ai_provider import get_provider
                    ai_provider = get_provider()
                except Exception:
                    pass

            doc = _Document(str(saved_path))
            candidates = _detect(doc, ai_provider=ai_provider, threshold=opts["heading_threshold"])
            allowed_levels = opts["allowed_heading_levels"]
            for candidate in candidates:
                candidate.suggested_level = _closest_allowed_level(candidate.suggested_level, allowed_levels)

            if candidates:
                form_fields = {}
                for key in request.form:
                    values = request.form.getlist(key)
                    form_fields[key] = values[0] if len(values) == 1 else values
                from ..upload import get_upload_expiry
                import os as _os
                _max_age = int(_os.environ.get("UPLOAD_MAX_AGE_HOURS", "1"))
                doc_expiry = get_upload_expiry(token, max_age_hours=_max_age)
                return render_template(
                    "fix_review_headings.html",
                    candidates=candidates,
                    allowed_heading_levels=allowed_levels,
                    token=token,
                    filename=saved_path.name,
                    form_fields=form_fields,
                    ai_used=ai_provider is not None,
                    doc_expiry=doc_expiry,
                )

        return _run_fix_and_render(saved_path, token, opts)

    except Exception:
        # Preserve prefill context so the user keeps their session reference on error
        return (
            render_template(
                "fix_form.html",
                error="An error occurred while fixing the document. Please try again.",
                ai_available=ai_heading_fix_enabled(),
                rules_by_format=get_rules_by_format(),
                prefill_token=token if get_temp_dir(token) is not None else None,
                prefill_filename=saved_path.name,
            ),
            500,
        )
    # Note: do NOT clean up here -- token stays alive for download
