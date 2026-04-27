"""Agentic Q&A handler with accessibility-focused tool calling for document interrogation.

Users ask questions about uploaded documents. Llama 3 decides which tools to use
based on the question. Tools are organized into four named agent categories:

  Compliance Agent  -- runs the GLOW audit engine; returns scores, findings, fixability
  Structure Agent   -- analyses heading hierarchy, lists, reading order, faux headings
  Content Agent     -- checks emphasis patterns, link text, reading level, alignment
  Remediation Agent -- explains rules, suggests fixes, prioritises and estimates impact

Tool calling flow:
  1. User asks question
  2. Llama decides which agent tools to invoke (or answers directly)
  3. Tools execute against the document (engine or heuristic)
  4. Llama uses tool results to form a grounded, actionable answer
  5. Answer + tool calls are stored in conversation history
  6. History can be exported as Markdown, Word, or PDF

Original document tools (11):
  extract_table, find_section, get_section_content, search_text, get_document_stats,
  get_document_summary, get_decisions_and_actions, summarize_section, list_headings,
  get_images, get_what_passes

Compliance Agent tools (4):
  run_accessibility_audit, get_compliance_score,
  get_critical_findings, get_auto_fixable_findings

Structure Agent tools (4):
  check_heading_hierarchy, find_faux_headings,
  check_list_structure, estimate_reading_order

Content Agent tools (4):
  check_emphasis_patterns, check_link_text,
  check_reading_level, check_alignment_hints

Remediation Agent tools (5):
  explain_rule, suggest_fix, prioritize_findings,
  estimate_fix_impact, check_image_alt_text
"""

from __future__ import annotations

import logging
import re
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

log = logging.getLogger(__name__)

# Tool category labels used in UI and system prompt
CATEGORY_DOCUMENT = "Document"
CATEGORY_COMPLIANCE = "Compliance Agent"
CATEGORY_STRUCTURE = "Structure Agent"
CATEGORY_CONTENT = "Content Agent"
CATEGORY_REMEDIATION = "Remediation Agent"

# ACB rule quick-reference for the Remediation Agent (subset for text-based analysis)
_RULE_EXPLANATIONS: dict[str, dict[str, str]] = {
    "ACB-FONT-FAMILY": {
        "plain": "All text must use Arial font. Other fonts (Times New Roman, Calibri, etc.) are not permitted.",
        "why": "Arial is a sans-serif font proven to be more readable for people with low vision.",
        "fix": "Select all text (Ctrl+A) and change the font to Arial.",
    },
    "ACB-FONT-SIZE-BODY": {
        "plain": "Body text must be at least 18pt. Smaller text is a critical violation.",
        "why": "18pt is the minimum size that most people with moderate low vision can read without magnification.",
        "fix": "Select all body paragraphs and set font size to 18pt.",
    },
    "ACB-FONT-SIZE-H1": {
        "plain": "Heading 1 must be 22pt.",
        "why": "Headings must be visually distinct from body text to support navigation.",
        "fix": "Apply Heading 1 style and set to 22pt.",
    },
    "ACB-FONT-SIZE-H2": {
        "plain": "Heading 2 must be 20pt.",
        "why": "Each heading level should be visually distinct and larger than body text.",
        "fix": "Apply Heading 2 style and set to 20pt.",
    },
    "ACB-NO-ITALIC": {
        "plain": "Italic formatting is completely prohibited. This is a critical violation.",
        "why": "Italic text is harder to read for people with dyslexia and low vision. ACB requires underline for emphasis instead.",
        "fix": "Remove all italic formatting. Use underline (Ctrl+U) for emphasis instead.",
    },
    "ACB-BOLD-HEADINGS-ONLY": {
        "plain": "Bold should not be used for emphasis in body text. Use underline instead.",
        "why": "Bold text can be hard to distinguish from normal weight text for some low-vision readers. ACB reserves bold for headings only.",
        "fix": "Replace bold emphasis in body paragraphs with underline formatting.",
    },
    "ACB-ALIGNMENT": {
        "plain": "All text must be flush left (left-aligned). Justified, centered, and right-aligned text are violations.",
        "why": "Justified text creates uneven word spacing ('rivers of white space') that makes reading harder for people with low vision and dyslexia.",
        "fix": "Select all text and apply Left alignment (Ctrl+L).",
    },
    "ACB-LINE-SPACING": {
        "plain": "Line spacing must be 1.5x (150%). Single-spaced text is a violation.",
        "why": "Increased line spacing reduces crowding and makes it easier to track lines while reading.",
        "fix": "Select all text, open Paragraph formatting, and set line spacing to 1.5.",
    },
    "ACB-MARGINS": {
        "plain": "All page margins must be 1 inch. Narrower margins reduce readability.",
        "why": "1-inch margins provide adequate white space and prevent text from feeling cramped.",
        "fix": "Go to Layout > Margins and select Normal (1 inch on all sides).",
    },
    "ACB-NO-HYPHENATION": {
        "plain": "Automatic hyphenation must be disabled.",
        "why": "Hyphenated words at line breaks are harder to read, especially for screen reader users and people with cognitive disabilities.",
        "fix": "Go to Layout > Hyphenation and turn it off.",
    },
    "ACB-HEADING-HIERARCHY": {
        "plain": "Heading levels must not skip. For example, you cannot go from Heading 1 directly to Heading 3.",
        "why": "Screen readers and navigation tools use heading hierarchy to help users jump between sections. Skipped levels break navigation.",
        "fix": "Review your heading structure and ensure each level follows the previous without gaps.",
    },
    "ACB-FAUX-HEADING": {
        "plain": "A paragraph formatted as bold text appears to be acting as a heading but uses normal body text style.",
        "why": "Screen readers announce true headings so users can navigate. Bold-formatted 'fake' headings are invisible to assistive technology.",
        "fix": "Apply the appropriate Word Heading style (Heading 1, 2, 3) instead of just making text bold.",
    },
    "ACB-LINK-TEXT": {
        "plain": "Hyperlinks should have descriptive text, not bare URLs or phrases like 'click here' or 'here'.",
        "why": "Screen reader users often navigate by links. 'Click here' is meaningless out of context. Bare URLs are not human-readable.",
        "fix": "Select the link text and replace it with a description of where the link goes (e.g., 'ACB Large Print Guidelines').",
    },
    "ACB-MISSING-ALT-TEXT": {
        "plain": "Images must have alternative text that describes their content.",
        "why": "Screen reader users cannot see images. Alt text is the only way they can access visual information.",
        "fix": "Right-click each image, select 'Edit Alt Text', and write a meaningful description.",
    },
    "ACB-DOC-TITLE": {
        "plain": "The document must have a title set in its properties.",
        "why": "Screen readers announce the document title when opening a file. Missing titles leave blind users without context.",
        "fix": "Go to File > Properties and enter a descriptive title.",
    },
    "ACB-DOC-LANGUAGE": {
        "plain": "The document language must be set (English or appropriate language).",
        "why": "Screen readers use the document language to select the correct voice and pronunciation rules.",
        "fix": "Go to Review > Language > Set Proofing Language and select the correct language.",
    },
    "ACB-TABLE-HEADERS": {
        "plain": "Tables must have defined header rows.",
        "why": "Screen readers announce column headers for each data cell in a table. Without headers, table data is meaningless to blind users.",
        "fix": "Click in the first row of the table, go to Table Design, and check 'Header Row'.",
    },
    "ACB-WIDOW-ORPHAN": {
        "plain": "Widow and orphan control should be enabled to prevent isolated lines at page tops and bottoms.",
        "why": "Single lines stranded at page boundaries are disorienting and break reading flow.",
        "fix": "Select all paragraphs, open Paragraph > Line and Page Breaks, and check Widow/Orphan control.",
    },
}


class DocumentContext:
    """Holds extracted document content and metadata for Q&A."""

    def __init__(
        self,
        text: str,
        filename: str,
        doc_path: Path | None = None,
        doc_type: str = "text",
    ) -> None:
        """Initialize document context.

        Args:
            text: Full extracted text from document.
            filename: Original filename.
            doc_path: Path to the original file (used by Compliance Agent).
            doc_type: "text" (Word/PDF/Markdown), "image" (scanned), or "mixed".
        """
        self.text = text
        self.filename = filename
        self.doc_path = doc_path  # None when not available; Compliance Agent checks
        self.doc_type = doc_type
        self.headings = self._extract_headings()
        self.tables = self._extract_tables()
        self.stats = self._compute_stats()
        self._audit_cache: dict[str, Any] | None = None  # Cache live audit results

    def _extract_headings(self) -> list[dict[str, Any]]:
        """Extract heading hierarchy and positions from text."""
        headings = []
        lines = self.text.split("\n")
        for i, line in enumerate(lines):
            if line.startswith("# "):
                headings.append({"level": 1, "text": line[2:].strip(), "line": i})
            elif line.startswith("## "):
                headings.append({"level": 2, "text": line[3:].strip(), "line": i})
            elif line.startswith("### "):
                headings.append({"level": 3, "text": line[4:].strip(), "line": i})
            elif line.startswith("#### "):
                headings.append({"level": 4, "text": line[5:].strip(), "line": i})
        return headings

    def _extract_tables(self) -> list[dict[str, Any]]:
        """Extract simple table-like structures (pipe-delimited markdown tables)."""
        tables = []
        lines = self.text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            if "|" in line and i + 1 < len(lines) and "|" in lines[i + 1]:
                # Likely a markdown table
                table_lines = [line]
                i += 1
                table_lines.append(lines[i])  # header separator
                i += 1
                while i < len(lines) and "|" in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                tables.append({"lines": table_lines, "start": len(tables)})
            else:
                i += 1
        return tables

    def _compute_stats(self) -> dict[str, Any]:
        """Compute document statistics."""
        words = len(self.text.split())
        lines = len(self.text.split("\n"))
        chars = len(self.text)
        reading_time_minutes = max(1, words // 200)  # ~200 wpm
        return {
            "words": words,
            "lines": lines,
            "characters": chars,
            "reading_time_minutes": reading_time_minutes,
            "filename": self.filename,
        }


class ToolRegistry:
    """Registry of callable tools for the agent."""

    def __init__(self, context: DocumentContext) -> None:
        self.context = context

    def extract_table(self, table_name: str) -> str:
        """Extract a specific table by name or index."""
        try:
            idx = int(table_name)
            if 0 <= idx < len(self.context.tables):
                return "\n".join(self.context.tables[idx]["lines"])
        except (ValueError, IndexError):
            pass
        return f"Table '{table_name}' not found. Available tables: {len(self.context.tables)}"

    def find_section(self, section_name: str) -> str:
        """Find and return a section by heading name."""
        for heading in self.context.headings:
            if section_name.lower() in heading["text"].lower():
                # Find text until next heading
                lines = self.context.text.split("\n")
                start = heading["line"]
                end = len(lines)
                for h in self.context.headings:
                    if h["line"] > start:
                        end = h["line"]
                        break
                return "\n".join(lines[start:end]).strip()
        return f"Section '{section_name}' not found."

    def search_text(self, keyword: str) -> str:
        """Search for keyword in document and return context."""
        if not keyword:
            return "No keyword provided."
        matches = []
        lines = self.context.text.split("\n")
        for i, line in enumerate(lines):
            if keyword.lower() in line.lower():
                matches.append(f"Line {i + 1}: {line}")
        if not matches:
            return f"No matches found for '{keyword}'."
        return "\n".join(matches[:10])  # Return first 10 matches

    def get_document_stats(self) -> str:
        """Return document statistics."""
        stats = self.context.stats
        return (
            f"Document: {stats['filename']}\n"
            f"Words: {stats['words']}\n"
            f"Lines: {stats['lines']}\n"
            f"Characters: {stats['characters']}\n"
            f"Reading time: ~{stats['reading_time_minutes']} minutes\n"
            f"Headings: {len(self.context.headings)}\n"
            f"Tables: {len(self.context.tables)}"
        )

    def summarize_section(self, section_name: str) -> str:
        """Summarize a section (via find_section + message to Llama)."""
        section_text = self.find_section(section_name)
        if "not found" in section_text:
            return section_text
        # Summaries are done by Llama, this just returns the section
        return f"Section '{section_name}':\n\n{section_text[:500]}..."

    def list_headings(self) -> str:
        """List all headings in the document."""
        if not self.context.headings:
            return "No headings found."
        headings_list = []
        for h in self.context.headings:
            indent = "  " * (h["level"] - 1)
            headings_list.append(f"{indent}• {h['text']}")
        return "\n".join(headings_list)

    def get_images(self) -> str:
        """Return list of images embedded in the document."""
        return (
            "Inline image extraction is not available in this version. "
            "If image alt text or figure captions are important, use the Audit tab "
            "to check for missing alt text findings."
        )

    def get_document_summary(self) -> str:
        """Return a plain-language overview of what the document is and what it covers."""
        text = self.context.text
        headings = self.context.headings
        stats = self.context.stats

        # Derive top-level title and top-level sections
        title = headings[0]["text"] if headings else self.context.filename
        h2s = [h["text"] for h in headings if h["level"] == 2][:6]
        h2_list = ", ".join(h2s) if h2s else "(no top-level sections detected)"

        # Detect document type keywords
        lower = text.lower()
        doc_type_hints: list[str] = []
        if any(kw in lower for kw in ("agenda", "minutes", "action item", "attendance")):
            doc_type_hints.append("meeting document (agenda or minutes)")
        if any(kw in lower for kw in ("budget", "revenue", "expenses", "reserve", "grant")):
            doc_type_hints.append("contains financial data")
        if any(kw in lower for kw in ("policy", "guideline", "procedure", "standard")):
            doc_type_hints.append("policy or procedure document")
        if any(kw in lower for kw in ("report", "findings", "recommendation")):
            doc_type_hints.append("report with findings or recommendations")
        type_str = "; ".join(doc_type_hints) if doc_type_hints else "general document"

        return (
            f"Document: {title}\n"
            f"Type: {type_str}\n"
            f"Length: {stats['words']} words across {len(headings)} headings\n"
            f"Main sections: {h2_list}\n"
            f"Tables: {len(self.context.tables)}"
        )

    def get_section_content(self, section_name: str) -> str:
        """Return the full text content of a named section (not just its existence)."""
        if not section_name:
            return "Please provide a section name."
        for heading in self.context.headings:
            if section_name.lower() in heading["text"].lower():
                lines = self.context.text.split("\n")
                start = heading["line"]
                end = len(lines)
                for h in self.context.headings:
                    if h["line"] > start:
                        end = h["line"]
                        break
                content = "\n".join(lines[start:end]).strip()
                # Cap at 1500 chars so it doesn't swamp the context
                if len(content) > 1500:
                    content = content[:1500] + "\n...[section truncated]"
                return f"Section: {heading['text']}\n\n{content}"
        # Fuzzy fallback: search for heading-like text anywhere
        return f"Section '{section_name}' not found. Available sections: " + ", ".join(
            f"\"{h['text']}\"" for h in self.context.headings[:12]
        )

    def get_decisions_and_actions(self) -> str:
        """Extract voted decisions, motions, and action items from meeting documents."""
        text = self.context.text
        lines = text.split("\n")
        decisions: list[str] = []
        actions: list[str] = []

        for line in lines:
            stripped = line.strip()
            low = stripped.lower()
            # Voted motions / resolutions
            if any(kw in low for kw in (
                "voted", "motion", "resolved", "unanimously", "seconded",
                "approved", "passed", "agreed", "board approved", "board voted",
            )):
                if len(stripped) > 15:
                    decisions.append(f"  • {stripped}")
            # Action items: assigned tasks / owners / deadlines
            elif any(kw in low for kw in (
                "action item", "action:", "will ", "assigned to", "responsible",
                "by q", "by april", "by may", "by june", "by july",
                "deadline", "owner:", "follow up", "follow-up",
            )):
                if len(stripped) > 10:
                    actions.append(f"  • {stripped}")

        parts: list[str] = []
        if decisions:
            parts.append(f"Decisions / motions ({len(decisions)} found):")
            parts.extend(decisions[:10])
        else:
            parts.append("No voted motions or decisions detected.")

        if actions:
            parts.append(f"\nAction items ({len(actions)} found):")
            parts.extend(actions[:10])
        else:
            parts.append("\nNo explicit action items detected.")

        if not decisions and not actions:
            parts.append(
                "\nTip: this tool is most useful for meeting minutes and board documents."
            )
        return "\n".join(parts)

    def get_what_passes(self) -> str:
        """Return what the document already does correctly to balance the audit picture."""
        context = self.context
        passes: list[str] = []

        # Heading hierarchy
        headings = context.headings
        if headings:
            has_skips = False
            prev = 0
            for h in headings:
                if h["level"] - prev > 1 and prev != 0:
                    has_skips = True
                    break
                prev = h["level"]
            if not has_skips:
                passes.append(
                    f"Heading hierarchy: PASS — {len(headings)} headings, no skipped levels."
                )

        # No italic
        text = context.text
        italic_matches = re.findall(r"(?<!\*)\*[^*\n]{3,60}\*(?!\*)|_[^_\n]{3,60}_", text)
        if not italic_matches:
            passes.append("Italic usage: PASS — no italic text detected.")

        # Link text
        bare_urls = re.findall(r"(?<!\()(https?://\S{10,})", text)
        generic = re.findall(
            r"\[(click here|here|read more|more|learn more|this link|link)\]\(",
            text, re.IGNORECASE,
        )
        if not bare_urls and not generic:
            passes.append("Link text: PASS — no bare URLs or generic link text detected.")

        # Alt text
        all_images = re.findall(r"!\[([^\]]*)\]\([^)]+\)", text)
        if all_images:
            missing = [img for img in all_images if not img.strip()]
            if not missing:
                passes.append(
                    f"Alt text: PASS — all {len(all_images)} image(s) have alt text."
                )
        else:
            passes.append("Alt text: N/A — no images in document.")

        # Faux headings
        faux = [line.strip() for line in text.split("\n")
                if not line.strip().startswith("#")
                and re.fullmatch(r"\*\*[^*]{3,60}\*\*", line.strip())]
        if not faux:
            passes.append("Faux headings: PASS — no bold-body-text fake headings detected.")

        # List structure
        bullet_lines = [l for l in text.split("\n") if re.match(r"^\s*[-*+] ", l)]
        deeply_nested = [l for l in bullet_lines if l.startswith("      ")]
        if bullet_lines and not deeply_nested:
            passes.append(
                f"List structure: PASS — {len(bullet_lines)} list items, no excessive nesting."
            )

        if not passes:
            return "No passing checks could be confirmed from plain-text analysis."
        return (
            f"What the document already does correctly ({len(passes)} checks passing):\n"
            + "\n".join(f"  {p}" for p in passes)
        )

    # ------------------------------------------------------------------
    # Compliance Agent tools
    # ------------------------------------------------------------------

    def run_accessibility_audit(self) -> str:
        """Run live GLOW audit engine and return a summary of all findings."""
        if self.context.doc_path is None:
            return self._heuristic_compliance_summary()
        try:
            from acb_large_print.auditor import audit_document

            result = audit_document(str(self.context.doc_path))
            self.context._audit_cache = {
                "findings": [
                    {
                        "rule_id": f.rule_id,
                        "severity": f.severity,
                        "description": f.description,
                        "auto_fixable": f.auto_fixable,
                    }
                    for f in result.findings
                ],
                "score": result.score,
            }
            total = len(result.findings)
            critical = sum(1 for f in result.findings if f.severity == "critical")
            high = sum(1 for f in result.findings if f.severity == "high")
            medium = sum(1 for f in result.findings if f.severity == "medium")
            low = sum(1 for f in result.findings if f.severity == "low")
            fixable = sum(1 for f in result.findings if f.auto_fixable)
            return (
                f"Audit complete. Score: {result.score}/100\n"
                f"Total findings: {total} "
                f"(Critical: {critical}, High: {high}, Medium: {medium}, Low: {low})\n"
                f"Auto-fixable: {fixable} of {total}\n"
                f"Use get_critical_findings or get_auto_fixable_findings for detail."
            )
        except Exception as e:
            log.warning("Live audit failed, falling back to heuristic: %s", e)
            return self._heuristic_compliance_summary()

    def _heuristic_compliance_summary(self) -> str:
        """Return a heuristic compliance overview from text content."""
        text = self.context.text
        issues: list[str] = []
        if re.search(r"\*\*[^*]{20,}\*\*", text):
            issues.append("• Possible bold-emphasis abuse (ACB-BOLD-HEADINGS-ONLY)")
        if re.search(r"_[^_]{5,}_|(?<!\*)\*[^*]{5,}\*(?!\*)", text):
            issues.append("• Possible italic usage (ACB-NO-ITALIC — critical violation)")
        if re.search(r"https?://\S{10,}", text):
            issues.append("• Possible bare URLs used as link text (ACB-LINK-TEXT)")
        heading_levels = [h["level"] for h in self.context.headings]
        for i in range(1, len(heading_levels)):
            if heading_levels[i] - heading_levels[i - 1] > 1:
                issues.append("• Skipped heading level detected (ACB-HEADING-HIERARCHY)")
                break
        if not issues:
            issues.append("No obvious issues detected in plain-text analysis.")
        return "Heuristic compliance summary (no live audit available):\n" + "\n".join(issues)

    def get_compliance_score(self) -> str:
        """Return the compliance score and severity distribution."""
        if self.context._audit_cache:
            cache = self.context._audit_cache
            findings = cache["findings"]
            score = cache["score"]
            c = sum(1 for f in findings if f["severity"] == "critical")
            h = sum(1 for f in findings if f["severity"] == "high")
            m = sum(1 for f in findings if f["severity"] == "medium")
            lo = sum(1 for f in findings if f["severity"] == "low")
            return (
                f"Compliance score: {score}/100\n"
                f"Critical: {c}  High: {h}  Medium: {m}  Low: {lo}"
            )
        # If no live audit yet, run one
        self.run_accessibility_audit()
        if self.context._audit_cache:
            return self.get_compliance_score()
        return "Score not available. Run run_accessibility_audit first."

    def get_critical_findings(self) -> str:
        """List critical and high severity findings."""
        if self.context._audit_cache is None:
            self.run_accessibility_audit()
        if self.context._audit_cache is None:
            return "No audit data available."
        findings = [
            f for f in self.context._audit_cache["findings"]
            if f["severity"] in ("critical", "high")
        ]
        if not findings:
            return "No critical or high findings. Document is in good shape."
        lines = [f"Critical/High findings ({len(findings)} total):"]
        for f in findings[:15]:
            fixable = "auto-fixable" if f["auto_fixable"] else "manual fix required"
            lines.append(f"  [{f['severity'].upper()}] {f['rule_id']}: {f['description']} ({fixable})")
        return "\n".join(lines)

    def get_auto_fixable_findings(self) -> str:
        """List findings that can be auto-fixed by GLOW."""
        if self.context._audit_cache is None:
            self.run_accessibility_audit()
        if self.context._audit_cache is None:
            return "No audit data available."
        findings = [f for f in self.context._audit_cache["findings"] if f["auto_fixable"]]
        if not findings:
            return "No auto-fixable findings detected."
        lines = [f"Auto-fixable findings ({len(findings)} total — upload to Fix to correct):"]
        for f in findings[:15]:
            lines.append(f"  {f['rule_id']}: {f['description']}")
        return "\n".join(lines)

    # ------------------------------------------------------------------
    # Structure Agent tools
    # ------------------------------------------------------------------

    def check_heading_hierarchy(self) -> str:
        """Detect heading hierarchy gaps and report skipped levels."""
        headings = self.context.headings
        if not headings:
            return "No headings found in document."
        issues: list[str] = []
        prev_level = 0
        for h in headings:
            if h["level"] - prev_level > 1 and prev_level != 0:
                issues.append(
                    f"  Skipped heading level: H{prev_level} → H{h['level']} "
                    f"at \"{h['text']}\""
                )
            prev_level = h["level"]
        if issues:
            return (
                f"Heading hierarchy issues found ({len(issues)}):\n"
                + "\n".join(issues)
                + "\nFix: do not skip heading levels (e.g. H1 → H3). Use consecutive levels."
            )
        return (
            f"Heading hierarchy is valid. {len(headings)} headings, no skipped levels."
        )

    def find_faux_headings(self) -> str:
        """Find paragraphs that look like headings but are not styled as headings."""
        lines = self.context.text.split("\n")
        candidates: list[str] = []
        for i, line in enumerate(lines):
            stripped = line.strip()
            # Short lines (< 80 chars), all-bold Markdown, NOT already a heading
            if (
                not stripped.startswith("#")
                and re.fullmatch(r"\*\*[^*]{3,60}\*\*", stripped)
            ):
                candidates.append(f"  Line {i + 1}: {stripped}")
        if not candidates:
            return (
                "No obvious faux headings found in Markdown analysis. "
                "For Word documents, upload to Fix for full faux-heading detection."
            )
        return (
            f"Possible faux headings ({len(candidates)} found):\n"
            + "\n".join(candidates[:10])
            + "\nFix: apply a proper heading style instead of bold body text."
        )

    def check_list_structure(self) -> str:
        """Inspect list structure for nesting and consistency."""
        text = self.context.text
        bullet_lines = [l for l in text.split("\n") if re.match(r"^\s*[-*+] ", l)]
        numbered_lines = [l for l in text.split("\n") if re.match(r"^\s*\d+\. ", l)]
        issues: list[str] = []
        # Check for deep nesting (> 2 levels)
        deeply_nested = [l for l in bullet_lines if l.startswith("      ")]
        if deeply_nested:
            issues.append(f"Deep list nesting detected ({len(deeply_nested)} items, >2 levels)")
        summary = (
            f"Lists: {len(bullet_lines)} bullet items, {len(numbered_lines)} numbered items."
        )
        if issues:
            return summary + "\nIssues:\n" + "\n".join(f"  • {i}" for i in issues)
        return summary + "\nNo structural list issues found."

    def estimate_reading_order(self) -> str:
        """Estimate reading-order risk by checking for multi-column or table-heavy layouts."""
        tables = self.context.tables
        text = self.context.text
        # Multi-column heuristic: pipe chars that aren't tables
        pipe_density = text.count("|") / max(len(text.split("\n")), 1)
        risks: list[str] = []
        if len(tables) > 5:
            risks.append(f"High table count ({len(tables)}) — confirm reading order in each table")
        if pipe_density > 3:
            risks.append("High pipe-character density — possible complex table or layout")
        if not risks:
            return (
                f"Reading order risk: LOW. "
                f"{len(tables)} table(s) detected. No complex layout signals."
            )
        return "Reading order risk: MEDIUM.\n" + "\n".join(f"  • {r}" for r in risks)

    # ------------------------------------------------------------------
    # Content Agent tools
    # ------------------------------------------------------------------

    def check_emphasis_patterns(self) -> str:
        """Check italic, bold, and underline emphasis for ACB compliance."""
        text = self.context.text
        issues: list[str] = []
        italic_matches = re.findall(r"(?<!\*)\*[^*\n]{3,60}\*(?!\*)|_[^_\n]{3,60}_", text)
        if italic_matches:
            issues.append(
                f"CRITICAL — {len(italic_matches)} italic instance(s) found "
                f"(ACB-NO-ITALIC). Example: {italic_matches[0][:60]}"
            )
        long_bold = re.findall(r"\*\*[^*\n]{40,}\*\*", text)
        if long_bold:
            issues.append(
                f"HIGH — {len(long_bold)} long bold phrase(s) that may be body emphasis "
                f"(ACB-BOLD-HEADINGS-ONLY). Bold is permitted for headings only."
            )
        if not issues:
            return (
                "Emphasis patterns: OK. No italic found. "
                "Bold usage appears limited to headings."
            )
        return "Emphasis issues found:\n" + "\n".join(f"  • {i}" for i in issues)

    def check_link_text(self) -> str:
        """Evaluate link text quality for ACB-LINK-TEXT compliance."""
        text = self.context.text
        issues: list[str] = []
        # Bare URLs
        bare_urls = re.findall(r"(?<!\()(https?://\S{10,})", text)
        if bare_urls:
            issues.append(
                f"Bare URLs ({len(bare_urls)} found): "
                + ", ".join(bare_urls[:3])
            )
        # Generic link text in Markdown [click here](...)
        generic = re.findall(
            r"\[(click here|here|read more|more|learn more|this link|link)\]\(",
            text,
            re.IGNORECASE,
        )
        if generic:
            issues.append(f"Generic link text ({len(generic)} found): {', '.join(generic[:5])}")
        if not issues:
            return "Link text: OK. No bare URLs or generic link text detected."
        return (
            "Link text issues (ACB-LINK-TEXT):\n"
            + "\n".join(f"  • {i}" for i in issues)
            + "\nFix: replace with descriptive text like 'ACB Large Print Guidelines'."
        )

    def check_reading_level(self) -> str:
        """Estimate reading level indicators (sentence and word length heuristics)."""
        text = self.context.text
        sentences = re.split(r"[.!?]+", text)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 10]
        if not sentences:
            return "Not enough text to estimate reading level."
        avg_words = sum(len(s.split()) for s in sentences) / len(sentences)
        long_sentences = [s for s in sentences if len(s.split()) > 30]
        words = text.split()
        long_words = [w for w in words if len(w) > 12]
        level = "LOW (plain language)" if avg_words < 18 else "MEDIUM" if avg_words < 25 else "HIGH (may be complex)"
        return (
            f"Reading level estimate: {level}\n"
            f"Average sentence length: {avg_words:.1f} words\n"
            f"Long sentences (>30 words): {len(long_sentences)}\n"
            f"Long words (>12 chars): {len(long_words)}\n"
            f"Total sentences analyzed: {len(sentences)}"
        )

    def check_alignment_hints(self) -> str:
        """Look for alignment signals in Markdown that suggest ACB-ALIGNMENT issues."""
        text = self.context.text
        issues: list[str] = []
        # HTML alignment overrides sometimes appear in Markdown
        center_tags = re.findall(r"<center>|style=[\"'][^\"']*text-align:\s*center", text, re.IGNORECASE)
        right_tags = re.findall(r"style=[\"'][^\"']*text-align:\s*right", text, re.IGNORECASE)
        if center_tags:
            issues.append(f"Center alignment override found ({len(center_tags)} instance(s)) — ACB requires flush-left")
        if right_tags:
            issues.append(f"Right alignment override found ({len(right_tags)} instance(s)) — ACB requires flush-left")
        if not issues:
            return "Alignment: No explicit alignment overrides detected in Markdown. Word documents should be verified with a full audit."
        return "Alignment issues (ACB-ALIGNMENT):\n" + "\n".join(f"  • {i}" for i in issues)

    # ------------------------------------------------------------------
    # Remediation Agent tools
    # ------------------------------------------------------------------

    def explain_rule(self, rule_id: str) -> str:
        """Explain an ACB rule in plain language with fix guidance."""
        rule_id = rule_id.upper().strip()
        entry = _RULE_EXPLANATIONS.get(rule_id)
        if entry:
            return (
                f"Rule: {rule_id}\n\n"
                f"What it means: {entry['plain']}\n\n"
                f"Why it matters: {entry['why']}\n\n"
                f"How to fix: {entry['fix']}"
            )
        return (
            f"Rule '{rule_id}' not found in quick reference. "
            f"Known rules: {', '.join(_RULE_EXPLANATIONS.keys())}. "
            f"For full rule details see the GLOW Guidelines page."
        )

    def suggest_fix(self, rule_id: str) -> str:
        """Return targeted fix instructions for a specific rule."""
        rule_id = rule_id.upper().strip()
        entry = _RULE_EXPLANATIONS.get(rule_id)
        if entry:
            return f"Fix for {rule_id}:\n{entry['fix']}"
        return f"No fix template found for '{rule_id}'. Use explain_rule to look it up."

    def prioritize_findings(self) -> str:
        """Rank findings by impact: critical first, then high, then auto-fixable."""
        if self.context._audit_cache is None:
            self.run_accessibility_audit()
        if self.context._audit_cache is None:
            # Fall back to heuristic ordering
            return (
                "Priority order (no live audit available):\n"
                "  1. Fix ACB-NO-ITALIC (critical — remove all italic)\n"
                "  2. Fix ACB-FAUX-HEADING (high — apply real heading styles)\n"
                "  3. Fix ACB-LINK-TEXT (high — replace generic/bare-URL links)\n"
                "  4. Fix ACB-FONT-SIZE-BODY (critical — raise body text to 18pt)\n"
                "  5. Fix ACB-MISSING-ALT-TEXT (high — add alt text to images)\n"
                "Upload to GLOW Fix for full auto-fix."
            )
        findings = self.context._audit_cache["findings"]
        order = {"critical": 0, "high": 1, "medium": 2, "low": 3}
        sorted_f = sorted(findings, key=lambda f: (order.get(f["severity"], 4), not f["auto_fixable"]))
        lines = ["Prioritized findings (by severity, auto-fixable first):"]
        for i, f in enumerate(sorted_f[:10], 1):
            tag = "AUTO-FIX" if f["auto_fixable"] else "MANUAL"
            lines.append(f"  {i}. [{f['severity'].upper()} / {tag}] {f['rule_id']}: {f['description']}")
        return "\n".join(lines)

    def estimate_fix_impact(self) -> str:
        """Estimate how much the score would improve after auto-fix."""
        if self.context._audit_cache is None:
            self.run_accessibility_audit()
        if self.context._audit_cache is None:
            return "Cannot estimate without audit data."
        findings = self.context._audit_cache["findings"]
        total = len(findings)
        fixable = sum(1 for f in findings if f["auto_fixable"])
        score = self.context._audit_cache.get("score", 0)
        if total == 0:
            return "No findings to fix. Document is already compliant."
        pct = fixable / total * 100
        return (
            f"Current score: {score}/100\n"
            f"Auto-fixable findings: {fixable} of {total} ({pct:.0f}%)\n"
            f"Uploading to GLOW Fix would automatically resolve these {fixable} issue(s).\n"
            f"Remaining {total - fixable} finding(s) require manual attention."
        )

    def check_image_alt_text(self) -> str:
        """Check for Markdown image alt text completeness."""
        text = self.context.text
        all_images = re.findall(r"!\[([^\]]*)\]\([^)]+\)", text)
        missing = [img for img in all_images if not img.strip()]
        if not all_images:
            return "No Markdown images found in document."
        if missing:
            return (
                f"Alt text issues (ACB-MISSING-ALT-TEXT):\n"
                f"  {len(missing)} of {len(all_images)} image(s) have empty alt text.\n"
                f"Fix: add a meaningful description inside the brackets: ![Description](image.png)"
            )
        return (
            f"Alt text: OK. All {len(all_images)} image(s) have alt text. "
            f"Verify descriptions are meaningful, not just filenames."
        )

    # ------------------------------------------------------------------
    # Tool registry: definitions + dispatcher
    # ------------------------------------------------------------------

    def get_all_tools(self) -> dict[str, dict[str, Any]]:
        """Return tool definitions for Llama tool calling, grouped by agent category."""
        return {
            # Document tools
            "extract_table": {
                "category": CATEGORY_DOCUMENT,
                "description": "Extract a specific table by name or 0-based index",
                "parameters": {"table_name": "str"},
            },
            "find_section": {
                "category": CATEGORY_DOCUMENT,
                "description": "Confirm whether a named section exists (use get_section_content for full text)",
                "parameters": {"section_name": "str"},
            },
            "get_section_content": {
                "category": CATEGORY_DOCUMENT,
                "description": "Return the full paragraph text of a named section",
                "parameters": {"section_name": "str"},
            },
            "search_text": {
                "category": CATEGORY_DOCUMENT,
                "description": "Search for keyword and return matching lines with line numbers",
                "parameters": {"keyword": "str"},
            },
            "get_document_stats": {
                "category": CATEGORY_DOCUMENT,
                "description": "Return word count, line count, headings, tables, reading time",
                "parameters": {},
            },
            "get_document_summary": {
                "category": CATEGORY_DOCUMENT,
                "description": "Return a plain-language overview of the document type, title, and main sections",
                "parameters": {},
            },
            "get_decisions_and_actions": {
                "category": CATEGORY_DOCUMENT,
                "description": "Extract voted motions, decisions, and action items from meeting or board documents",
                "parameters": {},
            },
            "summarize_section": {
                "category": CATEGORY_DOCUMENT,
                "description": "Return the text of a named section for the model to summarize",
                "parameters": {"section_name": "str"},
            },
            "list_headings": {
                "category": CATEGORY_DOCUMENT,
                "description": "List all headings with hierarchy indentation",
                "parameters": {},
            },
            "get_images": {
                "category": CATEGORY_DOCUMENT,
                "description": "List images (scanned PDF / vision model path)",
                "parameters": {},
            },
            "get_what_passes": {
                "category": CATEGORY_COMPLIANCE,
                "description": "Return what the document already does correctly (heading hierarchy, links, alt text, etc.)",
                "parameters": {},
            },
            # Compliance Agent
            "run_accessibility_audit": {
                "category": CATEGORY_COMPLIANCE,
                "description": "Run full GLOW accessibility audit and return severity summary",
                "parameters": {},
            },
            "get_compliance_score": {
                "category": CATEGORY_COMPLIANCE,
                "description": "Return the compliance score and severity distribution",
                "parameters": {},
            },
            "get_critical_findings": {
                "category": CATEGORY_COMPLIANCE,
                "description": "List critical and high severity findings with fix type",
                "parameters": {},
            },
            "get_auto_fixable_findings": {
                "category": CATEGORY_COMPLIANCE,
                "description": "List findings that GLOW Fix can automatically correct",
                "parameters": {},
            },
            # Structure Agent
            "check_heading_hierarchy": {
                "category": CATEGORY_STRUCTURE,
                "description": "Detect skipped heading levels (ACB-HEADING-HIERARCHY)",
                "parameters": {},
            },
            "find_faux_headings": {
                "category": CATEGORY_STRUCTURE,
                "description": "Find bold body paragraphs acting as headings (ACB-FAUX-HEADING)",
                "parameters": {},
            },
            "check_list_structure": {
                "category": CATEGORY_STRUCTURE,
                "description": "Inspect list nesting depth and consistency",
                "parameters": {},
            },
            "estimate_reading_order": {
                "category": CATEGORY_STRUCTURE,
                "description": "Estimate reading-order risk from table count and layout signals",
                "parameters": {},
            },
            # Content Agent
            "check_emphasis_patterns": {
                "category": CATEGORY_CONTENT,
                "description": "Detect italic and bold-abuse violations (ACB-NO-ITALIC, ACB-BOLD-HEADINGS-ONLY)",
                "parameters": {},
            },
            "check_link_text": {
                "category": CATEGORY_CONTENT,
                "description": "Evaluate link text quality for bare URLs and generic phrases (ACB-LINK-TEXT)",
                "parameters": {},
            },
            "check_reading_level": {
                "category": CATEGORY_CONTENT,
                "description": "Estimate sentence and word complexity (plain-language indicator)",
                "parameters": {},
            },
            "check_alignment_hints": {
                "category": CATEGORY_CONTENT,
                "description": "Detect center/right alignment overrides in Markdown (ACB-ALIGNMENT)",
                "parameters": {},
            },
            # Remediation Agent
            "explain_rule": {
                "category": CATEGORY_REMEDIATION,
                "description": "Explain an ACB rule (e.g. ACB-NO-ITALIC) in plain language with fix steps",
                "parameters": {"rule_id": "str (e.g. ACB-NO-ITALIC)"},
            },
            "suggest_fix": {
                "category": CATEGORY_REMEDIATION,
                "description": "Return targeted fix instructions for a rule ID",
                "parameters": {"rule_id": "str"},
            },
            "prioritize_findings": {
                "category": CATEGORY_REMEDIATION,
                "description": "Rank all findings by severity and auto-fixability",
                "parameters": {},
            },
            "estimate_fix_impact": {
                "category": CATEGORY_REMEDIATION,
                "description": "Estimate score improvement after auto-fix",
                "parameters": {},
            },
            "check_image_alt_text": {
                "category": CATEGORY_REMEDIATION,
                "description": "Check Markdown images for missing alt text (ACB-MISSING-ALT-TEXT)",
                "parameters": {},
            },
        }

    def call(self, tool_name: str, **kwargs: Any) -> str:
        """Dispatch a tool call by name and return its string result."""
        dispatch: dict[str, Any] = {
            # Document
            "extract_table": lambda: self.extract_table(kwargs.get("table_name", "")),
            "find_section": lambda: self.find_section(kwargs.get("section_name", "")),
            "get_section_content": lambda: self.get_section_content(kwargs.get("section_name", "")),
            "search_text": lambda: self.search_text(kwargs.get("keyword", "")),
            "get_document_stats": self.get_document_stats,
            "get_document_summary": self.get_document_summary,
            "get_decisions_and_actions": self.get_decisions_and_actions,
            "summarize_section": lambda: self.summarize_section(kwargs.get("section_name", "")),
            "list_headings": self.list_headings,
            "get_images": self.get_images,
            "get_what_passes": self.get_what_passes,
            # Compliance Agent
            "run_accessibility_audit": self.run_accessibility_audit,
            "get_compliance_score": self.get_compliance_score,
            "get_critical_findings": self.get_critical_findings,
            "get_auto_fixable_findings": self.get_auto_fixable_findings,
            # Structure Agent
            "check_heading_hierarchy": self.check_heading_hierarchy,
            "find_faux_headings": self.find_faux_headings,
            "check_list_structure": self.check_list_structure,
            "estimate_reading_order": self.estimate_reading_order,
            # Content Agent
            "check_emphasis_patterns": self.check_emphasis_patterns,
            "check_link_text": self.check_link_text,
            "check_reading_level": self.check_reading_level,
            "check_alignment_hints": self.check_alignment_hints,
            # Remediation Agent
            "explain_rule": lambda: self.explain_rule(kwargs.get("rule_id", "")),
            "suggest_fix": lambda: self.suggest_fix(kwargs.get("rule_id", "")),
            "prioritize_findings": self.prioritize_findings,
            "estimate_fix_impact": self.estimate_fix_impact,
            "check_image_alt_text": self.check_image_alt_text,
        }
        fn = dispatch.get(tool_name)
        if fn is None:
            return f"Unknown tool: {tool_name}"
        return fn()

    # ------------------------------------------------------------------
    # Pre-flight keyword dispatcher (cloud-first: no LLM function-calling
    # needed — tools run locally and results are injected into the prompt)
    # ------------------------------------------------------------------

    def dispatch_for_question(self, question: str) -> str:
        """Select and run 1–3 relevant tools based on keyword analysis of the question.

        Returns a formatted block suitable for injection into the LLM system prompt,
        grounding the answer in real document facts rather than hallucinated content.
        This pattern works with any OpenRouter model regardless of whether it supports
        JSON function-calling.
        """
        q = question.lower()
        results: list[tuple[str, str]] = []  # (tool_name, result)

        # Rule explanation (must be first — most specific)
        acb_rule = re.search(r"acb-[a-z\-]+", q, re.IGNORECASE)
        if acb_rule:
            rule_id = acb_rule.group(0).upper()
            results.append(("explain_rule", self.explain_rule(rule_id)))

        # Compliance / audit questions
        if any(kw in q for kw in ("audit", "score", "compliance", "pass", "fail", "violation", "violation", "findings")):
            results.append(("run_accessibility_audit", self.run_accessibility_audit()))
            if len(results) < 3:
                results.append(("get_critical_findings", self.get_critical_findings()))

        # Fix / remediation questions
        elif any(kw in q for kw in ("fix", "repair", "correct", "remediat", "auto-fix", "auto fix")):
            results.append(("prioritize_findings", self.prioritize_findings()))
            if len(results) < 3:
                results.append(("estimate_fix_impact", self.estimate_fix_impact()))

        # Heading structure questions
        if any(kw in q for kw in ("heading", "headings", "structure", "hierarchy", "h1", "h2", "h3", "outline")):
            if not any(n == "check_heading_hierarchy" for n, _ in results):
                results.append(("check_heading_hierarchy", self.check_heading_hierarchy()))
            if len(results) < 3:
                results.append(("list_headings", self.list_headings()))

        # Emphasis / formatting
        if any(kw in q for kw in ("italic", "bold", "emphasis", "formatting", "font", "underline")):
            if not any(n == "check_emphasis_patterns" for n, _ in results):
                results.append(("check_emphasis_patterns", self.check_emphasis_patterns()))

        # Link text
        if any(kw in q for kw in ("link", "url", "click here", "hyperlink", "href")):
            if not any(n == "check_link_text" for n, _ in results):
                results.append(("check_link_text", self.check_link_text()))

        # Image / alt text
        if any(kw in q for kw in ("image", "img", "alt text", "alt-text", "picture", "photo", "figure")):
            if not any(n == "check_image_alt_text" for n, _ in results):
                results.append(("check_image_alt_text", self.check_image_alt_text()))

        # List structure
        if any(kw in q for kw in ("list", "bullet", "numbered", "ordered", "unordered")):
            if not any(n == "check_list_structure" for n, _ in results):
                results.append(("check_list_structure", self.check_list_structure()))

        # Reading level
        if any(kw in q for kw in ("reading level", "complexity", "plain language", "readability", "sentence")):
            if not any(n == "check_reading_level" for n, _ in results):
                results.append(("check_reading_level", self.check_reading_level()))

        # Section / content search
        section_match = re.search(r"section [\"']?([a-z0-9 ]{3,40})[\"']?", q)
        if section_match and not any(n == "find_section" for n, _ in results):
            results.append(("find_section", self.find_section(section_match.group(1))))

        keyword_match = re.search(r"(?:search|find|look for|where is)[^\w]*\"?([a-z][a-z0-9 ]{2,40})\"?", q)
        if keyword_match and not any(n == "search_text" for n, _ in results):
            results.append(("search_text", self.search_text(keyword_match.group(1))))

        # Table questions — extract_table first, fall back to reading order
        if any(kw in q for kw in ("table", "column", "row", "data")):
            if not any(n == "extract_table" for n, _ in results):
                # Try to extract the first table; if none found, assess reading order
                table_result = self.extract_table("0")
                if "not found" in table_result:
                    results.append(("estimate_reading_order", self.estimate_reading_order()))
                else:
                    results.append(("extract_table", table_result))

        # Section content questions ("what does X say", "what is in X", etc.)
        section_content_match = re.search(
            r"(?:what (?:does|is|did|are)|tell me about|show me|details? (?:on|about)|content of)"
            r"[^\w]*[\"']?([a-z0-9 ]{3,50})[\"']?",
            q,
        )
        if section_content_match and not any(n == "get_section_content" for n, _ in results):
            results.append(("get_section_content", self.get_section_content(section_content_match.group(1))))

        # Decisions, motions, action items
        if any(kw in q for kw in ("decision", "voted", "motion", "action item", "approved", "resolved", "assigned")):
            if not any(n == "get_decisions_and_actions" for n, _ in results):
                results.append(("get_decisions_and_actions", self.get_decisions_and_actions()))

        # What's good / what passes
        if any(kw in q for kw in ("what passes", "what's good", "what is good", "what works", "passing", "correct", "compliant")):
            if not any(n == "get_what_passes" for n, _ in results):
                results.append(("get_what_passes", self.get_what_passes()))

        # Default fallback: stats + heading summary + emphasis check
        if not results:
            results.append(("get_document_stats", self.get_document_stats()))
            results.append(("check_heading_hierarchy", self.check_heading_hierarchy()))
            results.append(("check_emphasis_patterns", self.check_emphasis_patterns()))

        # Cap at 3 tool results to keep prompt size manageable
        results = results[:3]

        lines = ["=== GLOW Document Analysis (pre-flight tool results) ==="]
        for name, output in results:
            lines.append(f"\n[{name}]\n{output}")
        lines.append("\n=== Use the above analysis to give a grounded, specific answer. ===")
        return "\n".join(lines)


class ConversationTurn:
    """A single Q&A turn with optional tool calls."""

    def __init__(self, turn_number: int, question: str) -> None:
        self.turn_number = turn_number
        self.question = question
        self.answer: str | None = None
        self.tool_calls: list[dict[str, Any]] = []
        self.tool_results: dict[str, str] = {}

    def to_dict(self) -> dict[str, Any]:
        """Convert to dictionary for session storage."""
        return {
            "turn_number": self.turn_number,
            "question": self.question,
            "answer": self.answer,
            "tool_calls": self.tool_calls,
            "tool_results": self.tool_results,
        }

    @staticmethod
    def from_dict(data: dict[str, Any]) -> ConversationTurn:
        """Create from dictionary (session load)."""
        turn = ConversationTurn(data["turn_number"], data["question"])
        turn.answer = data.get("answer")
        turn.tool_calls = data.get("tool_calls", [])
        turn.tool_results = data.get("tool_results", {})
        return turn


class ChatSession:
    """Manages conversation history for a document."""

    def __init__(self, token: str, filename: str) -> None:
        self.token = token
        self.filename = filename
        self.turns: list[ConversationTurn] = []
        self.created_at = datetime.now(UTC)

    def add_turn(self, question: str) -> ConversationTurn:
        """Add a new Q&A turn."""
        turn_number = len(self.turns) + 1
        turn = ConversationTurn(turn_number, question)
        self.turns.append(turn)
        return turn

    def to_dict(self) -> dict[str, Any]:
        """Convert to dictionary for session storage."""
        return {
            "token": self.token,
            "filename": self.filename,
            "created_at": self.created_at.isoformat(),
            "turns": [turn.to_dict() for turn in self.turns],
        }

    @staticmethod
    def from_dict(data: dict[str, Any]) -> ChatSession:
        """Create from dictionary (session load)."""
        session = ChatSession(data["token"], data["filename"])
        session.created_at = datetime.fromisoformat(data["created_at"])
        session.turns = [ConversationTurn.from_dict(t) for t in data.get("turns", [])]
        return session

    def export_markdown(self) -> str:
        """Export conversation to Markdown."""
        lines = [
            f"# Chat Session: {self.filename}",
            f"*Created: {self.created_at.isoformat()}*\n",
        ]
        for turn in self.turns:
            lines.append(f"## Turn {turn.turn_number}: {turn.question}\n")
            if turn.tool_calls:
                lines.append("**Tools used:**")
                for call in turn.tool_calls:
                    lines.append(f"- {call.get('name', 'unknown')}")
                lines.append("")
            if turn.answer:
                lines.append(turn.answer)
            lines.append("")
        return "\n".join(lines)

    def export_word(self, output_path: Path) -> None:
        """Export conversation to Word document."""
        from docx import Document
        doc = Document()
        doc.add_heading(f"Chat Session: {self.filename}", 0)
        doc.add_paragraph(f"Created: {self.created_at.isoformat()}").runs[0].italic = True

        for turn in self.turns:
            doc.add_heading(f"Turn {turn.turn_number}", level=2)
            doc.add_paragraph(turn.question).runs[0].bold = True

            if turn.tool_calls:
                doc.add_paragraph("Tools used:")
                for call in turn.tool_calls:
                    doc.add_paragraph(call.get("name", "unknown"), style="List Bullet")

            if turn.answer:
                doc.add_paragraph(turn.answer)

        doc.save(str(output_path))

    def export_pdf(self, output_path: Path) -> None:
        """Export conversation to PDF via markdown + Pandoc + WeasyPrint."""
        import subprocess
        from weasyprint import HTML

        md_content = self.export_markdown()
        md_path = output_path.with_suffix(".md")
        html_path = output_path.with_suffix(".html")

        # Write markdown
        md_path.write_text(md_content, encoding="utf-8")

        # Convert markdown to HTML via Pandoc
        try:
            subprocess.run(
                ["pandoc", str(md_path), "-o", str(html_path)],
                check=True,
                capture_output=True,
            )
        except (FileNotFoundError, subprocess.CalledProcessError) as e:
            log.error("Pandoc conversion failed: %s", e)
            return

        # Convert HTML to PDF via WeasyPrint
        try:
            HTML(str(html_path)).write_pdf(str(output_path))
            md_path.unlink()
            html_path.unlink()
            log.info("PDF export successful: %s", output_path)
        except Exception as e:
            log.error("WeasyPrint PDF generation failed: %s", e)
