"""End-to-end tests: heading detection + AI refinement + fixer pipeline.

These tests exercise the full flow: create a real .docx, detect headings
with a mocked AI provider, then fix the document and verify the output.
Simulates the complete user experience from paste-to-fixed-document.
"""

from __future__ import annotations

import json
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.shared import Pt

from acb_large_print import constants as C
from acb_large_print.ai_provider import AIResult, build_prompt
from acb_large_print.ai_providers.ollama_provider import OllamaProvider
from acb_large_print.fixer import _convert_faux_headings, fix_document
from acb_large_print.heading_detector import HeadingCandidate, detect_headings

# ===================================================================
# Helpers
# ===================================================================


def _add(doc, text, *, bold=False, size_pt=None, style="Normal"):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size_pt:
        run.font.size = Pt(size_pt)
    return p


def _body(n=3):
    return (
        "This is ordinary body text that demonstrates typical paragraph content. " * n
    ).strip()


def _mock_provider_confirming_all():
    """AI provider that confirms every candidate as a heading."""
    provider = MagicMock(spec=OllamaProvider)
    provider.system_prompt = None

    def classify(candidates, contexts, **kwargs):
        return [
            AIResult(
                is_heading=True,
                level=c.suggested_level,
                confidence=0.95,
                reasoning="Confirmed by AI",
            )
            for c in candidates
        ]

    provider.classify_candidates = classify
    return provider


def _mock_provider_rejecting_all():
    """AI provider that rejects every candidate."""
    provider = MagicMock(spec=OllamaProvider)
    provider.system_prompt = None

    def classify(candidates, contexts, **kwargs):
        return [
            AIResult(
                is_heading=False,
                level=None,
                confidence=0.90,
                reasoning="Not a heading",
            )
            for _ in candidates
        ]

    provider.classify_candidates = classify
    return provider


def _mock_provider_selective():
    """AI provider that confirms headings and rejects non-headings
    based on whether the text is Title Case."""
    provider = MagicMock(spec=OllamaProvider)
    provider.system_prompt = None

    def classify(candidates, contexts, **kwargs):
        results = []
        for c in candidates:
            if c.text.istitle() or c.is_all_caps:
                results.append(
                    AIResult(
                        is_heading=True,
                        level=c.suggested_level,
                        confidence=0.88,
                        reasoning="Title case indicates heading",
                    )
                )
            else:
                results.append(
                    AIResult(
                        is_heading=False,
                        level=None,
                        confidence=0.85,
                        reasoning="Not structured as heading",
                    )
                )
        return results

    provider.classify_candidates = classify
    return provider


def _mock_provider_that_relabels():
    """AI provider that confirms headings but changes their levels."""
    provider = MagicMock(spec=OllamaProvider)
    provider.system_prompt = None

    def classify(candidates, contexts, **kwargs):
        return [
            AIResult(
                is_heading=True,
                level=min(c.suggested_level + 1, 6),
                confidence=0.90,
                reasoning="Demoted one level by AI",
            )
            for c in candidates
        ]

    provider.classify_candidates = classify
    return provider


def _mock_provider_failing():
    """AI provider that throws an exception."""
    provider = MagicMock(spec=OllamaProvider)
    provider.system_prompt = None
    provider.classify_candidates.side_effect = RuntimeError("Ollama crashed")
    return provider


def _mock_provider_partial_failure():
    """AI provider where some candidates fail (return None)."""
    provider = MagicMock(spec=OllamaProvider)
    provider.system_prompt = None

    def classify(candidates, contexts, **kwargs):
        results = []
        for i, c in enumerate(candidates):
            if i % 2 == 0:
                results.append(
                    AIResult(
                        is_heading=True,
                        level=c.suggested_level,
                        confidence=0.85,
                        reasoning="OK",
                    )
                )
            else:
                results.append(None)
        return results

    provider.classify_candidates = classify
    return provider


# ===================================================================
# 1. NOTEPAD PASTE -> DETECT -> AI CONFIRM -> FIX
# ===================================================================


class TestNotepadToFixedPipeline:
    """Full pipeline: user pastes from Notepad, we detect, AI confirms, fix."""

    def _build_notepad_report(self, tmp_path):
        doc = Document()
        _add(doc, "QUARTERLY FINANCIAL REPORT", bold=True, size_pt=24)
        _add(doc, "")
        _add(doc, _body(5), size_pt=12)
        _add(doc, "")
        _add(doc, "Revenue Analysis", bold=True, size_pt=18)
        _add(doc, _body(4), size_pt=12)
        _add(doc, "")
        _add(doc, "Expense Breakdown", bold=True, size_pt=18)
        _add(doc, _body(4), size_pt=12)
        _add(doc, "")
        _add(doc, "Recommendations", bold=True, size_pt=18)
        _add(doc, _body(3), size_pt=12)
        path = tmp_path / "notepad_report.docx"
        doc.save(str(path))
        return path

    def test_detect_then_fix_with_confirmed(self, tmp_path):
        path = self._build_notepad_report(tmp_path)
        doc = Document(str(path))
        candidates = detect_headings(doc)
        assert len(candidates) >= 3

        # Simulate user confirming all detected headings
        confirmed = [(c.paragraph_index, c.suggested_level, c.text) for c in candidates]
        out = tmp_path / "fixed.docx"
        _, total, records, _, _ = fix_document(
            path,
            out,
            detect_headings=True,
            confirmed_headings=confirmed,
        )
        heading_recs = [r for r in records if r.rule_id == "ACB-FAUX-HEADING"]
        assert len(heading_recs) >= 3

        # Verify the output document has real heading styles
        fixed_doc = Document(str(out))
        styled = [p for p in fixed_doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(styled) >= 3

    def test_detect_with_ai_confirming(self, tmp_path):
        path = self._build_notepad_report(tmp_path)
        doc = Document(str(path))
        provider = _mock_provider_confirming_all()
        candidates = detect_headings(doc, ai_provider=provider)

        # Medium-confidence candidates should now be "high"
        for c in candidates:
            if c.ai_reasoning:
                assert "Confirmed" in c.ai_reasoning

    def test_detect_with_ai_rejecting_medium(self, tmp_path):
        path = self._build_notepad_report(tmp_path)
        doc = Document(str(path))
        # First get candidates without AI
        no_ai = detect_headings(doc)
        # Then with rejecting AI
        provider = _mock_provider_rejecting_all()
        with_ai = detect_headings(doc, ai_provider=provider)
        # AI should have removed medium-confidence candidates
        assert len(with_ai) <= len(no_ai)


# ===================================================================
# 2. EMAIL PASTE -> AI SELECTIVE -> FIX
# ===================================================================


class TestEmailToFixedPipeline:
    def _build_email_thread(self, tmp_path):
        doc = Document()
        _add(doc, "Project Status Update", bold=True, size_pt=16)
        _add(doc, "")
        _add(doc, "Hi all,", size_pt=11)
        _add(doc, _body(3), size_pt=11)
        _add(doc, "")
        _add(doc, "Timeline", bold=True, size_pt=14)
        _add(doc, _body(2), size_pt=11)
        _add(doc, "")
        _add(doc, "next steps for the team", bold=True, size_pt=11)
        _add(doc, _body(2), size_pt=11)
        _add(doc, "")
        _add(doc, "Best regards,", size_pt=11)
        _add(doc, "Alice Johnson", bold=True, size_pt=11)
        path = tmp_path / "email.docx"
        doc.save(str(path))
        return path

    def test_selective_ai_keeps_real_headings(self, tmp_path):
        path = self._build_email_thread(tmp_path)
        doc = Document(str(path))
        provider = _mock_provider_selective()
        candidates = detect_headings(doc, ai_provider=provider)

        texts = [c.text for c in candidates]
        # Title Case headings should survive
        assert "Project Status Update" in texts
        assert "Timeline" in texts

    def test_selective_ai_removes_lowercase(self, tmp_path):
        """'next steps for the team' is lowercase -- selective AI rejects it."""
        path = self._build_email_thread(tmp_path)
        doc = Document(str(path))
        provider = _mock_provider_selective()
        candidates = detect_headings(doc, ai_provider=provider)

        texts = [c.text for c in candidates]
        assert "next steps for the team" not in texts


# ===================================================================
# 3. LEGAL DOC -> AI RELABELS LEVELS -> FIX
# ===================================================================


class TestLegalDocRelabelPipeline:
    def _build_legal(self, tmp_path):
        doc = Document()
        _add(doc, "CONTRACT FOR SERVICES", bold=True, size_pt=22)
        _add(doc, "")
        _add(doc, _body(4), size_pt=12)
        _add(doc, "")
        _add(doc, "ARTICLE I. SCOPE", bold=True, size_pt=16)
        _add(doc, _body(3), size_pt=12)
        _add(doc, "")
        _add(doc, "ARTICLE II. PAYMENT", bold=True, size_pt=16)
        _add(doc, _body(3), size_pt=12)
        path = tmp_path / "legal.docx"
        doc.save(str(path))
        return path

    def test_ai_relabels_levels(self, tmp_path):
        path = self._build_legal(tmp_path)
        doc = Document(str(path))
        provider = _mock_provider_that_relabels()
        candidates = detect_headings(doc, ai_provider=provider)

        # Medium-confidence candidates should have been relabeled
        for c in candidates:
            if c.ai_reasoning and "Demoted" in c.ai_reasoning:
                assert c.suggested_level >= 2

    def test_ai_failure_preserves_heuristic(self, tmp_path):
        path = self._build_legal(tmp_path)
        doc = Document(str(path))
        # Get heuristic-only result
        heuristic = detect_headings(doc)
        # Now try with failing AI
        provider = _mock_provider_failing()
        with_ai = detect_headings(doc, ai_provider=provider)

        # Should have same number of candidates (AI failure = fallback)
        assert len(with_ai) == len(heuristic)
        # Medium candidates should have fallback reasoning
        for c in with_ai:
            if c.confidence == "medium":
                assert c.ai_reasoning is not None


# ===================================================================
# 4. AI PARTIAL FAILURE
# ===================================================================


class TestPartialAIFailure:
    def _build_doc(self, tmp_path):
        doc = Document()
        _add(doc, "Title", bold=True, size_pt=24)
        _add(doc, _body(3), size_pt=12)
        _add(doc, "")
        _add(doc, "Section A", bold=True, size_pt=18)
        _add(doc, _body(3), size_pt=12)
        _add(doc, "")
        _add(doc, "Section B", bold=True, size_pt=18)
        _add(doc, _body(3), size_pt=12)
        _add(doc, "")
        _add(doc, "Section C", bold=True, size_pt=18)
        _add(doc, _body(3), size_pt=12)
        path = tmp_path / "partial.docx"
        doc.save(str(path))
        return path

    def test_partial_ai_results_mixed(self, tmp_path):
        """Some candidates get AI results, others get None."""
        path = self._build_doc(tmp_path)
        doc = Document(str(path))
        provider = _mock_provider_partial_failure()
        candidates = detect_headings(doc, ai_provider=provider)
        # Should not crash, all candidates should be present
        assert len(candidates) >= 1
        # Candidates with None results keep their heuristic scores
        assert all(isinstance(c, HeadingCandidate) for c in candidates)


# ===================================================================
# 5. CONFIRMED HEADINGS PIPELINE -- bypass detection
# ===================================================================


class TestConfirmedHeadingsPipeline:
    def _build_doc(self, tmp_path):
        doc = Document()
        _add(doc, "Report Title", size_pt=12)
        _add(doc, _body(2), size_pt=12)
        _add(doc, "Introduction", size_pt=12)
        _add(doc, _body(2), size_pt=12)
        _add(doc, "Conclusion", size_pt=12)
        _add(doc, _body(2), size_pt=12)
        path = tmp_path / "unformatted.docx"
        doc.save(str(path))
        return path

    def test_confirmed_headings_apply_without_detection(self, tmp_path):
        """User manually picks headings from plain text -- no visual signals."""
        path = self._build_doc(tmp_path)
        confirmed = [
            (0, 1, "Report Title"),
            (2, 2, "Introduction"),
            (4, 2, "Conclusion"),
        ]
        out = tmp_path / "fixed.docx"
        _, _, records, _, _ = fix_document(
            path,
            out,
            detect_headings=True,
            confirmed_headings=confirmed,
        )
        heading_recs = [r for r in records if r.rule_id == "ACB-FAUX-HEADING"]
        assert len(heading_recs) == 3

        fixed = Document(str(out))
        assert fixed.paragraphs[0].style.name == "Heading 1"
        assert fixed.paragraphs[2].style.name == "Heading 2"
        assert fixed.paragraphs[4].style.name == "Heading 2"

    def test_confirmed_heading_stale_text_skipped(self, tmp_path):
        """If paragraph text doesn't match confirmation, skip it."""
        path = self._build_doc(tmp_path)
        confirmed = [
            (0, 1, "WRONG TEXT"),
        ]
        out = tmp_path / "stale.docx"
        _, _, records, _, _ = fix_document(
            path,
            out,
            detect_headings=True,
            confirmed_headings=confirmed,
        )
        heading_recs = [r for r in records if r.rule_id == "ACB-FAUX-HEADING"]
        assert len(heading_recs) == 0

    def test_confirmed_paragraph_index_out_of_range(self, tmp_path):
        path = self._build_doc(tmp_path)
        confirmed = [(999, 1, "Ghost Heading")]
        out = tmp_path / "range.docx"
        _, _, records, _, _ = fix_document(
            path,
            out,
            detect_headings=True,
            confirmed_headings=confirmed,
        )
        heading_recs = [r for r in records if r.rule_id == "ACB-FAUX-HEADING"]
        assert len(heading_recs) == 0

    def test_empty_confirmed_list(self, tmp_path):
        path = self._build_doc(tmp_path)
        confirmed = []
        out = tmp_path / "empty.docx"
        _, _, records, _, _ = fix_document(
            path,
            out,
            detect_headings=True,
            confirmed_headings=confirmed,
        )
        heading_recs = [r for r in records if r.rule_id == "ACB-FAUX-HEADING"]
        assert len(heading_recs) == 0


# ===================================================================
# 6. FIX CLEARS DIRECT FORMATTING
# ===================================================================


class TestFixClearsFormatting:
    def test_bold_and_size_cleared_after_fix(self, tmp_path):
        """After converting to heading style, direct formatting is cleared."""
        doc = Document()
        _add(doc, "Chapter One", bold=True, size_pt=22)
        _add(doc, _body(3), size_pt=12)
        path = tmp_path / "direct_fmt.docx"
        doc.save(str(path))

        confirmed = [(0, 1, "Chapter One")]
        out = tmp_path / "cleared.docx"
        fix_document(path, out, detect_headings=True, confirmed_headings=confirmed)

        fixed = Document(str(out))
        heading_para = fixed.paragraphs[0]
        assert heading_para.style.name == "Heading 1"
        # Direct formatting should be cleared (set to None = inherit from style)
        for run in heading_para.runs:
            assert run.font.size is None
            assert run.font.bold is None
            assert run.font.name is None


# ===================================================================
# 7. DETECT + AI + FIX -- complete round-trip
# ===================================================================


class TestCompleteRoundTrip:
    def _build_report(self, tmp_path):
        doc = Document()
        _add(doc, "Annual Review 2025", bold=True, size_pt=26)
        _add(doc, "")
        _add(doc, _body(5), size_pt=12)
        _add(doc, "")
        _add(doc, "Objectives Met", bold=True, size_pt=20)
        _add(doc, _body(4), size_pt=12)
        _add(doc, "")
        _add(doc, "Areas for Improvement", bold=True, size_pt=20)
        _add(doc, _body(4), size_pt=12)
        _add(doc, "")
        _add(doc, "Training Completed", bold=True, size_pt=16)
        _add(doc, _body(3), size_pt=12)
        _add(doc, "")
        _add(doc, "Goals for 2026", bold=True, size_pt=20)
        _add(doc, _body(4), size_pt=12)
        path = tmp_path / "annual_review.docx"
        doc.save(str(path))
        return path

    def test_full_pipeline(self, tmp_path):
        """Detect -> AI confirms -> user confirms -> fix -> verify output."""
        path = self._build_report(tmp_path)

        # Step 1: Detect with AI confirmation
        doc = Document(str(path))
        provider = _mock_provider_confirming_all()
        candidates = detect_headings(doc, ai_provider=provider)
        assert len(candidates) >= 4

        # Step 2: Simulate user confirming detections
        confirmed = [(c.paragraph_index, c.suggested_level, c.text) for c in candidates]

        # Step 3: Fix the document
        out_path = tmp_path / "annual_review_fixed.docx"
        returned, total, records, audit, warnings = fix_document(
            path,
            out_path,
            detect_headings=True,
            confirmed_headings=confirmed,
        )

        # Step 4: Verify
        assert out_path.exists()
        heading_recs = [r for r in records if r.rule_id == "ACB-FAUX-HEADING"]
        assert len(heading_recs) >= 4

        fixed_doc = Document(str(out_path))
        h1s = [p for p in fixed_doc.paragraphs if p.style.name == "Heading 1"]
        assert len(h1s) >= 1
        assert "Annual Review" in h1s[0].text

    def test_multiple_rounds_idempotent(self, tmp_path):
        """Running detect on an already-fixed document yields no candidates."""
        path = self._build_report(tmp_path)
        doc = Document(str(path))
        candidates = detect_headings(doc)
        confirmed = [(c.paragraph_index, c.suggested_level, c.text) for c in candidates]

        # Fix once
        out1 = tmp_path / "fixed1.docx"
        fix_document(path, out1, detect_headings=True, confirmed_headings=confirmed)

        # Detect again on fixed document -- should find nothing new
        fixed_doc = Document(str(out1))
        new_candidates = detect_headings(fixed_doc)
        # The newly styled headings should NOT be re-detected
        for c in new_candidates:
            assert c.text not in [x[2] for x in confirmed]


# ===================================================================
# 8. SYSTEM PROMPT -- passed through AI pipeline
# ===================================================================


class TestSystemPromptPipeline:
    def test_custom_prompt_reaches_provider(self, tmp_path):
        doc = Document()
        _add(doc, "Custom Title", bold=True, size_pt=22)
        # Need a medium candidate for AI to process
        _add(doc, "maybe heading", bold=True)  # no size = likely medium
        _add(doc, _body(3), size_pt=12)

        provider = MagicMock(spec=OllamaProvider)
        provider.system_prompt = None
        provider.classify_candidates.return_value = [
            AIResult(is_heading=True, level=2, confidence=0.9, reasoning="OK")
        ]

        detect_headings(
            doc, ai_provider=provider, system_prompt="My custom: {paragraph_text}"
        )

        # Verify the custom prompt was set on the provider
        assert provider.system_prompt == "My custom: {paragraph_text}"


# ===================================================================
# 9. MIXED REAL + FAUX HEADINGS
# ===================================================================


class TestMixedRealAndFauxHeadings:
    def test_real_headings_not_re_detected(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Real H1", style="Heading 1")
        _add(doc, _body(3), size_pt=12)
        _add(doc, "Faux Section", bold=True, size_pt=20)
        _add(doc, _body(3), size_pt=12)
        doc.add_paragraph("Real H2", style="Heading 2")
        _add(doc, _body(3), size_pt=12)
        _add(doc, "Another Faux", bold=True, size_pt=20)
        _add(doc, _body(3), size_pt=12)

        candidates = detect_headings(doc)
        texts = [c.text for c in candidates]
        assert "Real H1" not in texts
        assert "Real H2" not in texts
        assert "Faux Section" in texts
        assert "Another Faux" in texts


# ===================================================================
# 10. STRESS TEST -- many heading levels and large document
# ===================================================================


class TestStressPipeline:
    def test_20_sections_with_3_subsections_each(self, tmp_path):
        doc = Document()
        _add(doc, "Master Document", bold=True, size_pt=28)
        _add(doc, _body(3), size_pt=12)

        for s in range(1, 21):
            _add(doc, "")
            _add(doc, f"Section {s}", bold=True, size_pt=22)
            _add(doc, _body(2), size_pt=12)
            for sub in range(1, 4):
                _add(doc, f"Section {s}.{sub} Detail", bold=True, size_pt=18)
                _add(doc, _body(2), size_pt=12)

        candidates = detect_headings(doc)
        # 1 title + 20 sections + 60 subsections = 81 total
        assert len(candidates) >= 60
        # Verify level structure
        assert candidates[0].suggested_level == 1
        for i in range(1, len(candidates)):
            assert (
                candidates[i].suggested_level <= candidates[i - 1].suggested_level + 1
            )

    def test_fix_large_document(self, tmp_path):
        doc = Document()
        _add(doc, "Big Report", bold=True, size_pt=24)
        for i in range(30):
            _add(doc, "")
            _add(doc, f"Chapter {i+1}", bold=True, size_pt=20)
            _add(doc, _body(4), size_pt=12)
        path = tmp_path / "big.docx"
        doc.save(str(path))

        doc = Document(str(path))
        candidates = detect_headings(doc)
        confirmed = [(c.paragraph_index, c.suggested_level, c.text) for c in candidates]

        out = tmp_path / "big_fixed.docx"
        _, _, records, _, _ = fix_document(
            path,
            out,
            detect_headings=True,
            confirmed_headings=confirmed,
        )
        heading_recs = [r for r in records if r.rule_id == "ACB-FAUX-HEADING"]
        assert len(heading_recs) >= 25
