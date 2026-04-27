from pathlib import Path

from docx import Document
from docx.shared import Pt

from acb_large_print import constants as C
from acb_large_print.auditor import audit_document
from acb_large_print.fixer import fix_document


def test_fix_document_normalizes_heading4_and_body_paragraph_spacing(tmp_path: Path):
    doc = Document()

    heading = doc.add_paragraph("Deep section heading")
    heading.style = doc.styles["Heading 4"]
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(0)
    heading_run = heading.runs[0]
    heading_run.font.name = "Calibri"
    heading_run.font.size = Pt(12)
    heading_run.font.italic = True

    body = doc.add_paragraph("Body paragraph with stale spacing.")
    body.style = doc.styles["Normal"]
    body.paragraph_format.space_after = Pt(8)

    source = tmp_path / "spacing-source.docx"
    output = tmp_path / "spacing-fixed.docx"
    doc.save(source)

    fix_document(source, output)

    fixed = Document(output)
    fixed_heading = fixed.paragraphs[0]
    fixed_heading_run = fixed_heading.runs[0]
    fixed_body = fixed.paragraphs[1]

    assert fixed_heading.style.name == "Heading 4"
    assert fixed_heading_run.font.name == C.FONT_FAMILY
    assert fixed_heading_run.font.size.pt == C.HEADING3_SIZE_PT
    assert fixed_heading_run.font.italic is False
    assert fixed_heading.paragraph_format.space_before.pt == C.SPACE_BEFORE_H3_PT
    assert fixed_heading.paragraph_format.space_after.pt == C.SPACE_AFTER_H3_PT
    assert fixed_body.paragraph_format.space_after.pt == C.SPACE_AFTER_BODY_PT


def test_audit_document_caps_repeated_non_arial_findings(tmp_path: Path):
    doc = Document()
    for index in range(5):
        para = doc.add_paragraph(f"Paragraph {index + 1}")
        para.runs[0].font.name = "Times New Roman"

    source = tmp_path / "font-noise.docx"
    doc.save(source)

    result = audit_document(source)
    font_findings = [finding for finding in result.findings if finding.rule_id == "ACB-FONT-FAMILY"]

    assert len(font_findings) == 4
    assert sum("Non-Arial font 'Times New Roman'" in finding.message for finding in font_findings) == 4
    assert any("showing the first 3 locations only" in finding.message for finding in font_findings)