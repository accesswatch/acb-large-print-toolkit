"""Large randomized stress tests for heading detection and fixer robustness."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from acb_large_print.ai_provider import AIResult
from acb_large_print.fixer import fix_document
from acb_large_print.heading_detector import detect_headings
from acb_large_print.stress_profiles import (
    DEFAULT_STRESS_SAMPLE_DOCUMENTS,
    STRESS_CASES_PER_DOCUMENT,
    STRESS_FIXER_DOCUMENTS,
    STRESS_TOTAL_HEADING_CASES,
    build_stress_document,
    describe_stress_corpus,
    generate_stress_documents,
)


class _ScenarioAIProvider:
    def __init__(self, expected: dict[str, tuple[bool, int | None]]):
        self.expected = expected
        self.system_prompt = None

    def classify_candidates(self, candidates, contexts, **kwargs):
        results: list[AIResult] = []
        for candidate in candidates:
            is_heading, level = self.expected[candidate.text]
            results.append(
                AIResult(
                    is_heading=is_heading,
                    level=level,
                    confidence=0.96 if is_heading else 0.93,
                    reasoning="Synthetic stress corpus oracle",
                )
            )
        return results


def _assert_fixed_docx(path: Path) -> None:
    doc = Document(path)
    for para in doc.paragraphs:
        if para.text.strip() and para.paragraph_format.alignment is not None:
            assert para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT

        style_name = para.style.name if para.style else "Normal"
        is_list = style_name.startswith("List ")
        is_heading = style_name.startswith("Heading ")
        if para.text.strip() and not is_list and not is_heading:
            left_indent = (
                para.paragraph_format.left_indent.inches
                if para.paragraph_format.left_indent
                else 0.0
            )
            first_line_indent = (
                para.paragraph_format.first_line_indent.inches
                if para.paragraph_format.first_line_indent
                else 0.0
            )
            assert abs(left_indent) <= 0.05
            assert abs(first_line_indent) <= 0.05


def test_stress_corpus_metadata_is_public_and_complete():
    summary = describe_stress_corpus()
    assert summary["document_count"] == STRESS_FIXER_DOCUMENTS
    assert summary["cases_per_document"] == STRESS_CASES_PER_DOCUMENT
    assert summary["total_heading_cases"] == STRESS_TOTAL_HEADING_CASES
    assert len(summary["families"]) >= 10


def test_sampled_two_phase_heading_detection_stress():
    total_cases = 0
    total_headings = 0

    for document_index in range(DEFAULT_STRESS_SAMPLE_DOCUMENTS):
        doc, scenarios = build_stress_document(document_index, cases_per_document=24)
        expected = {
            scenario.text: (scenario.final_is_heading, scenario.expected_level)
            for scenario in scenarios
            if scenario.is_candidate
        }
        provider = _ScenarioAIProvider(expected)
        detected = detect_headings(doc, ai_provider=provider)

        actual = {candidate.text for candidate in detected}
        wanted = {scenario.text for scenario in scenarios if scenario.final_is_heading}

        assert actual == wanted
        prev_level = 0
        for candidate in detected:
            assert 1 <= candidate.suggested_level <= 6
            if prev_level:
                assert candidate.suggested_level <= prev_level + 1
            prev_level = candidate.suggested_level

        total_cases += len(scenarios)
        total_headings += len(wanted)

    assert total_cases == DEFAULT_STRESS_SAMPLE_DOCUMENTS * 24
    assert total_headings > 0


def test_sampled_fixer_stress_enforces_flush_left_and_heading_conversion(
    tmp_path: Path,
):
    generated = generate_stress_documents(
        tmp_path / "sample-corpus",
        document_count=8,
        cases_per_document=18,
    )
    assert len(generated) == 8

    for generated_doc in generated:
        out_path = generated_doc.path.with_name(f"fixed-{generated_doc.path.name}")
        confirmed = [
            (scenario.case_index * 3 + 3, scenario.expected_level, scenario.text)
            for scenario in generated_doc.scenarios
            if scenario.final_is_heading and scenario.expected_level is not None
        ]
        fix_document(
            generated_doc.path,
            out_path,
            detect_headings=True,
            confirmed_headings=confirmed,
        )
        fixed_doc = Document(out_path)
        for scenario in generated_doc.scenarios:
            if scenario.final_is_heading and scenario.expected_level is not None:
                fixed_para = fixed_doc.paragraphs[scenario.case_index * 3 + 3]
                assert fixed_para.style.name == f"Heading {scenario.expected_level}"
        _assert_fixed_docx(out_path)


@pytest.mark.stress
def test_full_stress_corpus_generates_one_thousand_documents(tmp_path: Path):
    generated = generate_stress_documents(tmp_path / "full-corpus")
    assert len(generated) == STRESS_FIXER_DOCUMENTS
    assert sum(len(doc.scenarios) for doc in generated) == STRESS_TOTAL_HEADING_CASES


@pytest.mark.stress
def test_full_fixer_stress_corpus_repairs_all_documents(tmp_path: Path):
    generated = generate_stress_documents(tmp_path / "full-fixer")

    for generated_doc in generated:
        out_path = generated_doc.path.with_name(f"fixed-{generated_doc.path.name}")
        confirmed = [
            (scenario.case_index * 3 + 3, scenario.expected_level, scenario.text)
            for scenario in generated_doc.scenarios
            if scenario.final_is_heading and scenario.expected_level is not None
        ]
        fix_document(
            generated_doc.path,
            out_path,
            detect_headings=True,
            confirmed_headings=confirmed,
        )
        _assert_fixed_docx(out_path)
