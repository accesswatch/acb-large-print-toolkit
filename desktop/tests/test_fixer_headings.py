"""Tests for faux-heading conversion in the fixer module."""

from __future__ import annotations

import shutil
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.shared import Pt

from acb_large_print import constants as C
from acb_large_print.fixer import _convert_faux_headings, fix_document

# ---------------------------------------------------------------------------
# Integration helpers
# ---------------------------------------------------------------------------


def _make_docx_with_faux_heading(tmp_path: Path) -> Path:
    """Create a .docx file that has bold text styled as Normal (a faux heading)."""
    doc = Document()
    # Add a bold, large-font paragraph that should be detected as a heading
    p = doc.add_paragraph()
    run = p.add_run("Chapter One")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Arial"
    p.style = doc.styles["Normal"]

    # Add body text
    body = doc.add_paragraph("This is the body text that follows the heading. " * 5)
    for r in body.runs:
        r.font.size = Pt(18)

    path = tmp_path / "faux_heading.docx"
    doc.save(str(path))
    return path


def _make_plain_docx(tmp_path: Path) -> Path:
    """Create a basic .docx with no faux headings."""
    doc = Document()
    doc.add_paragraph("Normal paragraph one.")
    doc.add_paragraph("Normal paragraph two.")
    path = tmp_path / "plain.docx"
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# _convert_faux_headings -- unit tests
# ---------------------------------------------------------------------------


class TestConvertFauxHeadings:
    def test_confirmed_headings_applied(self, tmp_path):
        """Confirmed headings bypass detection and convert directly."""
        doc = Document()
        p1 = doc.add_paragraph("Introduction")
        p2 = doc.add_paragraph("Body text goes here.")
        path = tmp_path / "confirmed.docx"
        doc.save(str(path))

        doc = Document(str(path))
        records: list[C.FixRecord] = []
        confirmed = [(0, 1, "Introduction")]
        count = _convert_faux_headings(doc, records, confirmed_headings=confirmed)

        assert count == 1
        assert doc.paragraphs[0].style.name == "Heading 1"
        assert len(records) == 1
        assert records[0].rule_id == "ACB-FAUX-HEADING"

    def test_confirmed_heading_wrong_text_skipped(self, tmp_path):
        """If paragraph text changed, the conversion is safely skipped."""
        doc = Document()
        doc.add_paragraph("Original Title")
        path = tmp_path / "stale.docx"
        doc.save(str(path))

        doc = Document(str(path))
        records: list[C.FixRecord] = []
        confirmed = [(0, 1, "Different Title")]
        count = _convert_faux_headings(doc, records, confirmed_headings=confirmed)

        assert count == 0
        assert doc.paragraphs[0].style.name == "Normal"

    def test_confirmed_heading_out_of_range_skipped(self, tmp_path):
        """Paragraph index beyond document length is safely skipped."""
        doc = Document()
        doc.add_paragraph("Only paragraph")
        path = tmp_path / "short.docx"
        doc.save(str(path))

        doc = Document(str(path))
        records: list[C.FixRecord] = []
        confirmed = [(999, 1, "Only paragraph")]
        count = _convert_faux_headings(doc, records, confirmed_headings=confirmed)
        assert count == 0

    def test_no_candidates_returns_zero(self, tmp_path):
        """When detection finds nothing, zero conversions reported."""
        path = _make_plain_docx(tmp_path)
        doc = Document(str(path))
        records: list[C.FixRecord] = []
        count = _convert_faux_headings(doc, records)
        assert count == 0

    def test_multiple_confirmed_headings(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Chapter 1")
        doc.add_paragraph("Body text.")
        doc.add_paragraph("Section A")
        path = tmp_path / "multi.docx"
        doc.save(str(path))

        doc = Document(str(path))
        records: list[C.FixRecord] = []
        confirmed = [
            (0, 1, "Chapter 1"),
            (2, 2, "Section A"),
        ]
        count = _convert_faux_headings(doc, records, confirmed_headings=confirmed)
        assert count == 2
        assert doc.paragraphs[0].style.name == "Heading 1"
        assert doc.paragraphs[2].style.name == "Heading 2"


# ---------------------------------------------------------------------------
# fix_document -- integration tests with heading detection
# ---------------------------------------------------------------------------


class TestFixDocumentHeadings:
    def test_fix_with_detect_headings_off(self, tmp_path):
        path = _make_plain_docx(tmp_path)
        out = tmp_path / "output.docx"
        _, total_fixes, records, audit, warnings = fix_document(
            path,
            out,
            detect_headings=False,
        )
        assert isinstance(total_fixes, int)
        assert out.exists()
        # No heading records expected
        heading_records = [r for r in records if r.rule_id == "ACB-FAUX-HEADING"]
        assert len(heading_records) == 0

    def test_fix_with_confirmed_headings(self, tmp_path):
        doc = Document()
        doc.add_paragraph("My Title")
        doc.add_paragraph("Some body text here that is reasonably long for testing.")
        path = tmp_path / "with_heading.docx"
        doc.save(str(path))

        out = tmp_path / "fixed.docx"
        _, total_fixes, records, audit, warnings = fix_document(
            path,
            out,
            detect_headings=True,
            confirmed_headings=[(0, 1, "My Title")],
        )
        assert out.exists()
        heading_recs = [r for r in records if r.rule_id == "ACB-FAUX-HEADING"]
        assert len(heading_recs) == 1

    def test_fix_output_defaults_to_input(self, tmp_path):
        path = _make_plain_docx(tmp_path)
        returned_path, _, _, _, _ = fix_document(path)
        assert returned_path == path

    def test_fix_creates_output_dir(self, tmp_path):
        path = _make_plain_docx(tmp_path)
        out = tmp_path / "subdir" / "output.docx"
        returned_path, _, _, _, _ = fix_document(path, out)
        assert out.exists()
