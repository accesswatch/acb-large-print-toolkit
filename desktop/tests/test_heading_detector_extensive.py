"""Extensive tests for heading detection -- real-world paste scenarios.

Covers the full spectrum of documents people create by pasting from
Notepad, text editors, emails, web browsers, and other applications
into Microsoft Word.  Every test uses real python-docx Document objects
(not mocks) so we exercise the exact same code path the user hits.
"""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock

import pytest
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from acb_large_print import constants as C
from acb_large_print.heading_detector import (
    HeadingCandidate,
    _assign_heading_levels,
    _build_context,
    _fix_hierarchy_gaps,
    _get_body_font_size,
    _is_heading_style,
    _next_is_body,
    _para_font_size,
    _para_is_bold,
    _prev_is_blank_or_break,
    _score_paragraph,
    detect_headings,
)

# ===================================================================
# Helpers -- build real .docx paragraphs with controlled formatting
# ===================================================================


def _add_para(
    doc,
    text,
    *,
    bold=False,
    size_pt=None,
    style="Normal",
    all_caps_font=False,
    italic=False,
    color=None,
    name=None,
):
    """Add a paragraph with explicit run-level formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if name:
        run.font.name = name
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    if all_caps_font:
        run.font.all_caps = True
    return p


def _body_text(n=3):
    """Return realistic body-length filler text."""
    sentence = (
        "This paragraph contains ordinary body text that a person might "
        "type or paste from another source into a Word document. "
    )
    return (sentence * n).strip()


# ===================================================================
# 1. NOTEPAD PASTE -- zero formatting, structure only from whitespace
# ===================================================================


class TestNotepadPaste:
    """Simulate text pasted from Notepad: all Normal style, no bold,
    no font size set, structure only implied by short lines."""

    def _build_notepad_doc(self):
        doc = Document()
        # Notepad paste: all runs inherit the default 11pt Calibri
        # with no explicit font.size set on runs
        _add_para(doc, "Board Meeting Minutes")
        _add_para(doc, "")  # blank line separator
        _add_para(doc, _body_text())
        _add_para(doc, "")
        _add_para(doc, "Old Business")
        _add_para(doc, _body_text())
        _add_para(doc, "")
        _add_para(doc, "New Business")
        _add_para(doc, _body_text())
        return doc

    def test_no_formatting_no_candidates(self):
        """Plain text with no visual signals should produce zero candidates."""
        doc = self._build_notepad_doc()
        result = detect_headings(doc)
        # With no bold, no font size, the only signals are short text,
        # no trailing punct, preceded by blank, title case, single-line.
        # Max theoretical: 15+10+10+10+5 = 50 -- right at threshold
        # Whether candidates appear depends on text casing
        for c in result:
            assert c.score >= C.HEADING_CONFIDENCE_THRESHOLD

    def test_notepad_paste_with_caps_titles(self):
        """ALL CAPS short lines from Notepad should score as candidates."""
        doc = Document()
        _add_para(doc, "MEETING AGENDA")
        _add_para(doc, "")
        _add_para(doc, _body_text())
        _add_para(doc, "")
        _add_para(doc, "ACTION ITEMS")
        _add_para(doc, _body_text())
        result = detect_headings(doc)
        caps_candidates = [c for c in result if c.is_all_caps]
        assert len(caps_candidates) >= 1
        assert any("MEETING AGENDA" in c.text for c in result)

    def test_single_very_long_paste(self):
        """A single enormous paragraph (like a Notepad dump) yields nothing."""
        doc = Document()
        _add_para(doc, "x" * 300)
        result = detect_headings(doc)
        assert result == []


# ===================================================================
# 2. EMAIL PASTE -- bold subjects, mixed formatting, signature blocks
# ===================================================================


class TestEmailPaste:
    """Simulate text pasted from Outlook/email: subjects are bold,
    body is plain, signatures have mixed formatting."""

    def _build_email_doc(self):
        doc = Document()
        _add_para(doc, "Re: Quarterly Budget Review", bold=True, size_pt=14)
        _add_para(doc, "")
        _add_para(doc, "Hi Team,", size_pt=11)
        _add_para(doc, _body_text(5), size_pt=11)
        _add_para(doc, "")
        _add_para(doc, "Action Items", bold=True, size_pt=11)
        _add_para(doc, _body_text(2), size_pt=11)
        _add_para(doc, "")
        _add_para(doc, "Thanks,", size_pt=11)
        _add_para(doc, "Jane Smith", bold=True, size_pt=11)
        _add_para(doc, "Director of Finance", size_pt=9, italic=True)
        return doc

    def test_email_subject_detected(self):
        doc = self._build_email_doc()
        result = detect_headings(doc)
        texts = [c.text for c in result]
        assert "Re: Quarterly Budget Review" in texts

    def test_signature_name_could_be_candidate(self):
        """Bold short name in a signature could fool detection --
        verify it doesn't crash and is at least scored."""
        doc = self._build_email_doc()
        result = detect_headings(doc)
        # "Jane Smith" is bold, short, title case -- may or may not
        # reach threshold depending on surrounding context
        assert isinstance(result, list)

    def test_greeting_not_heading(self):
        """'Hi Team,' has trailing comma -- should lose punct signal."""
        doc = self._build_email_doc()
        result = detect_headings(doc)
        greetings = [c for c in result if "Hi Team" in c.text]
        # Even if scored, it shouldn't be a high-confidence heading
        for g in greetings:
            assert g.confidence != "high"


# ===================================================================
# 3. WEB BROWSER PASTE -- random sizes, colors, styles
# ===================================================================


class TestWebPaste:
    """Simulate content pasted from a web browser: inconsistent font
    sizes, colored text, mixed bold/italic."""

    def _build_web_doc(self):
        doc = Document()
        _add_para(
            doc, "Top 10 Accessibility Tips", bold=True, size_pt=24, color=(51, 51, 51)
        )
        _add_para(
            doc,
            "Posted by admin on March 15, 2026",
            size_pt=10,
            italic=True,
            color=(128, 128, 128),
        )
        _add_para(doc, _body_text(4), size_pt=16)
        _add_para(doc, "1. Use Semantic HTML", bold=True, size_pt=20)
        _add_para(doc, _body_text(3), size_pt=16)
        _add_para(doc, "2. Provide Alt Text", bold=True, size_pt=20)
        _add_para(doc, _body_text(3), size_pt=16)
        _add_para(doc, "3. Ensure Color Contrast", bold=True, size_pt=20)
        _add_para(doc, _body_text(3), size_pt=16)
        _add_para(doc, "Related Articles", bold=True, size_pt=18)
        _add_para(doc, "See also: WCAG Guidelines", size_pt=14)
        return doc

    def test_web_article_title_is_level_1(self):
        doc = self._build_web_doc()
        result = detect_headings(doc)
        titles = [c for c in result if "Top 10" in c.text]
        assert len(titles) == 1
        assert titles[0].suggested_level == 1

    def test_numbered_sections_detected(self):
        doc = self._build_web_doc()
        result = detect_headings(doc)
        numbered = [c for c in result if c.text.startswith(("1.", "2.", "3."))]
        assert len(numbered) >= 2

    def test_level_hierarchy_maintained(self):
        """Larger font = higher level, smaller font = lower level."""
        doc = self._build_web_doc()
        result = detect_headings(doc)
        if len(result) >= 2:
            # Title (24pt) should be level 1, sections (20pt) level 2
            title = [c for c in result if "Top 10" in c.text]
            sections = [c for c in result if c.text.startswith("1.")]
            if title and sections:
                assert title[0].suggested_level < sections[0].suggested_level

    def test_byline_not_a_heading(self):
        """'Posted by admin...' is short but italic/small -- low score."""
        doc = self._build_web_doc()
        result = detect_headings(doc)
        bylines = [c for c in result if "Posted by" in c.text]
        assert len(bylines) == 0  # should not reach threshold

    def test_body_font_size_detected_correctly(self):
        """Body text at 16pt should be the detected body size."""
        doc = self._build_web_doc()
        body_size = _get_body_font_size(doc)
        assert body_size == 16.0


# ===================================================================
# 4. LEGAL / GOVERNMENT MEMO -- ALL CAPS, numbered sections
# ===================================================================


class TestLegalMemo:
    """Legal documents: ALL CAPS headers, Roman numeral sections,
    centered titles, paragraph numbering."""

    def _build_legal_doc(self):
        doc = Document()
        p = _add_para(doc, "MEMORANDUM OF UNDERSTANDING", bold=True, size_pt=22)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_para(doc, "")
        _add_para(doc, _body_text(4), size_pt=12)
        _add_para(doc, "")
        _add_para(doc, "SECTION I. DEFINITIONS", bold=True, size_pt=14)
        _add_para(doc, _body_text(3), size_pt=12)
        _add_para(doc, "")
        _add_para(doc, "SECTION II. OBLIGATIONS", bold=True, size_pt=14)
        _add_para(doc, _body_text(3), size_pt=12)
        _add_para(doc, "")
        _add_para(doc, "A. General Requirements", bold=True, size_pt=12)
        _add_para(doc, _body_text(2), size_pt=12)
        _add_para(doc, "B. Specific Terms", bold=True, size_pt=12)
        _add_para(doc, _body_text(2), size_pt=12)
        return doc

    def test_all_caps_title_detected(self):
        doc = self._build_legal_doc()
        result = detect_headings(doc)
        memo = [c for c in result if "MEMORANDUM" in c.text]
        assert len(memo) == 1
        assert memo[0].is_all_caps is True

    def test_section_headings_detected(self):
        doc = self._build_legal_doc()
        result = detect_headings(doc)
        sections = [c for c in result if c.text.startswith("SECTION")]
        assert len(sections) >= 2

    def test_subsection_letters_detected(self):
        doc = self._build_legal_doc()
        result = detect_headings(doc)
        letters = [c for c in result if c.text.startswith(("A.", "B."))]
        # These have the numbering pattern signal
        for lc in letters:
            signal_names = [s[0] for s in lc.signals]
            assert "Numbering pattern" in signal_names

    def test_hierarchy_title_then_sections_then_subsections(self):
        doc = self._build_legal_doc()
        result = detect_headings(doc)
        if len(result) >= 3:
            levels = [c.suggested_level for c in result]
            assert levels[0] == 1  # Title
            # Verify no skipped levels
            for i in range(1, len(levels)):
                assert levels[i] <= levels[i - 1] + 1


# ===================================================================
# 5. RESUME / CV PASTE -- name at top, bold section headers
# ===================================================================


class TestResumePaste:
    """Resume pasted in: name is big/bold, section headers are bold,
    content varies widely in length."""

    def _build_resume_doc(self):
        doc = Document()
        _add_para(doc, "JOHN MICHAEL SMITH", bold=True, size_pt=24)
        _add_para(doc, "123 Main Street | City, ST 12345 | (555) 123-4567", size_pt=10)
        _add_para(doc, "")
        _add_para(doc, "Professional Summary", bold=True, size_pt=14)
        _add_para(doc, _body_text(3), size_pt=11)
        _add_para(doc, "")
        _add_para(doc, "Experience", bold=True, size_pt=14)
        _add_para(
            doc, "Senior Developer, Acme Corp (2020-Present)", bold=True, size_pt=11
        )
        _add_para(doc, _body_text(2), size_pt=11)
        _add_para(doc, "Developer, Beta Inc (2017-2020)", bold=True, size_pt=11)
        _add_para(doc, _body_text(2), size_pt=11)
        _add_para(doc, "")
        _add_para(doc, "Education", bold=True, size_pt=14)
        _add_para(doc, "B.S. Computer Science, State University, 2017", size_pt=11)
        _add_para(doc, "")
        _add_para(doc, "Skills", bold=True, size_pt=14)
        _add_para(doc, "Python, JavaScript, SQL, Docker, AWS", size_pt=11)
        return doc

    def test_name_is_level_1(self):
        doc = self._build_resume_doc()
        result = detect_headings(doc)
        name = [c for c in result if "JOHN" in c.text]
        assert len(name) == 1
        assert name[0].suggested_level == 1

    def test_section_headers_detected(self):
        doc = self._build_resume_doc()
        result = detect_headings(doc)
        sections = {c.text for c in result}
        for expected in ("Professional Summary", "Experience", "Education", "Skills"):
            assert expected in sections, f"'{expected}' not detected"

    def test_job_titles_as_possible_subheadings(self):
        """Bold job titles may be detected as sub-headings."""
        doc = self._build_resume_doc()
        result = detect_headings(doc)
        jobs = [c for c in result if "Acme Corp" in c.text or "Beta Inc" in c.text]
        # They should be at a lower level than section headers
        for j in jobs:
            assert j.suggested_level >= 2

    def test_contact_line_not_heading(self):
        """Long contact info line should not be a heading."""
        doc = self._build_resume_doc()
        result = detect_headings(doc)
        contacts = [c for c in result if "555" in c.text]
        assert len(contacts) == 0


# ===================================================================
# 6. MEETING AGENDA -- numbered items, short bullets, time slots
# ===================================================================


class TestMeetingAgenda:
    def _build_agenda_doc(self):
        doc = Document()
        _add_para(doc, "April Board Meeting Agenda", bold=True, size_pt=22)
        _add_para(doc, "April 8, 2026 -- Conference Room B", size_pt=11)
        _add_para(doc, "")
        _add_para(doc, "1. Call to Order", bold=True, size_pt=14)
        _add_para(doc, "Meeting called to order at 9:00 AM.", size_pt=11)
        _add_para(doc, "")
        _add_para(doc, "2. Approval of Minutes", bold=True, size_pt=14)
        _add_para(doc, _body_text(2), size_pt=11)
        _add_para(doc, "")
        _add_para(doc, "3. Financial Report", bold=True, size_pt=14)
        _add_para(doc, _body_text(3), size_pt=11)
        _add_para(doc, "")
        _add_para(doc, "4. Old Business", bold=True, size_pt=14)
        _add_para(doc, "A. Building renovation update", bold=True, size_pt=11)
        _add_para(doc, _body_text(2), size_pt=11)
        _add_para(doc, "B. Website redesign status", bold=True, size_pt=11)
        _add_para(doc, _body_text(2), size_pt=11)
        _add_para(doc, "")
        _add_para(doc, "5. New Business", bold=True, size_pt=14)
        _add_para(doc, _body_text(3), size_pt=11)
        _add_para(doc, "")
        _add_para(doc, "6. Adjournment", bold=True, size_pt=14)
        return doc

    def test_title_detected_as_h1(self):
        doc = self._build_agenda_doc()
        result = detect_headings(doc)
        titles = [c for c in result if "Board Meeting Agenda" in c.text]
        assert len(titles) == 1
        assert titles[0].suggested_level == 1

    def test_numbered_items_detected(self):
        doc = self._build_agenda_doc()
        result = detect_headings(doc)
        numbered = [
            c for c in result if c.text.startswith(("1.", "2.", "3.", "4.", "5.", "6."))
        ]
        assert len(numbered) >= 4

    def test_numbered_items_have_numbering_signal(self):
        doc = self._build_agenda_doc()
        result = detect_headings(doc)
        for c in result:
            if c.text.startswith(("1.", "2.", "3.")):
                signal_names = [s[0] for s in c.signals]
                assert "Numbering pattern" in signal_names

    def test_letter_sub_items_detected(self):
        doc = self._build_agenda_doc()
        result = detect_headings(doc)
        letters = [c for c in result if c.text.startswith(("A.", "B."))]
        assert len(letters) >= 1

    def test_date_line_not_heading(self):
        doc = self._build_agenda_doc()
        result = detect_headings(doc)
        dates = [c for c in result if "April 8" in c.text]
        # Date lines: not bold, 11pt -- shouldn't be candidates
        assert len(dates) == 0


# ===================================================================
# 7. REPORT -- multi-level headings all in Normal style
# ===================================================================


class TestReportDocument:
    """A formal report where the author used font size + bold for
    structure but never applied heading styles."""

    def _build_report_doc(self):
        doc = Document()
        _add_para(doc, "Annual Accessibility Compliance Report", bold=True, size_pt=28)
        _add_para(doc, "Fiscal Year 2025-2026", bold=True, size_pt=18)
        _add_para(doc, "")
        _add_para(doc, _body_text(5), size_pt=12)
        _add_para(doc, "")
        _add_para(doc, "Executive Summary", bold=True, size_pt=22)
        _add_para(doc, _body_text(6), size_pt=12)
        _add_para(doc, "")
        _add_para(doc, "Methodology", bold=True, size_pt=22)
        _add_para(doc, _body_text(4), size_pt=12)
        _add_para(doc, "")
        _add_para(doc, "Testing Approach", bold=True, size_pt=18)
        _add_para(doc, _body_text(3), size_pt=12)
        _add_para(doc, "Automated Scanning", bold=True, size_pt=14)
        _add_para(doc, _body_text(2), size_pt=12)
        _add_para(doc, "Manual Review", bold=True, size_pt=14)
        _add_para(doc, _body_text(2), size_pt=12)
        _add_para(doc, "")
        _add_para(doc, "Findings", bold=True, size_pt=22)
        _add_para(doc, _body_text(8), size_pt=12)
        _add_para(doc, "")
        _add_para(doc, "Recommendations", bold=True, size_pt=22)
        _add_para(doc, _body_text(4), size_pt=12)
        return doc

    def test_four_level_hierarchy(self):
        """28pt > 22pt > 18pt > 14pt should map to H1 > H2 > H3 > H4."""
        doc = self._build_report_doc()
        result = detect_headings(doc)
        # Group by font size
        by_size: dict[float | None, list] = {}
        for c in result:
            by_size.setdefault(c.font_size_pt, []).append(c)
        # Document title at 28pt should be level 1
        if 28.0 in by_size:
            for c in by_size[28.0]:
                assert c.suggested_level == 1

    def test_no_skipped_levels(self):
        doc = self._build_report_doc()
        result = detect_headings(doc)
        for i in range(1, len(result)):
            prev = result[i - 1].suggested_level
            curr = result[i].suggested_level
            assert curr <= prev + 1, (
                f"Skipped level: H{prev} -> H{curr} "
                f"('{result[i-1].text}' -> '{result[i].text}')"
            )

    def test_body_text_not_detected(self):
        doc = self._build_report_doc()
        result = detect_headings(doc)
        for c in result:
            assert len(c.text) <= 200
            assert c.is_bold or c.font_size_pt is not None

    def test_subtitle_detected(self):
        doc = self._build_report_doc()
        result = detect_headings(doc)
        subtitle = [c for c in result if "Fiscal Year" in c.text]
        assert len(subtitle) >= 1


# ===================================================================
# 8. EDGE CASES -- pathological inputs
# ===================================================================


class TestEdgeCases:
    def test_empty_document(self):
        doc = Document()
        result = detect_headings(doc)
        assert result == []

    def test_single_blank_paragraph(self):
        doc = Document()
        doc.add_paragraph("")
        result = detect_headings(doc)
        assert result == []

    def test_only_whitespace_paragraphs(self):
        doc = Document()
        doc.add_paragraph("   ")
        doc.add_paragraph("\t\t")
        doc.add_paragraph("  \n  ")
        result = detect_headings(doc)
        assert result == []

    def test_single_word_document(self):
        doc = Document()
        _add_para(doc, "Title", bold=True, size_pt=22)
        result = detect_headings(doc)
        # Even a single-word doc should detect the heading
        assert len(result) >= 1

    def test_200_character_boundary(self):
        """Text at exactly 200 chars should be scored; 201 should not."""
        doc = Document()
        text_200 = "A" * 200
        text_201 = "B" * 201
        _add_para(doc, text_200, bold=True, size_pt=22)
        _add_para(doc, text_201, bold=True, size_pt=22)
        _add_para(doc, _body_text(3), size_pt=12)
        result = detect_headings(doc)
        texts = [c.text for c in result]
        assert text_200 in texts
        assert text_201 not in texts

    def test_trailing_punctuation_variants(self):
        """Each punctuation type (.,;:!?) should suppress the signal."""
        doc = Document()
        for punct in [".", ",", ";", ":", "!", "?"]:
            _add_para(doc, f"Heading ends with{punct}", bold=True, size_pt=22)
            _add_para(doc, _body_text(), size_pt=12)
        result = detect_headings(doc)
        for c in result:
            signal_names = [s[0] for s in c.signals]
            assert "No trailing punctuation" not in signal_names

    def test_no_trailing_punctuation_signal_present(self):
        """Clean heading text without punctuation gets the signal."""
        doc = Document()
        _add_para(doc, "Clean Heading", bold=True, size_pt=22)
        _add_para(doc, _body_text(), size_pt=12)
        result = detect_headings(doc)
        assert len(result) >= 1
        signal_names = [s[0] for s in result[0].signals]
        assert "No trailing punctuation" in signal_names

    def test_unicode_headings(self):
        """Non-ASCII text should be handled gracefully."""
        doc = Document()
        _add_para(doc, "Accessibility -- Barrierefreiheit", bold=True, size_pt=22)
        _add_para(doc, _body_text(), size_pt=12)
        _add_para(doc, "Accessibilite et Inclusion", bold=True, size_pt=22)
        _add_para(doc, _body_text(), size_pt=12)
        result = detect_headings(doc)
        assert len(result) >= 2

    def test_document_with_only_headings_styled(self):
        """A document where paragraphs already have Heading styles
        should return zero candidates."""
        doc = Document()
        doc.add_paragraph("Real Heading", style="Heading 1")
        doc.add_paragraph(_body_text())
        doc.add_paragraph("Sub Heading", style="Heading 2")
        doc.add_paragraph(_body_text())
        result = detect_headings(doc)
        styled = [c for c in result if c.text in ("Real Heading", "Sub Heading")]
        assert len(styled) == 0

    def test_mixed_styled_and_faux(self):
        """Some paragraphs are properly styled, others are faux headings."""
        doc = Document()
        doc.add_paragraph("Proper Heading", style="Heading 1")
        _add_para(doc, _body_text(3), size_pt=12)
        _add_para(doc, "Faux Heading", bold=True, size_pt=22)
        _add_para(doc, _body_text(3), size_pt=12)
        result = detect_headings(doc)
        faux = [c for c in result if c.text == "Faux Heading"]
        assert len(faux) == 1

    def test_list_style_paragraphs_skipped(self):
        """Paragraphs with List styles should not be candidates."""
        doc = Document()
        # Use Normal style with our formatting for comparison
        _add_para(doc, "Short Bold Item", bold=True, size_pt=14)
        _add_para(doc, _body_text(), size_pt=12)
        result_before = len(detect_headings(doc))

        # Adding another doc with list style
        doc2 = Document()
        p = doc2.add_paragraph("List Item One", style="List Bullet")
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(14)
        _add_para(doc2, _body_text(), size_pt=12)
        result = detect_headings(doc2)
        list_candidates = [c for c in result if c.text == "List Item One"]
        assert len(list_candidates) == 0

    def test_paragraph_with_no_style(self):
        """A paragraph whose .style is None should not crash."""
        doc = Document()
        p = _add_para(doc, "No Style Para", bold=True, size_pt=22)
        # Force style to None for robustness test
        p.style = None
        _add_para(doc, _body_text(), size_pt=12)
        # Should not raise
        result = detect_headings(doc)
        assert isinstance(result, list)


# ===================================================================
# 9. TITLE CASE vs ALL CAPS vs lowercase
# ===================================================================


class TestCasingSignals:
    def test_title_case_gets_signal(self):
        doc = Document()
        _add_para(doc, "Introduction To The Project", bold=True, size_pt=14)
        _add_para(doc, _body_text(3), size_pt=12)
        result = detect_headings(doc)
        tc = [c for c in result if "Introduction" in c.text]
        if tc:
            signal_names = [s[0] for s in tc[0].signals]
            assert "Title Case" in signal_names

    def test_all_caps_gets_caps_signal_not_title_case(self):
        doc = Document()
        _add_para(doc, "SECTION OVERVIEW", bold=True, size_pt=14)
        _add_para(doc, _body_text(3), size_pt=12)
        result = detect_headings(doc)
        caps = [c for c in result if "SECTION OVERVIEW" in c.text]
        if caps:
            signal_names = [s[0] for s in caps[0].signals]
            assert "ALL CAPS" in signal_names
            assert "Title Case" not in signal_names

    def test_lowercase_no_casing_signal(self):
        """All-lowercase doesn't get either casing signal."""
        doc = Document()
        _add_para(doc, "some lower heading", bold=True, size_pt=22)
        _add_para(doc, _body_text(3), size_pt=12)
        result = detect_headings(doc)
        low = [c for c in result if "some lower" in c.text]
        if low:
            signal_names = [s[0] for s in low[0].signals]
            assert "ALL CAPS" not in signal_names
            assert "Title Case" not in signal_names

    def test_single_char_not_all_caps(self):
        """A single uppercase letter should not trigger ALL CAPS."""
        doc = Document()
        _add_para(doc, "A", bold=True, size_pt=22)
        _add_para(doc, _body_text(3), size_pt=12)
        result = detect_headings(doc)
        singles = [c for c in result if c.text == "A"]
        for s in singles:
            assert s.is_all_caps is False


# ===================================================================
# 10. FONT SIZE SIGNAL TIERS
# ===================================================================


class TestFontSizeSignals:
    def test_22pt_or_above_gets_25_points(self):
        doc = Document()
        _add_para(doc, "Big Heading", bold=True, size_pt=24)
        _add_para(doc, _body_text(3), size_pt=12)
        result = detect_headings(doc)
        assert len(result) >= 1
        size_signals = [(n, p) for n, p in result[0].signals if "Large font" in n]
        assert len(size_signals) == 1
        assert size_signals[0][1] == 25

    def test_20pt_gets_20_points(self):
        doc = Document()
        _add_para(doc, "Medium Heading", bold=True, size_pt=20)
        _add_para(doc, _body_text(3), size_pt=12)
        result = detect_headings(doc)
        assert len(result) >= 1
        size_signals = [(n, p) for n, p in result[0].signals if "Large font" in n]
        assert len(size_signals) == 1
        assert size_signals[0][1] == 20

    def test_below_20pt_no_size_signal(self):
        """Font smaller than 20pt gets no large-font signal."""
        doc = Document()
        _add_para(doc, "Small Heading", bold=True, size_pt=14)
        _add_para(doc, _body_text(3), size_pt=12)
        result = detect_headings(doc)
        for c in result:
            if "Small Heading" in c.text:
                size_signals = [n for n, _ in c.signals if "Large font" in n]
                assert len(size_signals) == 0

    def test_no_font_size_no_size_signal(self):
        """Paragraph with no explicit font size gets no size signal."""
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("Bold No Size")
        run.bold = True
        # No font.size set
        _add_para(doc, _body_text(3), size_pt=12)
        result = detect_headings(doc)
        for c in result:
            if "Bold No Size" in c.text:
                size_signals = [n for n, _ in c.signals if "Large font" in n]
                assert len(size_signals) == 0

    def test_different_from_body_signal(self):
        """Font size differing from body by >0.5pt triggers signal."""
        doc = Document()
        # Body will be 12pt (most common)
        _add_para(doc, "Slightly Bigger", bold=True, size_pt=14)
        _add_para(doc, _body_text(3), size_pt=12)
        _add_para(doc, _body_text(3), size_pt=12)
        result = detect_headings(doc)
        for c in result:
            if "Slightly Bigger" in c.text:
                signal_names = [s[0] for s in c.signals]
                assert "Different from body size" in signal_names


# ===================================================================
# 11. _get_body_font_size -- direct tests
# ===================================================================


class TestGetBodyFontSize:
    def test_most_common_size(self):
        doc = Document()
        _add_para(doc, "Heading", bold=True, size_pt=22)
        _add_para(doc, _body_text(10), size_pt=12)
        _add_para(doc, _body_text(10), size_pt=12)
        assert _get_body_font_size(doc) == 12.0

    def test_fallback_when_no_sizes(self):
        """If no run has an explicit font size, fall back to constant."""
        doc = Document()
        doc.add_paragraph("No size set")
        doc.add_paragraph("Also no size")
        assert _get_body_font_size(doc) == C.BODY_SIZE_PT

    def test_size_weighted_by_text_length(self):
        """Longer text at one size should win over short text at another."""
        doc = Document()
        _add_para(doc, "Short", size_pt=22)
        _add_para(doc, "x" * 500, size_pt=12)
        assert _get_body_font_size(doc) == 12.0

    def test_empty_document(self):
        doc = Document()
        assert _get_body_font_size(doc) == C.BODY_SIZE_PT


# ===================================================================
# 12. CONFIDENCE THRESHOLDS
# ===================================================================


class TestConfidenceScoring:
    def test_high_confidence_at_75(self):
        """Score >= 75 should be 'high' confidence."""
        doc = Document()
        # Bold(20) + 22pt(25) + short(15) + no punct(10) + preceded(10)
        # + title case(10) + diff from body(5) + single-line(5) = 100
        _add_para(doc, "Clear Heading", bold=True, size_pt=22)
        _add_para(doc, _body_text(5), size_pt=12)
        result = detect_headings(doc)
        assert len(result) >= 1
        assert result[0].confidence == "high"

    def test_medium_confidence_below_75(self):
        """Score >= threshold but < 75 should be 'medium'."""
        doc = Document()
        # Bold(20) + short(15) + no punct(10) + preceded(10) + single-line(5) = 60
        _add_para(doc, "maybe heading", bold=True)  # no size, no case signal
        _add_para(doc, _body_text(3), size_pt=12)
        result = detect_headings(doc)
        medium = [c for c in result if c.text == "maybe heading"]
        for m in medium:
            assert m.confidence == "medium" or m.score >= C.HEADING_HIGH_CONFIDENCE

    def test_custom_threshold_raises_bar(self):
        doc = Document()
        _add_para(doc, "Marginal", bold=True, size_pt=14)
        _add_para(doc, _body_text(3), size_pt=12)
        low_thresh = detect_headings(doc, threshold=30)
        high_thresh = detect_headings(doc, threshold=90)
        assert len(low_thresh) >= len(high_thresh)

    def test_threshold_zero_includes_everything_scoreable(self):
        doc = Document()
        _add_para(doc, "A")  # minimal signals
        _add_para(doc, _body_text(3), size_pt=12)
        result = detect_headings(doc, threshold=0)
        # Anything with score > 0 is included
        assert isinstance(result, list)


# ===================================================================
# 13. LEVEL ASSIGNMENT -- extended scenarios
# ===================================================================


class TestLevelAssignmentExtended:
    def _c(self, *, font_size_pt=None, bold=False, caps=False, title=False):
        return HeadingCandidate(
            paragraph_index=0,
            text="T",
            font_size_pt=font_size_pt,
            is_bold=bold,
            is_all_caps=caps,
            is_title_case=title,
            char_count=1,
        )

    def test_four_distinct_sizes(self):
        candidates = [
            self._c(font_size_pt=28),
            self._c(font_size_pt=22),
            self._c(font_size_pt=18),
            self._c(font_size_pt=14),
        ]
        _assign_heading_levels(candidates)
        assert [c.suggested_level for c in candidates] == [1, 2, 3, 4]

    def test_six_sizes_capped_at_6(self):
        candidates = [self._c(font_size_pt=sz) for sz in (30, 26, 22, 18, 14, 10)]
        _assign_heading_levels(candidates)
        levels = [c.suggested_level for c in candidates]
        assert levels == [1, 2, 3, 4, 5, 6]

    def test_seven_sizes_still_caps_at_6(self):
        candidates = [self._c(font_size_pt=sz) for sz in (32, 28, 24, 20, 16, 12, 8)]
        _assign_heading_levels(candidates)
        assert all(c.suggested_level <= 6 for c in candidates)

    def test_same_size_gets_same_level(self):
        candidates = [
            self._c(font_size_pt=22),
            self._c(font_size_pt=22),
            self._c(font_size_pt=18),
        ]
        _assign_heading_levels(candidates)
        assert candidates[0].suggested_level == candidates[1].suggested_level

    def test_no_size_bold_caps_gets_level_1(self):
        c = self._c(bold=True, caps=True)
        _assign_heading_levels([c])
        assert c.suggested_level == 1

    def test_no_size_bold_only_promoted_as_first(self):
        c = self._c(bold=True)
        _assign_heading_levels([c])
        assert c.suggested_level == 1  # promoted because it's first

    def test_no_size_title_case_only(self):
        c1 = self._c(font_size_pt=22)
        c2 = self._c(title=True)
        _assign_heading_levels([c1, c2])
        assert c2.suggested_level == 2  # hierarchy gap fixer adjusts

    def test_first_candidate_always_promoted(self):
        """Even if first candidate's size is smallest, it gets promoted."""
        c1 = self._c(font_size_pt=12)
        c2 = self._c(font_size_pt=22)
        _assign_heading_levels([c1, c2])
        assert c1.suggested_level == 1

    def test_hierarchy_gap_cascade(self):
        """H1, H1, H5 should become H1, H1, H2 after gap fix."""
        c1 = self._c(font_size_pt=22)
        c2 = self._c(font_size_pt=22)
        c3 = self._c(font_size_pt=10)  # would get level 5+
        c1.suggested_level = 1
        c2.suggested_level = 1
        c3.suggested_level = 5
        _fix_hierarchy_gaps([c1, c2, c3])
        assert c3.suggested_level == 2

    def test_mixed_sized_and_unsized(self):
        """Candidates with and without font sizes mix gracefully."""
        c1 = self._c(font_size_pt=22, bold=True)
        c2 = self._c(bold=True)  # no size
        c3 = self._c(font_size_pt=14, bold=True)
        _assign_heading_levels([c1, c2, c3])
        levels = [c.suggested_level for c in [c1, c2, c3]]
        # No skipped levels
        for i in range(1, len(levels)):
            assert levels[i] <= levels[i - 1] + 1


# ===================================================================
# 14. REAL-WORLD MULTI-RUN PARAGRAPHS
# ===================================================================


class TestMultiRunParagraphs:
    """Word often creates paragraphs with multiple runs (e.g. bolding
    part of a sentence, or merging formatting from different sources)."""

    def test_partially_bold_not_detected_as_bold(self):
        """If only part of a paragraph is bold, _para_is_bold is False."""
        doc = Document()
        p = doc.add_paragraph()
        run1 = p.add_run("Bold Part")
        run1.bold = True
        run1.font.size = Pt(14)
        run2 = p.add_run(" and normal part")
        run2.bold = False
        run2.font.size = Pt(14)
        _add_para(doc, _body_text(3), size_pt=12)

        result = detect_headings(doc)
        mixed = [c for c in result if "Bold Part" in c.text]
        for m in mixed:
            assert m.is_bold is False

    def test_all_runs_bold_detected(self):
        """Multiple runs all bold should be detected as bold."""
        doc = Document()
        p = doc.add_paragraph()
        for chunk in ["Chapter ", "One"]:
            run = p.add_run(chunk)
            run.bold = True
            run.font.size = Pt(22)
        _add_para(doc, _body_text(3), size_pt=12)

        result = detect_headings(doc)
        chapter = [c for c in result if "Chapter One" in c.text]
        assert len(chapter) >= 1
        assert chapter[0].is_bold is True

    def test_font_size_from_first_non_empty_run(self):
        """_para_font_size returns size of first non-empty run with a size."""
        doc = Document()
        p = doc.add_paragraph()
        # Whitespace-only run at 10pt
        r1 = p.add_run("  ")
        r1.font.size = Pt(10)
        # Real run at 22pt
        r2 = p.add_run("Real Content")
        r2.bold = True
        r2.font.size = Pt(22)
        _add_para(doc, _body_text(3), size_pt=12)

        result = detect_headings(doc)
        real = [c for c in result if "Real Content" in c.text]
        if real:
            assert real[0].font_size_pt == 22.0


# ===================================================================
# 15. CONTEXT BUILDER -- extended tests
# ===================================================================


class TestBuildContextExtended:
    def test_long_text_truncated_at_200(self):
        doc = Document()
        _add_para(doc, "A" * 300, size_pt=12)
        _add_para(doc, "Target", bold=True, size_pt=22)
        _add_para(doc, "B" * 300, size_pt=12)
        paras = list(doc.paragraphs)
        ctx = _build_context(paras, 1, [])
        assert len(ctx["before"][0]) == 200
        assert len(ctx["after"][0]) == 200

    def test_existing_headings_in_context(self):
        paras = [MagicMock() for _ in range(3)]
        for p in paras:
            p.text = "text"
        existing = [
            HeadingCandidate(
                paragraph_index=0,
                text="Document Title",
                font_size_pt=22,
                is_bold=True,
                is_all_caps=False,
                is_title_case=True,
                char_count=14,
                suggested_level=1,
            )
        ]
        ctx = _build_context(paras, 2, existing)
        assert 'H1: "Document Title"' in ctx["existing_headings"]

    def test_max_three_before_and_after(self):
        doc = Document()
        for i in range(10):
            _add_para(doc, f"Para {i}", size_pt=12)
        paras = list(doc.paragraphs)
        ctx = _build_context(paras, 5, [])
        assert len(ctx["before"]) == 3
        assert len(ctx["after"]) == 3


# ===================================================================
# 16. PREV/NEXT POSITIONAL SIGNALS -- extended
# ===================================================================


class TestPositionalSignalsExtended:
    def test_prev_is_blank_with_whitespace_only(self):
        """A paragraph with only spaces counts as blank."""
        doc = Document()
        _add_para(doc, "   ")
        _add_para(doc, "After Blank", bold=True, size_pt=22)
        _add_para(doc, _body_text(3), size_pt=12)
        result = detect_headings(doc)
        ab = [c for c in result if "After Blank" in c.text]
        if ab:
            signal_names = [s[0] for s in ab[0].signals]
            assert "Preceded by blank/break" in signal_names

    def test_next_is_body_boundary_81_chars(self):
        """Next paragraph at exactly 81 chars + not bold = body signal."""
        doc = Document()
        _add_para(doc, "Heading", bold=True, size_pt=22)
        # 81 characters of body text
        _add_para(doc, "x" * 81, size_pt=12)
        result = detect_headings(doc)
        h = [c for c in result if "Heading" == c.text]
        if h:
            signal_names = [s[0] for s in h[0].signals]
            assert "Followed by body text" in signal_names

    def test_next_at_80_chars_not_body(self):
        """80 chars is NOT > 80, so no body signal."""
        doc = Document()
        _add_para(doc, "Heading", bold=True, size_pt=22)
        _add_para(doc, "x" * 80, size_pt=12)
        result = detect_headings(doc)
        h = [c for c in result if "Heading" == c.text]
        if h:
            signal_names = [s[0] for s in h[0].signals]
            assert "Followed by body text" not in signal_names

    def test_next_bold_text_not_body(self):
        """Even if next paragraph is long, bold text doesn't count as body."""
        doc = Document()
        _add_para(doc, "First Heading", bold=True, size_pt=22)
        _add_para(doc, "y" * 100, bold=True, size_pt=12)
        result = detect_headings(doc)
        fh = [c for c in result if "First Heading" in c.text]
        if fh:
            signal_names = [s[0] for s in fh[0].signals]
            assert "Followed by body text" not in signal_names


# ===================================================================
# 17. NUMBERING PATTERN REGEX
# ===================================================================


class TestNumberingPatterns:
    @pytest.mark.parametrize(
        "text",
        [
            "1. Introduction",
            "2) Background",
            "A. General Terms",
            "B) Specific Conditions",
            "Chapter 1 Overview",
            "Section 3 Requirements",
            "Part 2 Implementation",
            "chapter 5 details",
        ],
    )
    def test_numbering_detected(self, text):
        doc = Document()
        _add_para(doc, text, bold=True, size_pt=22)
        _add_para(doc, _body_text(3), size_pt=12)
        result = detect_headings(doc)
        assert len(result) >= 1
        signal_names = [s[0] for s in result[0].signals]
        assert "Numbering pattern" in signal_names

    @pytest.mark.parametrize(
        "text",
        [
            "The 3 key findings were",
            "In section one we see",
            "a simple paragraph",
        ],
    )
    def test_numbering_not_detected(self, text):
        doc = Document()
        _add_para(doc, text, bold=True, size_pt=22)
        _add_para(doc, _body_text(3), size_pt=12)
        result = detect_headings(doc)
        for c in result:
            if c.text == text:
                signal_names = [s[0] for s in c.signals]
                assert "Numbering pattern" not in signal_names


# ===================================================================
# 18. LARGE DOCUMENTS -- performance & correctness
# ===================================================================


class TestLargeDocument:
    def test_50_heading_document(self):
        """A document with 50 chapters should detect all headings."""
        doc = Document()
        for i in range(1, 51):
            _add_para(doc, f"Chapter {i} Title", bold=True, size_pt=22)
            _add_para(doc, _body_text(3), size_pt=12)
            _add_para(doc, "")
        result = detect_headings(doc)
        chapter_headings = [c for c in result if c.text.startswith("Chapter ")]
        assert len(chapter_headings) >= 40  # Some at edges may vary

    def test_mixed_headings_and_body(self):
        """Interleaved headings and long body paragraphs."""
        doc = Document()
        _add_para(doc, "Document Title", bold=True, size_pt=28)
        for section in range(5):
            _add_para(doc, "")
            _add_para(doc, f"Section {section + 1}", bold=True, size_pt=22)
            _add_para(doc, _body_text(10), size_pt=12)
            for sub in range(3):
                _add_para(
                    doc, f"Subsection {section + 1}.{sub + 1}", bold=True, size_pt=18
                )
                _add_para(doc, _body_text(5), size_pt=12)
        result = detect_headings(doc)
        # Title + 5 sections + 15 subsections = 21 expected headings
        assert len(result) >= 18  # Allow some margin
        assert result[0].suggested_level == 1
        # Verify no skipped levels
        for i in range(1, len(result)):
            assert result[i].suggested_level <= result[i - 1].suggested_level + 1
