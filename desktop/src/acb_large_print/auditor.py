"""Audit Word documents for ACB Large Print Guidelines compliance."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# VML namespace not in python-docx's nsmap -- register manually
_VML_NS = "urn:schemas-microsoft-com:vml"
_WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"


def _vml_qn(tag: str) -> str:
    """Build a qualified name for VML namespace tags (v:shape, v:textbox, etc.)."""
    return f"{{{_VML_NS}}}{tag}"


from . import constants as C

# ---------------------------------------------------------------------------
# Module-level measurement tolerances
# ---------------------------------------------------------------------------
# 5% deviation allowed from the expected line-spacing multiple (e.g. 1.5×)
_LINE_SPACING_TOLERANCE: float = 0.05
# 0.1-inch tolerance for page-size checks (US Letter: 8.5 × 11 in)
_PAGE_SIZE_TOLERANCE: float = 0.1


@dataclass
class Finding:
    """A single audit finding."""

    rule_id: str
    severity: C.Severity
    message: str
    location: str = ""
    auto_fixable: bool = True

    @property
    def rule(self) -> C.RuleDef:
        return C.AUDIT_RULES[self.rule_id]

    @property
    def acb_reference(self) -> str:
        return self.rule.acb_reference


@dataclass
class AuditResult:
    """Complete audit results for a document."""

    file_path: str
    findings: list[Finding] = field(default_factory=list)
    total_paragraphs: int = 0
    total_runs: int = 0
    metadata_display: object | None = (
        None  # EpubAccessibilityDisplay when auditing EPUBs
    )

    @property
    def passed(self) -> bool:
        return len(self.findings) == 0

    @property
    def score(self) -> int:
        """Compliance score 0-100.

        Deduct per *unique* rule violation (full severity weight) plus a small
        diminishing-returns penalty for repeat occurrences of the same rule.
        This prevents score inflation when a single category (e.g. wrong font
        family) fires on every paragraph in the document.
        """
        if not self.findings:
            return 100
        deductions = {
            C.Severity.CRITICAL: 15,
            C.Severity.HIGH: 10,
            C.Severity.MEDIUM: 5,
            C.Severity.LOW: 2,
        }
        # Group findings by rule_id so each rule is charged once at full weight,
        # then a capped per-extra-occurrence penalty.
        by_rule: dict[str, list[Finding]] = {}
        for f in self.findings:
            by_rule.setdefault(f.rule_id, []).append(f)

        total = 0
        for rule_id, group in by_rule.items():
            sev = group[0].severity
            base = deductions.get(sev, 5)
            extras = max(0, len(group) - 1)
            # +1pt per extra occurrence, capped at the base weight per rule.
            total += base + min(extras, base)
        return max(0, 100 - total)

    @property
    def grade(self) -> str:
        """Letter grade from score."""
        s = self.score
        if s >= 90:
            return "A"
        if s >= 80:
            return "B"
        if s >= 70:
            return "C"
        if s >= 60:
            return "D"
        return "F"

    def add(self, rule_id: str, message: str, location: str = "") -> None:
        rule = C.AUDIT_RULES[rule_id]
        self.findings.append(
            Finding(
                rule_id=rule_id,
                severity=rule.severity,
                message=message,
                location=location,
                auto_fixable=rule.auto_fixable,
            )
        )

    @property
    def critical_count(self) -> int:
        return sum(1 for f in self.findings if f.severity == C.Severity.CRITICAL)

    @property
    def high_count(self) -> int:
        return sum(1 for f in self.findings if f.severity == C.Severity.HIGH)

    @property
    def medium_count(self) -> int:
        return sum(1 for f in self.findings if f.severity == C.Severity.MEDIUM)

    @property
    def low_count(self) -> int:
        return sum(1 for f in self.findings if f.severity == C.Severity.LOW)


def _effective_font_name(run) -> str | None:
    """Get the effective font name for a run (checks direct + style)."""
    if run.font.name:
        return run.font.name
    if run.style and run.style.font.name:
        return run.style.font.name
    para_style = run._element.getparent()
    if para_style is not None:
        # Walk up to paragraph, then check its style
        pass
    return None


def _effective_font_size(run) -> float | None:
    """Get effective font size in points for a run."""
    if run.font.size is not None:
        return run.font.size.pt
    if run.style and run.style.font.size:
        return run.style.font.size.pt
    return None


def _add_capped_font_family_findings(
    result: AuditResult,
    font_locations: dict[str, list[str]],
    *,
    max_examples: int = 3,
) -> None:
    """Add capped document findings for repeated non-Arial font usage."""
    for font_name, locations in sorted(font_locations.items()):
        for location in locations[:max_examples]:
            result.add(
                "ACB-FONT-FAMILY",
                f"Non-Arial font '{font_name}'",
                location,
            )
        if len(locations) > max_examples:
            result.add(
                "ACB-FONT-FAMILY",
                (
                    f"Non-Arial font '{font_name}' also appears in "
                    f"{len(locations) - max_examples} additional paragraph(s); "
                    f"showing the first {max_examples} locations only"
                ),
                "Document-wide",
            )


def _is_heading_style(style_name: str) -> bool:
    """Check if a style name is a heading style."""
    return bool(re.match(r"^Heading \d+$", style_name))


def _heading_level(style_name: str) -> int:
    """Extract heading level number from style name."""
    match = re.match(r"^Heading (\d+)$", style_name)
    return int(match.group(1)) if match else 0


def _list_level_from_style(style_name: str) -> int | None:
    """Infer list nesting level from built-in Word list style names."""
    if style_name in ("List Bullet", "List Number"):
        return 0
    if style_name in ("List Bullet 2", "List Number 2"):
        return 1
    if style_name in ("List Bullet 3", "List Number 3"):
        return 2
    return None


def _check_document_properties(doc: Document, result: AuditResult) -> None:
    """Check document-level properties."""
    # Title
    title = doc.core_properties.title
    if not title or not title.strip():
        result.add(
            "ACB-DOC-TITLE",
            "Document has no title set in properties",
            "Document Properties",
        )

    # Language
    styles_elem = doc.styles.element
    doc_defaults = styles_elem.find(qn("w:docDefaults"))
    rpr = None
    if doc_defaults is not None:
        rpr_default = doc_defaults.find(qn("w:rPrDefault"))
        if rpr_default is not None:
            rpr = rpr_default.find(qn("w:rPr"))
    if rpr is not None:
        lang = rpr.find(qn("w:lang"))
        if lang is None or not lang.get(qn("w:val")):
            result.add(
                "ACB-DOC-LANGUAGE",
                "Document language is not set",
                "Document Properties",
            )
    else:
        result.add(
            "ACB-DOC-LANGUAGE",
            "Document language is not set",
            "Document Properties",
        )


def _check_page_setup(doc: Document, result: AuditResult) -> None:
    """Check page margins and orientation."""
    for i, section in enumerate(doc.sections):
        loc = f"Section {i + 1}"
        margin_tolerance = Inches(0.05)

        if section.top_margin is not None:
            if abs(section.top_margin - Inches(C.MARGIN_TOP_IN)) > margin_tolerance:
                result.add(
                    "ACB-MARGINS",
                    f"Top margin is {section.top_margin.inches:.2f} inches (expected {C.MARGIN_TOP_IN})",
                    loc,
                )

        if section.bottom_margin is not None:
            if (
                abs(section.bottom_margin - Inches(C.MARGIN_BOTTOM_IN))
                > margin_tolerance
            ):
                result.add(
                    "ACB-MARGINS",
                    f"Bottom margin is {section.bottom_margin.inches:.2f} inches (expected {C.MARGIN_BOTTOM_IN})",
                    loc,
                )

        if section.left_margin is not None:
            if abs(section.left_margin - Inches(C.MARGIN_LEFT_IN)) > margin_tolerance:
                result.add(
                    "ACB-MARGINS",
                    f"Left margin is {section.left_margin.inches:.2f} inches (expected {C.MARGIN_LEFT_IN})",
                    loc,
                )

        if section.right_margin is not None:
            if abs(section.right_margin - Inches(C.MARGIN_RIGHT_IN)) > margin_tolerance:
                result.add(
                    "ACB-MARGINS",
                    f"Right margin is {section.right_margin.inches:.2f} inches (expected {C.MARGIN_RIGHT_IN})",
                    loc,
                )


def _check_styles(
    doc: Document,
    result: AuditResult,
    *,
    style_size_overrides: dict[str, float] | None = None,
) -> None:
    """Check that named styles match ACB specifications."""
    styles_table = C.effective_styles(style_size_overrides)
    for style_name, style_def in styles_table.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue

        loc = f"Style: {style_name}"
        font = style.font
        pf = style.paragraph_format

        # Font family
        if font.name and font.name != C.FONT_FAMILY:
            result.add(
                "ACB-FONT-FAMILY",
                f"{style_name} uses '{font.name}' instead of '{C.FONT_FAMILY}'",
                loc,
            )

        # Font size
        expected_size = style_def.font.size_pt
        if font.size and abs(font.size.pt - expected_size) > 0.5:
            size_rule = "ACB-FONT-SIZE-BODY"
            if style_name == "Heading 1":
                size_rule = "ACB-FONT-SIZE-H1"
            elif style_name == "Heading 2":
                size_rule = "ACB-FONT-SIZE-H2"
            result.add(
                size_rule,
                f"{style_name} is {font.size.pt}pt (expected {expected_size}pt)",
                loc,
            )

        # Italic check
        if font.italic:
            result.add(
                "ACB-NO-ITALIC",
                f"{style_name} style has italic enabled",
                loc,
            )

        # Alignment
        if pf.alignment is not None and pf.alignment != WD_ALIGN_PARAGRAPH.LEFT:
            result.add(
                "ACB-ALIGNMENT",
                f"{style_name} is not flush left",
                loc,
            )


def _check_paragraph_content(
    doc: Document,
    result: AuditResult,
    *,
    list_indent_in: float = C.LIST_INDENT_IN,
    list_level_indents: dict[int, float] | None = None,
    para_indent_in: float = C.PARA_INDENT_IN,
    first_line_indent_in: float = C.FIRST_LINE_INDENT_IN,
    style_size_overrides: dict[str, float] | None = None,
) -> None:
    """Check each paragraph and run for direct formatting violations."""
    prev_heading_level = 0
    paragraphs_since_heading = 0
    heading_texts: dict[int, list[str]] = {}  # level -> list of heading texts
    non_arial_font_locations: dict[str, list[str]] = {}

    for i, para in enumerate(doc.paragraphs):
        result.total_paragraphs += 1
        loc = f"Paragraph {i + 1}"
        style_name = para.style.name if para.style else "Normal"
        para_style_font = para.style.font.name if para.style is not None else None
        text_preview = para.text[:60].strip()
        if text_preview:
            loc = (
                f"Paragraph {i + 1}: '{text_preview}...'"
                if len(para.text) > 60
                else f"Paragraph {i + 1}: '{text_preview}'"
            )

        is_heading = _is_heading_style(style_name)

        # Heading hierarchy check
        if is_heading:
            level = _heading_level(style_name)
            if prev_heading_level > 0 and level > prev_heading_level + 1:
                result.add(
                    "ACB-HEADING-HIERARCHY",
                    f"Heading level skipped: H{prev_heading_level} to H{level}",
                    loc,
                )
            prev_heading_level = level

            # Duplicate heading text check
            heading_text = para.text.strip()
            if heading_text:
                if level not in heading_texts:
                    heading_texts[level] = []
                if heading_text in heading_texts[level]:
                    result.add(
                        "ACB-DUPLICATE-HEADING",
                        f"Duplicate H{level} heading: '{heading_text[:60]}'",
                        loc,
                    )
                heading_texts[level].append(heading_text)

        # Direct alignment override
        if para.paragraph_format.alignment is not None:
            if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                msg = "Direct alignment override (not flush left)"
                if is_heading:
                    msg = "Direct alignment override on heading paragraph"
                result.add(
                    "ACB-ALIGNMENT",
                    msg,
                    loc,
                )

        # List indentation check
        is_list = style_name in (
            "List Bullet",
            "List Number",
            "List Bullet 2",
            "List Number 2",
            "List Bullet 3",
            "List Number 3",
        )
        if is_list:
            pf = para.paragraph_format
            actual_indent = pf.left_indent.inches if pf.left_indent else 0.0
            level = _list_level_from_style(style_name)
            expected_indent = list_indent_in
            if list_level_indents is not None and level is not None:
                expected_indent = list_level_indents.get(level, list_indent_in)
            indent_tolerance = 0.05
            if abs(actual_indent - expected_indent) > indent_tolerance:
                result.add(
                    "ACB-LIST-INDENT",
                    f'List indent is {actual_indent:.2f}" (expected {expected_indent:.2f}")',
                    loc,
                )

        # Non-list, non-heading paragraph indentation checks
        if not is_list and not is_heading and para.text.strip():
            pf = para.paragraph_format
            indent_tolerance = 0.05

            # Left indent check
            actual_left = pf.left_indent.inches if pf.left_indent else 0.0
            if abs(actual_left - para_indent_in) > indent_tolerance:
                # Blockquote-style: body-length text with significant indent
                if (
                    actual_left > para_indent_in + 0.25 + indent_tolerance
                    and len(para.text.strip()) > 40
                ):
                    result.add(
                        "ACB-BLOCKQUOTE-INDENT",
                        f'Blockquote-style indent ({actual_left:.2f}") on body text',
                        loc,
                    )
                else:
                    result.add(
                        "ACB-PARA-INDENT",
                        f'Paragraph left indent is {actual_left:.2f}" (expected {para_indent_in:.2f}")',
                        loc,
                    )

            # First-line indent check
            actual_fli = pf.first_line_indent.inches if pf.first_line_indent else 0.0
            if abs(actual_fli - first_line_indent_in) > indent_tolerance:
                result.add(
                    "ACB-FIRST-LINE-INDENT",
                    f'First-line indent is {actual_fli:.2f}" (expected {first_line_indent_in:.2f}")',
                    loc,
                )

        paragraph_non_arial_fonts: set[str] = set()

        # Check each run
        for run in para.runs:
            result.total_runs += 1

            # Italic at run level
            if run.font.italic:
                result.add(
                    "ACB-NO-ITALIC",
                    f"Italic text found: '{run.text[:40]}'",
                    loc,
                )

            # Bold in body text (not headings)
            if run.font.bold and not is_heading and run.text.strip():
                # Only flag if the bold run is a substantial portion of the paragraph
                if len(run.text.strip()) > 5:
                    result.add(
                        "ACB-BOLD-HEADINGS-ONLY",
                        f"Bold used in body text: '{run.text[:40]}'",
                        loc,
                    )

            # Font family override at run level
            effective_font_name = _effective_font_name(run)
            if effective_font_name and effective_font_name != C.FONT_FAMILY:
                style_covers_font = (
                    style_name in C.ACB_STYLES
                    and para_style_font == effective_font_name
                    and para_style_font != C.FONT_FAMILY
                )
                if (
                    not style_covers_font
                    and effective_font_name not in paragraph_non_arial_fonts
                ):
                    paragraph_non_arial_fonts.add(effective_font_name)
                    non_arial_font_locations.setdefault(effective_font_name, []).append(loc)

            # Font size below user/ACB minimum at run level
            _min_pt = C.effective_min_body_pt(style_size_overrides)
            if run.font.size and run.font.size.pt < _min_pt - 0.5:
                result.add(
                    "ACB-FONT-SIZE-BODY",
                    f"Font size {run.font.size.pt}pt below {_min_pt}pt minimum",
                    loc,
                )

            # All caps at run level
            if run.font.all_caps:
                result.add(
                    "ACB-NO-ALLCAPS",
                    f"ALL CAPS formatting on: '{run.text[:40]}'",
                    loc,
                )

        # Check for all-caps text content (not formatting, but typed in caps)
        if (
            para.text.strip()
            and para.text.strip().isupper()
            and len(para.text.strip()) > 3
        ):
            if is_heading:
                result.add(
                    "ACB-NO-ALLCAPS",
                    f"Heading text is ALL CAPS",
                    loc,
                )

        # Repeated blank characters (consecutive spaces for visual layout)
        if "  " in para.text and not is_heading:
            # Flag runs of 3+ spaces as likely layout abuse
            space_runs = re.findall(r" {3,}", para.text)
            if space_runs:
                result.add(
                    "ACB-REPEATED-SPACES",
                    f"Paragraph contains {len(space_runs)} run(s) of consecutive spaces",
                    loc,
                )

        # Fake list detection (manually typed bullets/numbers)
        stripped = para.text.strip()
        if stripped and not is_heading:
            # Typed bullet characters: bullet, en-dash, em-dash, heavy bullet
            if stripped[0] in "\u2022\u2023\u25cf\u25cb\u25e6\u2043\u2219\u00b7":
                result.add(
                    "ACB-FAKE-LIST",
                    f"Manually typed bullet character '{stripped[0]}' used instead of built-in list style",
                    loc,
                )
            # Typed numbered list: "1." "2)" "a." at start
            elif (
                re.match(r"^(\d{1,3}[\.\)]\s|[a-z][\.\)]\s)", stripped)
                and style_name == "Normal"
            ):
                result.add(
                    "ACB-FAKE-LIST",
                    f"Manually typed list numbering detected",
                    loc,
                )

        # Track paragraphs since last heading for long-section check
        if is_heading:
            paragraphs_since_heading = 0
        else:
            paragraphs_since_heading += 1
            if paragraphs_since_heading == 21:
                result.add(
                    "ACB-LONG-SECTION",
                    f"20+ paragraphs without a heading (last heading was before paragraph {i + 1 - 20})",
                    loc,
                )

    _add_capped_font_family_findings(result, non_arial_font_locations)


def _check_hyphenation(doc: Document, result: AuditResult) -> None:
    """Check whether auto-hyphenation is enabled."""
    settings = doc.settings.element
    auto_hyph = settings.find(qn("w:autoHyphenation"))
    if auto_hyph is not None:
        val = auto_hyph.get(qn("w:val"), "true")
        if val.lower() not in ("false", "0", "off"):
            result.add(
                "ACB-NO-HYPHENATION",
                "Automatic hyphenation is enabled",
                "Document Settings",
            )


def _check_line_spacing(doc: Document, result: AuditResult) -> None:
    """Check paragraph line spacing at the style and paragraph levels.

    ACB guidelines require line spacing to be set to the prescribed multiple
    (``C.LINE_SPACING_MULTIPLE``, typically 1.5× or "Multiple 1.5").  The
    check inspects:
      1. Named styles (Normal, Heading 1-6, Body Text, etc.)
      2. Direct paragraph formatting overrides
    """
    from docx.oxml.ns import qn as _qn
    from docx.shared import Pt

    # Helper: read a spacing element's w:line and w:lineRule
    def _spacing_vals(elem):
        sp = elem.find(_qn("w:spacing"))
        if sp is None:
            return None, None
        line = sp.get(_qn("w:line"))
        rule = sp.get(_qn("w:lineRule"), "auto")
        return line, rule

    _TOLERANCE = _LINE_SPACING_TOLERANCE  # allow 5% deviation from expected multiple
    _EXPECTED = C.LINE_SPACING_MULTIPLE  # e.g. 1.5
    # Word stores "Multiple N" as N * 240 twips (w:lineRule="auto")
    _EXPECTED_TWIPS = int(_EXPECTED * 240)

    def _is_bad_spacing(line_str: str | None, rule: str | None) -> bool:
        if line_str is None:
            return False
        rule = (rule or "auto").lower()
        try:
            line_val = int(line_str)
        except (ValueError, TypeError):
            return False
        if rule in ("auto", "atLeast", ""):
            # "auto" means multiple; stored as N * 240
            actual_multiple = line_val / 240.0
            return abs(actual_multiple - _EXPECTED) > _TOLERANCE
        # "exact" means exact pt spacing -- always flag as wrong
        return True

    # --- 1. Check named styles ---
    for style in doc.styles:
        try:
            pPr = style.element.find(_qn("w:pPr"))
        except Exception:
            continue
        if pPr is None:
            continue
        line_str, rule = _spacing_vals(pPr)
        if _is_bad_spacing(line_str, rule):
            try:
                actual = int(line_str) / 240.0
            except Exception:
                actual = 0.0
            result.add(
                "ACB-LINE-SPACING",
                f"Style '{style.name}' line spacing is {actual:.2f}× (expected {_EXPECTED}×)",
                f"Style: {style.name}",
            )

    # --- 2. Check paragraph-level overrides ---
    for i, para in enumerate(doc.paragraphs):
        pPr = para._p.find(_qn("w:pPr"))
        if pPr is None:
            continue
        line_str, rule = _spacing_vals(pPr)
        if _is_bad_spacing(line_str, rule):
            try:
                actual = int(line_str) / 240.0
            except Exception:
                actual = 0.0
            preview = para.text[:60].strip()
            loc = (
                f"Paragraph {i + 1}: '{preview}'"
                if preview
                else f"Paragraph {i + 1}"
            )
            result.add(
                "ACB-LINE-SPACING",
                f"Direct paragraph line spacing is {actual:.2f}× (expected {_EXPECTED}×)",
                loc,
            )


def _check_widow_orphan(doc: Document, result: AuditResult) -> None:
    """Check that widow and orphan protection is enabled.

    Inspects both named styles (Normal, Heading 1-6, Body Text) and direct
    paragraph formatting overrides.  Widow/orphan control keeps at least two
    lines of a paragraph together at page boundaries.
    """
    from docx.oxml.ns import qn as _qn

    def _widow_orphan_off(elem) -> bool:
        """Return True if widow/orphan control is explicitly disabled."""
        pPr = elem.find(_qn("w:pPr"))
        if pPr is None:
            return False
        wo = pPr.find(_qn("w:widowControl"))
        if wo is None:
            return False  # absent = Word default (on)
        val = wo.get(_qn("w:val"), "true").lower()
        return val in ("false", "0", "off")

    # --- 1. Named styles ---
    for style in doc.styles:
        try:
            if _widow_orphan_off(style.element):
                result.add(
                    "ACB-WIDOW-ORPHAN",
                    f"Style '{style.name}' has widow/orphan control disabled",
                    f"Style: {style.name}",
                )
        except Exception:
            continue

    # --- 2. Paragraph-level overrides ---
    for i, para in enumerate(doc.paragraphs):
        if _widow_orphan_off(para._p):
            preview = para.text[:60].strip()
            loc = (
                f"Paragraph {i + 1}: '{preview}'"
                if preview
                else f"Paragraph {i + 1}"
            )
            result.add(
                "ACB-WIDOW-ORPHAN",
                "Widow/orphan control disabled on paragraph",
                loc,
            )


def _check_page_layout(doc: Document, result: AuditResult) -> None:
    """Check page size (US Letter) and multi-column layout.

    ACB large print documents must use US Letter (8.5 × 11 in) in portrait
    orientation and must not use multi-column layouts.
    """
    _US_LETTER_WIDTH = 8.5   # inches
    _US_LETTER_HEIGHT = 11.0  # inches
    _TOLERANCE = _PAGE_SIZE_TOLERANCE    # 0.1-inch tolerance

    for i, section in enumerate(doc.sections):
        loc = f"Section {i + 1}"

        # --- Page size ---
        try:
            page_width_in = section.page_width.inches
            page_height_in = section.page_height.inches
            # Accept both portrait and landscape orientations
            dims = sorted([page_width_in, page_height_in])
            expected = sorted([_US_LETTER_WIDTH, _US_LETTER_HEIGHT])
            if (
                abs(dims[0] - expected[0]) > _TOLERANCE
                or abs(dims[1] - expected[1]) > _TOLERANCE
            ):
                result.add(
                    "ACB-PAGE-SIZE",
                    f"Page size is {page_width_in:.2f}\" × {page_height_in:.2f}\" "
                    f"(expected {_US_LETTER_WIDTH}\" × {_US_LETTER_HEIGHT}\")",
                    loc,
                )
        except Exception:
            pass

        # --- Multi-column layout ---
        try:
            sectPr = section._sectPr
            from docx.oxml.ns import qn as _qn
            cols_elem = sectPr.find(_qn("w:cols"))
            if cols_elem is not None:
                num_cols_str = cols_elem.get(_qn("w:num"), "1")
                try:
                    num_cols = int(num_cols_str)
                except (ValueError, TypeError):
                    num_cols = 1
                if num_cols > 1:
                    result.add(
                        "ACB-MULTIPLE-COLUMNS",
                        f"Section uses {num_cols}-column layout — ACB requires single column",
                        loc,
                    )
        except Exception:
            pass


def _check_run_formatting(doc: Document, result: AuditResult) -> None:
    """Check run-level formatting: font color and strikethrough.

    ACB large print documents must use black (automatic) text color to ensure
    WCAG 4.5:1 contrast.  Strikethrough text should be removed as it is
    difficult to read for low-vision users.

    Findings are capped per-document to avoid flooding the report with
    hundreds of identical issues.
    """
    from docx.oxml.ns import qn as _qn

    _AUTO_COLOR = "auto"
    _BLACK_HEX = "000000"
    _MAX_COLOR_FINDINGS = 15
    _MAX_STRIKE_FINDINGS = 15
    color_count = 0
    strike_count = 0

    for i, para in enumerate(doc.paragraphs):
        loc_prefix = f"Paragraph {i + 1}"
        preview = para.text[:40].strip()
        loc = f"{loc_prefix}: '{preview}'" if preview else loc_prefix

        for run in para.runs:
            if not run.text.strip():
                continue

            # --- Font color ---
            if color_count < _MAX_COLOR_FINDINGS:
                rPr = run._r.find(_qn("w:rPr"))
                if rPr is not None:
                    color_elem = rPr.find(_qn("w:color"))
                    if color_elem is not None:
                        color_val = color_elem.get(_qn("w:val"), _AUTO_COLOR)
                        if color_val.lower() not in (_AUTO_COLOR, _BLACK_HEX, "000000"):
                            result.add(
                                "ACB-FONT-COLOR",
                                f"Colored text (#{color_val}): '{run.text[:30]}'",
                                loc,
                            )
                            color_count += 1

            # --- Strikethrough ---
            if strike_count < _MAX_STRIKE_FINDINGS:
                if run.font.strike or run.font.double_strike:
                    result.add(
                        "ACB-STRIKETHROUGH",
                        f"Strikethrough text: '{run.text[:30]}'",
                        loc,
                    )
                    strike_count += 1


def _check_page_numbers(doc: Document, result: AuditResult) -> None:
    """Check for page number fields in footer."""
    has_page_numbers = False
    for section in doc.sections:
        footer = section.footer
        if footer is None:
            continue
        footer_xml = footer._element.xml
        if "PAGE" in footer_xml or "w:fldChar" in footer_xml:
            has_page_numbers = True
            break

    if not has_page_numbers:
        result.add(
            "ACB-PAGE-NUMBERS",
            "No page numbers found in document footer",
            "Footer",
        )


# ---------------------------------------------------------------------------
# Microsoft Accessibility Checker (MSAC) rule checks
# ---------------------------------------------------------------------------

_NON_DESCRIPTIVE_LINK_RE = re.compile(
    r"^(click here|here|read more|learn more|more|link|this link|details|info|"
    r"download|submit|go|open|visit|see|view|page|website)\.?$",
    re.IGNORECASE,
)

_URL_LIKE_RE = re.compile(r"^https?://", re.IGNORECASE)


def _check_alt_text(doc: Document, result: AuditResult) -> None:
    """Check images and shapes for missing alt text."""
    body = doc.element.body
    # Inline images (w:drawing > wp:inline and wp:anchor with pic:pic)
    for drawing in body.iter(qn("w:drawing")):
        # Look for docPr which holds the alt text
        for doc_pr in drawing.iter(qn("wp:docPr")):
            name = doc_pr.get("name", "")
            descr = doc_pr.get("descr", "").strip()
            # Also check for decorative mark (a14:decorative)
            decorative = False
            for ext_lst in drawing.iter(qn("a:extLst")):
                xml_str = ext_lst.xml if hasattr(ext_lst, "xml") else ""
                if "decorative" in str(xml_str).lower():
                    decorative = True
                    break
            if not descr and not decorative:
                result.add(
                    "ACB-MISSING-ALT-TEXT",
                    f"Image or shape '{name}' has no alternative text",
                    f"Inline image: {name}",
                )

    # Also check v:shape elements (legacy shapes)
    for shape in body.iter(_vml_qn("shape")):
        alt = shape.get("alt", None)
        name = shape.get("id", "unknown shape")
        # In VML, alt="" means explicitly decorative.
        # Missing alt means no alt text.
        if alt is None:
            result.add(
                "ACB-MISSING-ALT-TEXT",
                f"Shape '{name}' has no alternative text",
                f"Shape: {name}",
            )


def _check_tables(doc: Document, result: AuditResult) -> None:
    """Check tables for header rows, complex headers, and empty cells."""
    for ti, table in enumerate(doc.tables):
        loc = f"Table {ti + 1}"
        rows = table.rows
        if not rows:
            continue

        # Check header row designation via w:tblLook
        tbl_elem = table._tbl
        tbl_pr = tbl_elem.find(qn("w:tblPr"))
        has_header = False
        if tbl_pr is not None:
            tbl_look = tbl_pr.find(qn("w:tblLook"))
            if tbl_look is not None:
                first_row = tbl_look.get(qn("w:firstRow"), "0")
                if first_row in ("1", "true"):
                    has_header = True
            # Also check if first row is marked as header via trPr/tblHeader
            first_row_elem = rows[0]._tr
            tr_pr = first_row_elem.find(qn("w:trPr"))
            if tr_pr is not None:
                tbl_header = tr_pr.find(qn("w:tblHeader"))
                if tbl_header is not None:
                    has_header = True

        if not has_header:
            result.add(
                "ACB-TABLE-HEADER-ROW",
                "Table does not have a designated header row",
                loc,
            )

        # Check for complex table (both first row and first column have bold/header styling)
        if has_header and len(rows) > 1 and len(table.columns) > 1:
            first_col_has_headers = True
            for row in rows[1:]:
                cell = row.cells[0]
                cell_text = cell.text.strip()
                if not cell_text:
                    first_col_has_headers = False
                    break
                # Check if first cell in each row is bold
                has_bold = False
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.bold:
                            has_bold = True
                            break
                if not has_bold:
                    first_col_has_headers = False
                    break
            if first_col_has_headers:
                result.add(
                    "ACB-COMPLEX-TABLE",
                    "Table has both row and column headers -- verify screen reader navigation",
                    loc,
                )

        # Check for empty cells
        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row.cells):
                if not cell.text.strip():
                    result.add(
                        "ACB-EMPTY-TABLE-CELL",
                        f"Empty cell at row {ri + 1}, column {ci + 1}",
                        loc,
                    )


def _check_hyperlinks(doc: Document, result: AuditResult) -> None:
    """Check hyperlinks for non-descriptive text and missing underline."""
    for para in doc.paragraphs:
        for hyp in para._element.findall(qn("w:hyperlink")):
            # Build display text from runs inside the hyperlink
            runs = hyp.findall(qn("w:r"))
            display_text = "".join(
                (r.find(qn("w:t")).text or "")
                for r in runs
                if r.find(qn("w:t")) is not None
            ).strip()
            loc = (
                f"Hyperlink: '{display_text[:60]}'"
                if display_text
                else "Hyperlink (empty)"
            )

            # Non-descriptive text
            if display_text and _NON_DESCRIPTIVE_LINK_RE.match(display_text):
                result.add(
                    "ACB-LINK-TEXT",
                    f"Non-descriptive hyperlink text: '{display_text}'",
                    loc,
                )
            elif display_text and _URL_LIKE_RE.match(display_text):
                result.add(
                    "ACB-LINK-TEXT",
                    f"Raw URL used as hyperlink text: '{display_text[:60]}'",
                    loc,
                )

            # Underline check
            for r in runs:
                rpr = r.find(qn("w:rPr"))
                if rpr is not None:
                    u_elem = rpr.find(qn("w:u"))
                    if u_elem is not None:
                        u_val = u_elem.get(qn("w:val"), "single")
                        if u_val == "none":
                            result.add(
                                "ACB-LINK-UNDERLINE",
                                f"Hyperlink is not underlined: '{display_text[:60]}'",
                                loc,
                            )
                            break


def _check_form_fields(doc: Document, result: AuditResult) -> None:
    """Check legacy form fields for missing help text / labels."""
    body = doc.element.body
    for fld in body.iter(qn("w:fldChar")):
        fld_type = fld.get(qn("w:fldCharType"), "")
        if fld_type != "begin":
            continue
        ff_data = fld.find(qn("w:ffData"))
        if ff_data is None:
            continue
        # Check for status/help text
        help_text = ff_data.find(qn("w:helpText"))
        status_text = ff_data.find(qn("w:statusText"))
        name_elem = ff_data.find(qn("w:name"))
        field_name = (
            name_elem.get(qn("w:val"), "unnamed")
            if name_elem is not None
            else "unnamed"
        )

        has_label = False
        if help_text is not None and help_text.get(qn("w:val"), "").strip():
            has_label = True
        if status_text is not None and status_text.get(qn("w:val"), "").strip():
            has_label = True

        if not has_label:
            result.add(
                "ACB-FORM-FIELD-LABEL",
                f"Form field '{field_name}' has no help text",
                f"Form field: {field_name}",
            )


def _check_floating_content(doc: Document, result: AuditResult) -> None:
    """Check for floating text boxes (wp:anchor with textbox content)."""
    body = doc.element.body
    for drawing in body.iter(qn("w:drawing")):
        for anchor in drawing.iter(qn("wp:anchor")):
            # Look for text box content (wps:txbx or w:txbxContent)
            has_textbox = False
            for child in anchor.iter():
                if child.tag.endswith("}txbx") or child.tag.endswith("}txbxContent"):
                    has_textbox = True
                    break
            if has_textbox:
                doc_pr_list = list(anchor.iter(qn("wp:docPr")))
                name = (
                    doc_pr_list[0].get("name", "text box")
                    if doc_pr_list
                    else "text box"
                )
                result.add(
                    "ACB-FLOATING-CONTENT",
                    f"Floating text box '{name}' -- reading order may differ from visual order",
                    f"Floating: {name}",
                )

    # Also check VML floating text boxes (v:textbox inside w:pict)
    for pict in body.iter(qn("w:pict")):
        for textbox in pict.iter(_vml_qn("textbox")):
            result.add(
                "ACB-FLOATING-CONTENT",
                "Legacy floating text box found -- reading order may differ from visual order",
                "Floating: VML text box",
            )


def _check_author(doc: Document, result: AuditResult) -> None:
    """Check document author property."""
    author = doc.core_properties.author
    if not author or not author.strip():
        result.add(
            "ACB-DOC-AUTHOR",
            "Document has no author set in properties",
            "Document Properties",
        )


def _check_faux_headings(doc: Document, result: AuditResult) -> None:
    """Detect paragraphs formatted like headings but without heading styles."""
    from .heading_detector import detect_headings

    candidates = detect_headings(doc)
    for c in candidates:
        preview = c.text[:60]
        loc = f"Paragraph {c.paragraph_index + 1}: '{preview}'"
        result.add(
            "ACB-FAUX-HEADING",
            f"Looks like Heading {c.suggested_level} but has no heading style "
            f"(score {c.score}, {c.confidence} confidence)",
            loc,
        )


def audit_document(
    file_path: str | Path,
    *,
    list_indent_in: float | None = None,
    list_level_indents: dict[int, float] | None = None,
    para_indent_in: float | None = None,
    first_line_indent_in: float | None = None,
    style_size_overrides: dict[str, float] | None = None,
) -> AuditResult:
    """Run a full ACB Large Print compliance audit on a Word document.

    Args:
        file_path: Path to the .docx file to audit.
        list_indent_in: Expected list left-indent in inches.  Defaults to
            ``constants.LIST_INDENT_IN`` (flush left).
        para_indent_in: Expected non-list paragraph left-indent in inches.
            Defaults to ``constants.PARA_INDENT_IN`` (flush left).
        first_line_indent_in: Expected first-line indent in inches.
            Defaults to ``constants.FIRST_LINE_INDENT_IN`` (zero).
        style_size_overrides: Optional mapping of style name (``"Normal"``,
            ``"Heading 1"``, ... ``"Heading 6"``) to font size in points.
            When provided, these become the expected sizes for the
            corresponding styles instead of the ACB defaults.

    Returns:
        AuditResult with all findings.
    """
    if list_indent_in is None:
        list_indent_in = C.LIST_INDENT_IN
    if para_indent_in is None:
        para_indent_in = C.PARA_INDENT_IN
    if first_line_indent_in is None:
        first_line_indent_in = C.FIRST_LINE_INDENT_IN
    file_path = Path(file_path)
    result = AuditResult(file_path=str(file_path))

    doc = Document(str(file_path))

    _check_document_properties(doc, result)
    _check_page_setup(doc, result)
    _check_page_layout(doc, result)
    _check_styles(doc, result, style_size_overrides=style_size_overrides)
    _check_paragraph_content(
        doc,
        result,
        list_indent_in=list_indent_in,
        list_level_indents=list_level_indents,
        para_indent_in=para_indent_in,
        first_line_indent_in=first_line_indent_in,
        style_size_overrides=style_size_overrides,
    )
    _check_hyphenation(doc, result)
    _check_line_spacing(doc, result)
    _check_widow_orphan(doc, result)
    _check_run_formatting(doc, result)
    _check_page_numbers(doc, result)
    _check_alt_text(doc, result)
    _check_tables(doc, result)
    _check_hyperlinks(doc, result)
    _check_form_fields(doc, result)
    _check_floating_content(doc, result)
    _check_author(doc, result)
    _check_faux_headings(doc, result)

    # Sort findings by severity
    severity_order = {
        C.Severity.CRITICAL: 0,
        C.Severity.HIGH: 1,
        C.Severity.MEDIUM: 2,
        C.Severity.LOW: 3,
    }
    result.findings.sort(key=lambda f: severity_order.get(f.severity, 99))

    return result
