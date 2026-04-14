"""Synthetic stress corpus builders for heading detection and fixer validation."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from random import Random
from typing import Iterator

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from . import constants as C

STRESS_RANDOM_SEED = 20260414
STRESS_FIXER_DOCUMENTS = 1000
STRESS_CASES_PER_DOCUMENT = 1
STRESS_TOTAL_HEADING_CASES = STRESS_FIXER_DOCUMENTS * STRESS_CASES_PER_DOCUMENT
DEFAULT_STRESS_SAMPLE_DOCUMENTS = 12


@dataclass(frozen=True)
class StressFamily:
    slug: str
    name: str
    why: str
    detection_focus: str
    fixer_focus: str


@dataclass(frozen=True)
class StressScenario:
    document_index: int
    case_index: int
    family: StressFamily
    text: str
    body_text: str
    is_candidate: bool
    final_is_heading: bool
    expected_level: int | None
    candidate_alignment: str
    body_alignment: str
    candidate_font_name: str
    body_font_name: str
    candidate_font_size_pt: float
    body_font_size_pt: float
    candidate_bold: bool
    body_italic: bool
    candidate_left_indent_in: float
    candidate_first_line_indent_in: float
    body_left_indent_in: float
    body_first_line_indent_in: float
    uses_acb_layout: bool
    pattern: str
    language_hint: str = "en"
    candidate_apply_direct_format: bool = True
    body_apply_direct_format: bool = True
    candidate_style_name: str = "Normal"
    body_style_name: str = "Normal"


@dataclass(frozen=True)
class GeneratedStressDocument:
    path: Path
    scenarios: tuple[StressScenario, ...]


STRESS_FAMILIES: tuple[StressFamily, ...] = (
    StressFamily(
        slug="notepad-paste",
        name="Notepad Paste",
        why="Plain-text paste removes styles, so the detector must infer structure from short lines, whitespace, and casing.",
        detection_focus="unstyled heading-like lines",
        fixer_focus="restoring real heading styles without disturbing body text",
    ),
    StressFamily(
        slug="email-thread",
        name="Email Thread",
        why="Email subjects, replies, greetings, and signatures routinely look like headings without actually being structural headings.",
        detection_focus="subjects, greetings, and signature false positives",
        fixer_focus="keeping signatures flush left after repair",
    ),
    StressFamily(
        slug="web-paste",
        name="Web Paste",
        why="Browser paste often imports mixed font sizes and emphasis that create noisy faux-heading signals.",
        detection_focus="mixed visual weights and numbered callouts",
        fixer_focus="normalizing font family, size, and italics",
    ),
    StressFamily(
        slug="meeting-agenda",
        name="Meeting Agenda",
        why="Agendas combine numbered sections, sub-sections, and short procedural lines that should become real headings.",
        detection_focus="numbered heading ladders",
        fixer_focus="left-aligned heading hierarchy",
    ),
    StressFamily(
        slug="policy-manual",
        name="Policy Manual",
        why="Policies mix long body paragraphs with terse section labels and frequent indentation drift from copied templates.",
        detection_focus="section labels versus body prose",
        fixer_focus="removing hanging and first-line indents",
    ),
    StressFamily(
        slug="newsletter",
        name="Newsletter",
        why="Newsletters introduce pull quotes, bylines, and centered titles that can confuse both detection and repair.",
        detection_focus="bylines and callouts",
        fixer_focus="repairing centered or justified copy to flush left",
    ),
    StressFamily(
        slug="training-handout",
        name="Training Handout",
        why="Handouts include short prompts, exercises, and answer labels that resemble headings but are not always structural.",
        detection_focus="exercise labels and prompts",
        fixer_focus="preserving readable emphasis while removing italic and bold body text",
    ),
    StressFamily(
        slug="legal-outline",
        name="Legal Outline",
        why="Legal outlines use nested numbering and indentation patterns that must be preserved semantically, not visually.",
        detection_focus="outline numbering",
        fixer_focus="converting visual hierarchy into real heading styles",
    ),
    StressFamily(
        slug="report-appendix",
        name="Report Appendix",
        why="Appendices often mix appendix labels, figure captions, and short notes that need strong discrimination.",
        detection_focus="appendix labels versus captions",
        fixer_focus="keeping appendix content in ACB-compliant body formatting",
    ),
    StressFamily(
        slug="community-flyer",
        name="Community Flyer",
        why="Flyers rely on visual emphasis, centered slogans, and short bursts of text that must be normalized for large print.",
        detection_focus="headline slogans and event details",
        fixer_focus="repairing decorative layout into strict ACB presentation",
    ),
)

_ALIGNMENTS = {
    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_TOPIC_WORDS: dict[str, tuple[str, ...]] = {
    "notepad-paste": ("Board Update", "Volunteer Notes", "Program Review"),
    "email-thread": ("Quarterly Review", "Action Items", "Reply Summary"),
    "web-paste": ("Accessibility Tips", "Design Notes", "Release Highlights"),
    "meeting-agenda": ("Call to Order", "Old Business", "New Business"),
    "policy-manual": ("Policy Scope", "Required Steps", "Control Matrix"),
    "newsletter": ("Member Spotlight", "Upcoming Events", "Feature Story"),
    "training-handout": ("Learning Goals", "Practice Task", "Review Questions"),
    "legal-outline": ("Definitions", "Duties", "Exceptions"),
    "report-appendix": ("Appendix A", "Supporting Data", "Reference Notes"),
    "community-flyer": ("Registration Info", "Featured Speaker", "Community Day"),
}

_LANGUAGE_PREFIX: dict[str, str] = {
    "en": "Section",
    "es": "Seccion",
    "fr": "Section",
    "de": "Abschnitt",
    "it": "Sezione",
    "pt": "Secao",
    "nl": "Onderdeel",
}

_LANGUAGE_BODY_LEAD: dict[str, str] = {
    "en": "This section explains the context clearly.",
    "es": "Esta seccion explica el contexto con claridad.",
    "fr": "Cette section explique le contexte clairement.",
    "de": "Dieser Abschnitt erklaert den Kontext klar.",
    "it": "Questa sezione spiega chiaramente il contesto.",
    "pt": "Esta secao explica claramente o contexto.",
    "nl": "Deze sectie legt de context duidelijk uit.",
}


def describe_stress_corpus() -> dict[str, object]:
    """Return public metadata describing the synthetic stress corpus."""
    return {
        "seed": STRESS_RANDOM_SEED,
        "document_count": STRESS_FIXER_DOCUMENTS,
        "cases_per_document": STRESS_CASES_PER_DOCUMENT,
        "total_heading_cases": STRESS_TOTAL_HEADING_CASES,
        "families": [
            {
                "slug": family.slug,
                "name": family.name,
                "why": family.why,
                "detection_focus": family.detection_focus,
                "fixer_focus": family.fixer_focus,
            }
            for family in STRESS_FAMILIES
        ],
    }


def iter_stress_documents(
    *,
    document_count: int = STRESS_FIXER_DOCUMENTS,
    cases_per_document: int = STRESS_CASES_PER_DOCUMENT,
    seed: int = STRESS_RANDOM_SEED,
) -> Iterator[tuple[Document, tuple[StressScenario, ...]]]:
    """Yield synthetic Word documents and their expected scenarios."""
    for document_index in range(document_count):
        yield build_stress_document(
            document_index,
            cases_per_document=cases_per_document,
            seed=seed,
        )


def build_stress_document(
    document_index: int,
    *,
    cases_per_document: int = STRESS_CASES_PER_DOCUMENT,
    seed: int = STRESS_RANDOM_SEED,
) -> tuple[Document, tuple[StressScenario, ...]]:
    """Build a deterministic synthetic Word document for stress testing."""
    rng = Random(seed + document_index)
    doc = Document()
    doc.core_properties.title = f"ACB Stress Corpus Document {document_index + 1}"
    doc.core_properties.author = "ACB Stress Harness"
    doc.core_properties.subject = "Synthetic heading detection and fixer validation"

    intro = doc.add_paragraph(style="Heading 1")
    intro.add_run(f"Stress Corpus Document {document_index + 1}")
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro_body = doc.add_paragraph(
        "This synthetic document exercises heading detection and ACB large print repair paths. "
        "It mixes clean, noisy, and intentionally broken formatting patterns so the product can learn from broad real-world variance."
    )
    intro_body.alignment = WD_ALIGN_PARAGRAPH.LEFT

    scenarios: list[StressScenario] = []
    for case_index in range(cases_per_document):
        scenario = _make_scenario(document_index, case_index, rng)
        scenarios.append(scenario)
        doc.add_paragraph("")
        _add_candidate_paragraph(doc, scenario)
        _add_body_paragraph(doc, scenario)

    return doc, tuple(scenarios)


def generate_stress_documents(
    output_dir: str | Path,
    *,
    document_count: int = STRESS_FIXER_DOCUMENTS,
    cases_per_document: int = STRESS_CASES_PER_DOCUMENT,
    seed: int = STRESS_RANDOM_SEED,
) -> list[GeneratedStressDocument]:
    """Generate and save a synthetic .docx corpus to disk."""
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    generated: list[GeneratedStressDocument] = []

    for document_index, (doc, scenarios) in enumerate(
        iter_stress_documents(
            document_count=document_count,
            cases_per_document=cases_per_document,
            seed=seed,
        )
    ):
        path = output_dir / f"stress-doc-{document_index + 1:04d}.docx"
        doc.save(path)
        generated.append(GeneratedStressDocument(path=path, scenarios=scenarios))

    return generated


def _make_scenario(document_index: int, case_index: int, rng: Random) -> StressScenario:
    family = STRESS_FAMILIES[
        (document_index + case_index + rng.randrange(0, len(STRESS_FAMILIES)))
        % len(STRESS_FAMILIES)
    ]
    topic = _TOPIC_WORDS[family.slug][
        (document_index + case_index) % len(_TOPIC_WORDS[family.slug])
    ]
    token = f"D{document_index + 1:04d}-C{case_index + 1:03d}"
    pattern = rng.choice((0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11))
    uses_acb_layout = (document_index + case_index) % 3 == 0
    language_hint = rng.choice(tuple(_LANGUAGE_PREFIX.keys()))

    if pattern == 0:
        return StressScenario(
            document_index=document_index,
            case_index=case_index,
            family=family,
            text=f"{topic} {token}",
            body_text=_body_text(family, token, language_hint),
            is_candidate=True,
            final_is_heading=True,
            expected_level=1,
            candidate_alignment="LEFT",
            body_alignment="LEFT",
            candidate_font_name=C.FONT_FAMILY,
            body_font_name=C.FONT_FAMILY,
            candidate_font_size_pt=C.HEADING1_SIZE_PT,
            body_font_size_pt=C.BODY_SIZE_PT,
            candidate_bold=True,
            body_italic=False,
            candidate_left_indent_in=0.0,
            candidate_first_line_indent_in=0.0,
            body_left_indent_in=0.0,
            body_first_line_indent_in=0.0,
            uses_acb_layout=True,
            pattern="high-heading-h1",
            language_hint=language_hint,
        )

    if pattern == 1:
        return StressScenario(
            document_index=document_index,
            case_index=case_index,
            family=family,
            text=f"{_LANGUAGE_PREFIX[language_hint]} {case_index + 1}. {topic} {token}",
            body_text=_body_text(family, token, language_hint),
            is_candidate=True,
            final_is_heading=True,
            expected_level=2,
            candidate_alignment="LEFT",
            body_alignment="LEFT",
            candidate_font_name=C.FONT_FAMILY,
            body_font_name=C.FONT_FAMILY,
            candidate_font_size_pt=C.HEADING2_SIZE_PT,
            body_font_size_pt=C.BODY_SIZE_PT,
            candidate_bold=True,
            body_italic=False,
            candidate_left_indent_in=0.0,
            candidate_first_line_indent_in=0.0,
            body_left_indent_in=0.0,
            body_first_line_indent_in=0.0,
            uses_acb_layout=True,
            pattern="high-heading-numbered",
            language_hint=language_hint,
        )

    if pattern == 2:
        return StressScenario(
            document_index=document_index,
            case_index=case_index,
            family=family,
            text=f"{family.name.upper()} {token}",
            body_text=_body_text(family, token, language_hint),
            is_candidate=True,
            final_is_heading=True,
            expected_level=2,
            candidate_alignment="LEFT",
            body_alignment="LEFT",
            candidate_font_name=C.FONT_FAMILY,
            body_font_name=C.FONT_FAMILY,
            candidate_font_size_pt=C.BODY_SIZE_PT,
            body_font_size_pt=C.BODY_SIZE_PT,
            candidate_bold=True,
            body_italic=False,
            candidate_left_indent_in=0.0,
            candidate_first_line_indent_in=0.0,
            body_left_indent_in=0.0,
            body_first_line_indent_in=0.0,
            uses_acb_layout=True,
            pattern="medium-heading-ai-confirmed",
            language_hint=language_hint,
        )

    if pattern == 3:
        return StressScenario(
            document_index=document_index,
            case_index=case_index,
            family=family,
            text=f"Jane Smith, {token}.",
            body_text=_body_text(family, token, language_hint),
            is_candidate=False,
            final_is_heading=False,
            expected_level=None,
            candidate_alignment="LEFT",
            body_alignment="LEFT",
            candidate_font_name=C.FONT_FAMILY,
            body_font_name=C.FONT_FAMILY,
            candidate_font_size_pt=C.BODY_SIZE_PT,
            body_font_size_pt=C.BODY_SIZE_PT,
            candidate_bold=False,
            body_italic=False,
            candidate_left_indent_in=0.0,
            candidate_first_line_indent_in=0.0,
            body_left_indent_in=0.0,
            body_first_line_indent_in=0.0,
            uses_acb_layout=True,
            pattern="false-positive-signature",
            language_hint=language_hint,
        )

    if pattern == 4:
        return StressScenario(
            document_index=document_index,
            case_index=case_index,
            family=family,
            text=f"Please Note: {token}.",
            body_text=_body_text(family, token, language_hint),
            is_candidate=False,
            final_is_heading=False,
            expected_level=None,
            candidate_alignment="CENTER" if not uses_acb_layout else "LEFT",
            body_alignment="JUSTIFY" if not uses_acb_layout else "LEFT",
            candidate_font_name="Calibri" if not uses_acb_layout else C.FONT_FAMILY,
            body_font_name="Times New Roman" if not uses_acb_layout else C.FONT_FAMILY,
            candidate_font_size_pt=C.BODY_SIZE_PT,
            body_font_size_pt=14.0 if not uses_acb_layout else C.BODY_SIZE_PT,
            candidate_bold=False,
            body_italic=not uses_acb_layout,
            candidate_left_indent_in=0.5 if not uses_acb_layout else 0.0,
            candidate_first_line_indent_in=-0.25 if not uses_acb_layout else 0.0,
            body_left_indent_in=0.5 if not uses_acb_layout else 0.0,
            body_first_line_indent_in=-0.25 if not uses_acb_layout else 0.0,
            uses_acb_layout=uses_acb_layout,
            pattern="false-positive-callout",
            language_hint=language_hint,
        )

    if pattern == 5:
        return StressScenario(
            document_index=document_index,
            case_index=case_index,
            family=family,
            text=(
                f"This ordinary paragraph {token} looks descriptive, ends with punctuation, "
                f"and continues long enough to stay below the heading threshold for {family.name.lower()}."
            ),
            body_text=_body_text(family, token, language_hint),
            is_candidate=False,
            final_is_heading=False,
            expected_level=None,
            candidate_alignment="LEFT",
            body_alignment="LEFT",
            candidate_font_name=C.FONT_FAMILY,
            body_font_name=C.FONT_FAMILY,
            candidate_font_size_pt=C.BODY_SIZE_PT,
            body_font_size_pt=C.BODY_SIZE_PT,
            candidate_bold=False,
            body_italic=False,
            candidate_left_indent_in=0.0,
            candidate_first_line_indent_in=0.0,
            body_left_indent_in=0.0,
            body_first_line_indent_in=0.0,
            uses_acb_layout=True,
            pattern="negative-body-line",
            language_hint=language_hint,
        )

    if pattern == 6:
        return StressScenario(
            document_index=document_index,
            case_index=case_index,
            family=family,
            text=(
                f"Reminder: {topic} {token}, please review the attached details before the next session "
                "because this line is intentionally too long and punctuated to count as a heading."
            ),
            body_text=_body_text(family, token, language_hint),
            is_candidate=False,
            final_is_heading=False,
            expected_level=None,
            candidate_alignment="LEFT",
            body_alignment="LEFT",
            candidate_font_name=C.FONT_FAMILY,
            body_font_name=C.FONT_FAMILY,
            candidate_font_size_pt=C.BODY_SIZE_PT,
            body_font_size_pt=C.BODY_SIZE_PT,
            candidate_bold=False,
            body_italic=False,
            candidate_left_indent_in=0.0,
            candidate_first_line_indent_in=0.0,
            body_left_indent_in=0.0,
            body_first_line_indent_in=0.0,
            uses_acb_layout=True,
            pattern="negative-long-body",
            language_hint=language_hint,
        )

    if pattern == 8:
        return StressScenario(
            document_index=document_index,
            case_index=case_index,
            family=family,
            text=(
                f"This plain text line about {topic} {token} intentionally reads like ordinary narrative content, "
                "includes sentence punctuation, and should remain below heading confidence."
            ),
            body_text=_body_text(family, token, language_hint),
            is_candidate=False,
            final_is_heading=False,
            expected_level=None,
            candidate_alignment="LEFT",
            body_alignment="LEFT",
            candidate_font_name=C.FONT_FAMILY,
            body_font_name=C.FONT_FAMILY,
            candidate_font_size_pt=C.BODY_SIZE_PT,
            body_font_size_pt=C.BODY_SIZE_PT,
            candidate_bold=False,
            body_italic=False,
            candidate_left_indent_in=0.0,
            candidate_first_line_indent_in=0.0,
            body_left_indent_in=0.0,
            body_first_line_indent_in=0.0,
            uses_acb_layout=True,
            pattern="plain-text-no-style-negative",
            language_hint=language_hint,
            candidate_apply_direct_format=False,
            body_apply_direct_format=False,
        )

    if pattern == 9:
        return StressScenario(
            document_index=document_index,
            case_index=case_index,
            family=family,
            text=f"{topic} {token}",
            body_text=_body_text(family, token, language_hint),
            is_candidate=True,
            final_is_heading=True,
            expected_level=2,
            candidate_alignment="LEFT",
            body_alignment="LEFT",
            candidate_font_name="Calibri" if not uses_acb_layout else C.FONT_FAMILY,
            body_font_name=C.FONT_FAMILY,
            candidate_font_size_pt=C.HEADING2_SIZE_PT,
            body_font_size_pt=C.BODY_SIZE_PT,
            candidate_bold=False,
            body_italic=False,
            candidate_left_indent_in=0.0,
            candidate_first_line_indent_in=0.0,
            body_left_indent_in=0.0,
            body_first_line_indent_in=0.0,
            uses_acb_layout=uses_acb_layout,
            pattern="font-size-only-heading",
            language_hint=language_hint,
        )

    if pattern == 10:
        return StressScenario(
            document_index=document_index,
            case_index=case_index,
            family=family,
            text=f"Thanks for your help {token}.",
            body_text=_body_text(family, token, language_hint),
            is_candidate=False,
            final_is_heading=False,
            expected_level=None,
            candidate_alignment="LEFT",
            body_alignment="LEFT",
            candidate_font_name=C.FONT_FAMILY,
            body_font_name=C.FONT_FAMILY,
            candidate_font_size_pt=C.BODY_SIZE_PT,
            body_font_size_pt=C.BODY_SIZE_PT,
            candidate_bold=False,
            body_italic=False,
            candidate_left_indent_in=0.0,
            candidate_first_line_indent_in=0.0,
            body_left_indent_in=0.0,
            body_first_line_indent_in=0.0,
            uses_acb_layout=True,
            pattern="multilingual-salutation-false-positive",
            language_hint=language_hint,
        )

    if pattern == 11:
        return StressScenario(
            document_index=document_index,
            case_index=case_index,
            family=family,
            text=f"{_LANGUAGE_PREFIX[language_hint]} {document_index + 1} {topic} {token}",
            body_text=_body_text(family, token, language_hint),
            is_candidate=True,
            final_is_heading=True,
            expected_level=2,
            candidate_alignment="LEFT" if uses_acb_layout else "CENTER",
            body_alignment="LEFT" if uses_acb_layout else "JUSTIFY",
            candidate_font_name=C.FONT_FAMILY if uses_acb_layout else "Georgia",
            body_font_name=C.FONT_FAMILY if uses_acb_layout else "Times New Roman",
            candidate_font_size_pt=C.HEADING2_SIZE_PT,
            body_font_size_pt=C.BODY_SIZE_PT if uses_acb_layout else 16.0,
            candidate_bold=True,
            body_italic=not uses_acb_layout,
            candidate_left_indent_in=0.0 if uses_acb_layout else 0.5,
            candidate_first_line_indent_in=0.0 if uses_acb_layout else -0.25,
            body_left_indent_in=0.0 if uses_acb_layout else 0.5,
            body_first_line_indent_in=0.0 if uses_acb_layout else -0.25,
            uses_acb_layout=uses_acb_layout,
            pattern="multilingual-numbered-heading",
            language_hint=language_hint,
        )

    return StressScenario(
        document_index=document_index,
        case_index=case_index,
        family=family,
        text=f"Part {document_index + 1}.{case_index + 1} {topic} {token}",
        body_text=_body_text(family, token, language_hint),
        is_candidate=True,
        final_is_heading=True,
        expected_level=2,
        candidate_alignment="LEFT" if uses_acb_layout else "RIGHT",
        body_alignment="LEFT" if uses_acb_layout else "CENTER",
        candidate_font_name=C.FONT_FAMILY if uses_acb_layout else "Georgia",
        body_font_name=C.FONT_FAMILY if uses_acb_layout else "Cambria",
        candidate_font_size_pt=C.HEADING2_SIZE_PT,
        body_font_size_pt=C.BODY_SIZE_PT if uses_acb_layout else 16.0,
        candidate_bold=True,
        body_italic=not uses_acb_layout,
        candidate_left_indent_in=0.0 if uses_acb_layout else -0.25,
        candidate_first_line_indent_in=0.0 if uses_acb_layout else -0.25,
        body_left_indent_in=0.0 if uses_acb_layout else -0.25,
        body_first_line_indent_in=0.0 if uses_acb_layout else 0.5,
        uses_acb_layout=uses_acb_layout,
        pattern="numbered-heading-with-layout-drift",
        language_hint=language_hint,
    )


def _body_text(family: StressFamily, token: str, language_hint: str = "en") -> str:
    lead = _LANGUAGE_BODY_LEAD.get(language_hint, _LANGUAGE_BODY_LEAD["en"])
    return (
        f"{family.name} body copy {token}. {lead} "
        "It is intentionally long enough to look like body text, provide a strong next-paragraph signal, "
        "and expose alignment, font, and indentation drift during repair."
    )


def _add_candidate_paragraph(doc: Document, scenario: StressScenario) -> None:
    para = doc.add_paragraph(style=scenario.candidate_style_name)
    run = para.add_run(scenario.text)
    if scenario.candidate_apply_direct_format:
        run.font.name = scenario.candidate_font_name
        run.font.size = Pt(scenario.candidate_font_size_pt)
        run.font.bold = scenario.candidate_bold
        para.alignment = _ALIGNMENTS[scenario.candidate_alignment]
        para.paragraph_format.left_indent = Inches(scenario.candidate_left_indent_in)
        para.paragraph_format.first_line_indent = Inches(
            scenario.candidate_first_line_indent_in
        )


def _add_body_paragraph(doc: Document, scenario: StressScenario) -> None:
    para = doc.add_paragraph(style=scenario.body_style_name)
    run = para.add_run(scenario.body_text)
    if scenario.body_apply_direct_format:
        run.font.name = scenario.body_font_name
        run.font.size = Pt(scenario.body_font_size_pt)
        run.font.italic = scenario.body_italic
        para.alignment = _ALIGNMENTS[scenario.body_alignment]
        para.paragraph_format.left_indent = Inches(scenario.body_left_indent_in)
        para.paragraph_format.first_line_indent = Inches(
            scenario.body_first_line_indent_in
        )
