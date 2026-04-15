"""Fix Word documents to comply with ACB Large Print Guidelines."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from . import constants as C
from .auditor import AuditResult, audit_document


def _apply_acb_font(font_obj, spec: C.FontDef) -> None:
    """Apply ACB font settings to a python-docx font object."""
    font_obj.name = spec.name
    font_obj.size = Pt(spec.size_pt)
    font_obj.bold = spec.bold
    font_obj.italic = False
    font_obj.all_caps = False
    font_obj.color.rgb = RGBColor(0, 0, 0)


def _list_level_from_style(style_name: str) -> int | None:
    """Infer list nesting level from built-in Word list style names."""
    if style_name in ("List Bullet", "List Number"):
        return 0
    if style_name in ("List Bullet 2", "List Number 2"):
        return 1
    if style_name in ("List Bullet 3", "List Number 3"):
        return 2
    return None


def _fix_styles(
    doc: Document,
    records: list[C.FixRecord],
    *,
    list_indent_in: float = C.LIST_INDENT_IN,
    list_hanging_in: float = C.LIST_HANGING_IN,
    list_level_indents: dict[int, float] | None = None,
    preserve_heading_alignment: bool = False,
) -> int:
    """Fix all named styles to match ACB specs. Returns count of fixes."""
    fixes = 0
    for style_name, style_def in C.ACB_STYLES.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue

        font = style.font
        pf = style.paragraph_format
        changed = False
        details: list[str] = []

        # Font
        if font.name != style_def.font.name:
            font.name = style_def.font.name
            details.append(f"font family -> {style_def.font.name}")
            changed = True
        if font.size is None or abs(font.size.pt - style_def.font.size_pt) > 0.5:
            font.size = Pt(style_def.font.size_pt)
            details.append(f"font size -> {style_def.font.size_pt}pt")
            changed = True
        if font.bold != style_def.font.bold:
            font.bold = style_def.font.bold
            changed = True
        if font.italic:
            font.italic = False
            details.append("removed italic")
            changed = True
        if font.all_caps:
            font.all_caps = False
            details.append("removed all caps")
            changed = True

        # Paragraph format
        if pf.alignment != WD_ALIGN_PARAGRAPH.LEFT and not (
            preserve_heading_alignment and style_name.startswith("Heading")
        ):
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            details.append("alignment -> flush left")
            changed = True
        if pf.line_spacing != style_def.para.line_spacing:
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = style_def.para.line_spacing
            changed = True
        if pf.widow_control is not True:
            pf.widow_control = True
            changed = True

        # Apply user-configurable list indentation to list styles
        list_level = _list_level_from_style(style_name)
        if list_level is not None:
            target_indent = list_indent_in
            if list_level_indents is not None:
                target_indent = list_level_indents.get(list_level, list_indent_in)
            target_left = Inches(target_indent)
            if pf.left_indent is None or abs(pf.left_indent - target_left) > Inches(
                0.05
            ):
                pf.left_indent = target_left
                changed = True
            if target_indent > 0 and list_hanging_in > 0:
                target_hang = Inches(-list_hanging_in)
                if pf.first_line_indent is None or abs(
                    pf.first_line_indent - target_hang
                ) > Inches(0.05):
                    pf.first_line_indent = target_hang
                    changed = True
            elif target_indent == 0:
                if pf.first_line_indent is not None and pf.first_line_indent != Inches(
                    0
                ):
                    pf.first_line_indent = None
                    changed = True

        # Remove theme colors from headings
        if style_name.startswith("Heading"):
            rPr = style.element.find(qn("w:rPr"))
            if rPr is not None:
                color_elem = rPr.find(qn("w:color"))
                if color_elem is not None:
                    theme = color_elem.get(qn("w:themeColor"))
                    if theme:
                        color_elem.attrib.pop(qn("w:themeColor"), None)
                        color_elem.attrib.pop(qn("w:themeShade"), None)
                        color_elem.set(qn("w:val"), "000000")
                        details.append("removed theme color")
                        changed = True

        if changed:
            fixes += 1
            desc = f"Fixed style '{style_name}'"
            if details:
                desc += ": " + ", ".join(details)
            records.append(
                C.FixRecord(
                    rule_id=(
                        "ACB-FONT-FAMILY"
                        if "font family" in desc
                        else "ACB-FONT-SIZE-BODY"
                    ),
                    category=C.FIX_CATEGORY_FONT,
                    description=desc,
                    location=f"Style: {style_name}",
                )
            )

    return fixes


def _fix_page_setup(
    doc: Document, records: list[C.FixRecord], bound: bool = False
) -> int:
    """Fix page margins and orientation. Returns count of fixes."""
    fixes = 0
    tolerance = Inches(0.05)
    for idx, section in enumerate(doc.sections):
        sec_loc = f"Section {idx + 1}"
        left_target = C.MARGIN_LEFT_IN + (C.BINDING_EXTRA_IN if bound else 0)

        if (
            section.top_margin is None
            or abs(section.top_margin - Inches(C.MARGIN_TOP_IN)) > tolerance
        ):
            section.top_margin = Inches(C.MARGIN_TOP_IN)
            fixes += 1
            records.append(
                C.FixRecord(
                    "ACB-MARGINS",
                    C.FIX_CATEGORY_PAGE,
                    f"Top margin -> {C.MARGIN_TOP_IN}in",
                    sec_loc,
                )
            )
        if (
            section.bottom_margin is None
            or abs(section.bottom_margin - Inches(C.MARGIN_BOTTOM_IN)) > tolerance
        ):
            section.bottom_margin = Inches(C.MARGIN_BOTTOM_IN)
            fixes += 1
            records.append(
                C.FixRecord(
                    "ACB-MARGINS",
                    C.FIX_CATEGORY_PAGE,
                    f"Bottom margin -> {C.MARGIN_BOTTOM_IN}in",
                    sec_loc,
                )
            )
        if (
            section.left_margin is None
            or abs(section.left_margin - Inches(left_target)) > tolerance
        ):
            section.left_margin = Inches(left_target)
            fixes += 1
            records.append(
                C.FixRecord(
                    "ACB-MARGINS",
                    C.FIX_CATEGORY_PAGE,
                    f"Left margin -> {left_target}in",
                    sec_loc,
                )
            )
        if (
            section.right_margin is None
            or abs(section.right_margin - Inches(C.MARGIN_RIGHT_IN)) > tolerance
        ):
            section.right_margin = Inches(C.MARGIN_RIGHT_IN)
            fixes += 1
            records.append(
                C.FixRecord(
                    "ACB-MARGINS",
                    C.FIX_CATEGORY_PAGE,
                    f"Right margin -> {C.MARGIN_RIGHT_IN}in",
                    sec_loc,
                )
            )

    return fixes


def _fix_paragraph_formatting(
    doc: Document,
    records: list[C.FixRecord],
    *,
    list_indent_in: float = C.LIST_INDENT_IN,
    list_hanging_in: float = C.LIST_HANGING_IN,
    list_level_indents: dict[int, float] | None = None,
    para_indent_in: float = C.PARA_INDENT_IN,
    first_line_indent_in: float = C.FIRST_LINE_INDENT_IN,
    preserve_heading_alignment: bool = False,
) -> int:
    """Fix direct formatting overrides in paragraph content. Returns fix count."""
    fixes = 0

    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else "Normal"
        is_heading = bool(re.match(r"^Heading \d+$", style_name))
        is_list = style_name in (
            "List Bullet",
            "List Number",
            "List Bullet 2",
            "List Number 2",
            "List Bullet 3",
            "List Number 3",
        )

        loc = f"Paragraph {i + 1}"

        # Fix alignment
        if para.paragraph_format.alignment is not None:
            if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                if not (preserve_heading_alignment and is_heading):
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    fixes += 1
                    records.append(
                        C.FixRecord(
                            "ACB-ALIGNMENT",
                            C.FIX_CATEGORY_ALIGNMENT,
                            "Alignment -> flush left",
                            loc,
                        )
                    )

        # Fix list indentation
        if is_list:
            pf = para.paragraph_format
            level = _list_level_from_style(style_name)
            expected_indent = list_indent_in
            if list_level_indents is not None and level is not None:
                expected_indent = list_level_indents.get(level, list_indent_in)
            actual_indent = pf.left_indent.inches if pf.left_indent else 0.0
            if abs(actual_indent - expected_indent) > 0.05:
                pf.left_indent = (
                    Inches(expected_indent) if expected_indent > 0 else Inches(0)
                )
                if list_hanging_in and expected_indent > 0:
                    pf.first_line_indent = Inches(-list_hanging_in)
                elif expected_indent == 0:
                    pf.first_line_indent = None
                fixes += 1
                records.append(
                    C.FixRecord(
                        "ACB-LIST-INDENT",
                        C.FIX_CATEGORY_ALIGNMENT,
                        f"List indent -> {expected_indent}in",
                        loc,
                    )
                )

        # Fix non-list, non-heading paragraph indentation
        if not is_list and not is_heading and para.text.strip():
            pf = para.paragraph_format
            # Left indent
            actual_left = pf.left_indent.inches if pf.left_indent else 0.0
            if abs(actual_left - para_indent_in) > 0.05:
                pf.left_indent = (
                    Inches(para_indent_in) if para_indent_in > 0 else Inches(0)
                )
                fixes += 1
                records.append(
                    C.FixRecord(
                        "ACB-PARA-INDENT",
                        C.FIX_CATEGORY_ALIGNMENT,
                        f"Left indent -> {para_indent_in}in",
                        loc,
                    )
                )
            # First-line indent
            actual_fli = pf.first_line_indent.inches if pf.first_line_indent else 0.0
            if abs(actual_fli - first_line_indent_in) > 0.05:
                pf.first_line_indent = (
                    Inches(first_line_indent_in)
                    if abs(first_line_indent_in) > 0.001
                    else None
                )
                fixes += 1
                records.append(
                    C.FixRecord(
                        "ACB-FIRST-LINE-INDENT",
                        C.FIX_CATEGORY_ALIGNMENT,
                        f"First-line indent -> {first_line_indent_in}in",
                        loc,
                    )
                )

        for run in para.runs:
            # Remove italic everywhere
            if run.font.italic:
                run.font.italic = False
                # Convert italic text to underline (ACB emphasis)
                run.font.underline = True
                fixes += 1
                records.append(
                    C.FixRecord(
                        "ACB-ITALIC",
                        C.FIX_CATEGORY_EMPHASIS,
                        "Italic -> underline",
                        loc,
                    )
                )

            # Fix bold in body text: convert to underline
            if run.font.bold and not is_heading and run.text.strip():
                if len(run.text.strip()) > 5:
                    run.font.bold = False
                    run.font.underline = True
                    fixes += 1
                    records.append(
                        C.FixRecord(
                            "ACB-BOLD-BODY",
                            C.FIX_CATEGORY_EMPHASIS,
                            "Bold body text -> underline",
                            loc,
                        )
                    )

            # Fix font family
            if run.font.name and run.font.name != C.FONT_FAMILY:
                old_name = run.font.name
                run.font.name = C.FONT_FAMILY
                fixes += 1
                records.append(
                    C.FixRecord(
                        "ACB-FONT-FAMILY",
                        C.FIX_CATEGORY_FONT,
                        f"Font '{old_name}' -> {C.FONT_FAMILY}",
                        loc,
                    )
                )

            # Fix font size below minimum
            if run.font.size and run.font.size.pt < C.MIN_SIZE_PT - 0.5:
                old_size = run.font.size.pt
                run.font.size = Pt(C.MIN_SIZE_PT)
                fixes += 1
                records.append(
                    C.FixRecord(
                        "ACB-FONT-SIZE-BODY",
                        C.FIX_CATEGORY_FONT,
                        f"Font size {old_size}pt -> {C.MIN_SIZE_PT}pt",
                        loc,
                    )
                )

            # Remove all caps
            if run.font.all_caps:
                run.font.all_caps = False
                fixes += 1
                records.append(
                    C.FixRecord(
                        "ACB-ALL-CAPS", C.FIX_CATEGORY_EMPHASIS, "Removed all caps", loc
                    )
                )

    return fixes


def _fix_hyphenation(doc: Document, records: list[C.FixRecord]) -> int:
    """Disable automatic hyphenation. Returns 1 if fixed, 0 if already correct."""
    settings = doc.settings.element
    auto_hyph = settings.find(qn("w:autoHyphenation"))
    if auto_hyph is not None:
        val = auto_hyph.get(qn("w:val"), "true")
        if val.lower() not in ("false", "0", "off"):
            auto_hyph.set(qn("w:val"), "false")
            records.append(
                C.FixRecord(
                    "ACB-HYPHENATION",
                    C.FIX_CATEGORY_PAGE,
                    "Disabled automatic hyphenation",
                    "Document settings",
                )
            )
            return 1
        return 0
    # Add element explicitly set to false
    elem = OxmlElement("w:autoHyphenation")
    elem.set(qn("w:val"), "false")
    settings.append(elem)
    records.append(
        C.FixRecord(
            "ACB-HYPHENATION",
            C.FIX_CATEGORY_PAGE,
            "Disabled automatic hyphenation",
            "Document settings",
        )
    )
    return 1


def _fix_page_numbers(doc: Document, records: list[C.FixRecord]) -> int:
    """Add page numbers to footer if missing. Returns 1 if added."""
    has_page_numbers = False
    for section in doc.sections:
        footer = section.footer
        if footer is None:
            continue
        footer_xml = footer._element.xml
        if "PAGE" in footer_xml or "w:fldChar" in footer_xml:
            has_page_numbers = True
            break

    if has_page_numbers:
        return 0

    # Add page numbers to first section footer (others will link)
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = paragraph.add_run()
    run.font.name = C.FONT_FAMILY
    run.font.size = Pt(C.FOOTER_SIZE_PT)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    records.append(
        C.FixRecord(
            "ACB-PAGE-NUMBERS",
            C.FIX_CATEGORY_PAGE,
            "Added page numbers to footer",
            "Section 1 footer",
        )
    )
    return 1


def _fix_document_title(
    doc: Document, file_path: Path, records: list[C.FixRecord]
) -> tuple[int, str | None]:
    """Set document title from first Heading 1 or filename. Returns (fix_count, warning).

    WARNING: The inferred title should be manually reviewed.
    """
    title = doc.core_properties.title
    if title and title.strip():
        return 0, None

    # Try first Heading 1
    for para in doc.paragraphs:
        if para.style and para.style.name == "Heading 1" and para.text.strip():
            doc.core_properties.title = para.text.strip()
            warn = (
                f"Document title was set from the first heading: "
                f'"{para.text.strip()}". Please review and correct if needed.'
            )
            records.append(
                C.FixRecord(
                    "ACB-DOC-TITLE",
                    C.FIX_CATEGORY_PROPS,
                    f"Title set from heading: {para.text.strip()}",
                    "Core properties",
                )
            )
            return 1, warn

    # Fallback to filename (without extension)
    doc.core_properties.title = file_path.stem
    warn = (
        f"Document title was set from the filename: "
        f'"{file_path.stem}". Please review and correct if needed.'
    )
    records.append(
        C.FixRecord(
            "ACB-DOC-TITLE",
            C.FIX_CATEGORY_PROPS,
            f"Title set from filename: {file_path.stem}",
            "Core properties",
        )
    )
    return 1, warn


def _fix_document_language(
    doc: Document, records: list[C.FixRecord], lang: str = "en-US"
) -> int:
    """Set document language if missing. Returns 1 if fixed."""
    styles_elem = doc.styles.element
    doc_defaults = styles_elem.find(qn("w:docDefaults"))
    rpr = None
    if doc_defaults is not None:
        rpr_default = doc_defaults.find(qn("w:rPrDefault"))
        if rpr_default is not None:
            rpr = rpr_default.find(qn("w:rPr"))
    if rpr is not None:
        lang_elem = rpr.find(qn("w:lang"))
        if lang_elem is None:
            lang_elem = OxmlElement("w:lang")
            rpr.append(lang_elem)
        existing = lang_elem.get(qn("w:val"))
        if existing and existing.strip():
            return 0
        lang_elem.set(qn("w:val"), lang)
        lang_elem.set(qn("w:eastAsia"), lang)
        lang_elem.set(qn("w:bidi"), lang)
        records.append(
            C.FixRecord(
                "ACB-DOC-LANG",
                C.FIX_CATEGORY_PROPS,
                f"Document language set to {lang}",
                "Document defaults",
            )
        )
        return 1
    return 0


def _convert_faux_headings(
    doc: Document,
    records: list[C.FixRecord],
    *,
    ai_provider: object | None = None,
    threshold: int | None = None,
    system_prompt: str | None = None,
    confirmed_headings: list | None = None,
    accuracy_level: str = "balanced",
) -> int:
    """Detect faux headings and convert them to real heading styles.

    This is a pre-pass that runs before other fixes so that subsequent
    style fixes can apply correct heading formatting.

    Args:
        confirmed_headings: Pre-confirmed list of
            ``(paragraph_index, suggested_level, text)`` tuples from an
            interactive review step.  When provided, detection is skipped
            and only the confirmed candidates are applied.
        accuracy_level: "conservative", "balanced", or "thorough" for detection strategy.

    Returns:
        Number of paragraphs converted to headings.
    """
    if confirmed_headings is not None:
        # Build lightweight candidate objects from confirmed data
        from .heading_detector import HeadingCandidate

        candidates = [
            HeadingCandidate(
                paragraph_index=idx,
                text=txt,
                font_size_pt=None,
                is_bold=False,
                is_all_caps=False,
                is_title_case=False,
                char_count=len(txt),
                score=100,
                suggested_level=level,
                confidence="confirmed",
            )
            for idx, level, txt in confirmed_headings
        ]
    else:
        from .heading_detector import detect_headings

        candidates = detect_headings(
            doc,
            ai_provider=ai_provider,
            threshold=threshold,
            system_prompt=system_prompt,
            accuracy_level=accuracy_level,
        )
    if not candidates:
        return 0

    fixes = 0
    para_list = list(doc.paragraphs)
    for candidate in candidates:
        if candidate.paragraph_index >= len(para_list):
            continue
        para = para_list[candidate.paragraph_index]
        # Guard against stale indices for auto-detected candidates.
        # For user-confirmed headings from web review, trust the index selection
        # because form round-trips can normalize whitespace and punctuation.
        para_text = para.text.strip()
        candidate_text = candidate.text.strip()
        if para_text != candidate_text:
            para_norm = " ".join(para_text.split())
            candidate_norm = " ".join(candidate_text.split())
            is_confirmed = candidate.confidence == "confirmed"
            if not is_confirmed and para_norm != candidate_norm:
                continue

        target_style = f"Heading {candidate.suggested_level}"
        try:
            _ = doc.styles[target_style]
        except KeyError:
            continue

        old_style = para.style.name if para.style else "Normal"
        para.style = doc.styles[target_style]
        # Clear direct formatting that conflicts with heading styles
        for run in para.runs:
            run.font.size = None
            run.font.bold = None
            run.font.name = None
        fixes += 1
        records.append(
            C.FixRecord(
                rule_id="ACB-FAUX-HEADING",
                category=C.FIX_CATEGORY_HEADINGS,
                description=f"Converted '{candidate.text[:60]}' from {old_style} to {target_style} (confidence {candidate.score}%)",
                location=f"Paragraph {candidate.paragraph_index + 1}",
            )
        )

    return fixes


def _normalize_heading_structure(doc: Document, records: list[C.FixRecord]) -> int:
    """Normalize heading text and hierarchy to satisfy ACB heading rules.

    - Converts ALL CAPS heading text to title case.
    - Prevents skipped heading levels (e.g., H1 -> H3).
    """
    fixes = 0
    heading_pattern = re.compile(r"^Heading (\d+)$")
    prev_level = 0

    for idx, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        m = heading_pattern.match(style_name)
        if not m or not para.text.strip():
            continue

        level = int(m.group(1))
        loc = f"Paragraph {idx + 1}"

        # Fix ALL CAPS heading text.
        text = para.text.strip()
        if text.isupper() and len(text) > 1:
            para.text = text.title()
            fixes += 1
            records.append(
                C.FixRecord(
                    "ACB-NO-ALLCAPS",
                    C.FIX_CATEGORY_HEADINGS,
                    "Converted ALL CAPS heading text to title case",
                    loc,
                )
            )

        # Fix skipped heading levels.
        if prev_level and level > prev_level + 1:
            new_level = prev_level + 1
            target_style = f"Heading {new_level}"
            try:
                para.style = doc.styles[target_style]
                level = new_level
                fixes += 1
                records.append(
                    C.FixRecord(
                        "ACB-HEADING-HIERARCHY",
                        C.FIX_CATEGORY_HEADINGS,
                        f"Adjusted heading level to avoid skip: Heading {new_level}",
                        loc,
                    )
                )
            except KeyError:
                # If style is missing, skip gracefully.
                pass

        prev_level = level

    return fixes


def fix_document(
    file_path: str | Path,
    output_path: str | Path | None = None,
    *,
    bound: bool = False,
    list_indent_in: float | None = None,
    list_hanging_in: float | None = None,
    list_level_indents: dict[int, float] | None = None,
    para_indent_in: float | None = None,
    first_line_indent_in: float | None = None,
    preserve_heading_alignment: bool = False,
    detect_headings: bool = False,
    ai_provider: object | None = None,
    heading_threshold: int | None = None,
    system_prompt: str | None = None,
    confirmed_headings: list | None = None,
    heading_accuracy_level: str = "balanced",
) -> tuple[Path, int, list[C.FixRecord], AuditResult, list[str]]:
    """Fix a Word document for ACB compliance.

    Args:
        file_path: Input .docx file.
        output_path: Where to save fixed document. Defaults to overwriting.
        bound: Add binding margin if True.
        list_indent_in: Left indent for list paragraphs (inches).
        list_hanging_in: Hanging indent for list paragraphs (inches).
        para_indent_in: Left indent for non-list paragraphs (inches).
        first_line_indent_in: First-line indent for paragraphs (inches).
        detect_headings: Run faux-heading detection and convert to real headings.
        ai_provider: Optional AI provider for heading detection refinement.
        heading_threshold: Confidence threshold for heading detection.
        system_prompt: Custom system prompt for AI heading detection.
        heading_accuracy_level: "conservative" (heuristics only), "balanced" (default),
                                or "thorough" (extra AI refinement).

    Returns:
        Tuple of (output_path, total_fixes, fix_records, post_fix_audit_result, warnings).
    """
    if list_indent_in is None:
        list_indent_in = C.LIST_INDENT_IN
    if list_hanging_in is None:
        list_hanging_in = C.LIST_HANGING_IN
    if para_indent_in is None:
        para_indent_in = C.PARA_INDENT_IN
    if first_line_indent_in is None:
        first_line_indent_in = C.FIRST_LINE_INDENT_IN

    file_path = Path(file_path)
    if output_path is None:
        output_path = file_path
    output_path = Path(output_path)

    doc = Document(str(file_path))

    total_fixes = 0
    warnings: list[str] = []
    records: list[C.FixRecord] = []

    # Phase 4: Heading conversion pre-pass (before other fixes)
    if detect_headings or confirmed_headings:
        total_fixes += _convert_faux_headings(
            doc,
            records,
            ai_provider=ai_provider,
            threshold=heading_threshold,
            system_prompt=system_prompt,
            confirmed_headings=confirmed_headings,
            accuracy_level=heading_accuracy_level,
        )

    total_fixes += _normalize_heading_structure(doc, records)

    total_fixes += _fix_styles(
        doc,
        records,
        list_indent_in=list_indent_in,
        list_hanging_in=list_hanging_in,
        list_level_indents=list_level_indents,
        preserve_heading_alignment=preserve_heading_alignment,
    )
    total_fixes += _fix_page_setup(doc, records, bound=bound)
    total_fixes += _fix_paragraph_formatting(
        doc,
        records,
        list_indent_in=list_indent_in,
        list_hanging_in=list_hanging_in,
        list_level_indents=list_level_indents,
        para_indent_in=para_indent_in,
        first_line_indent_in=first_line_indent_in,
        preserve_heading_alignment=preserve_heading_alignment,
    )
    total_fixes += _fix_hyphenation(doc, records)
    total_fixes += _fix_page_numbers(doc, records)
    total_fixes += _fix_document_language(doc, records)

    title_fixes, title_warning = _fix_document_title(doc, file_path, records)
    total_fixes += title_fixes
    if title_warning:
        warnings.append(title_warning)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    # Run post-fix audit to show remaining issues
    post_audit = audit_document(
        output_path,
        list_indent_in=list_indent_in,
        list_level_indents=list_level_indents,
        para_indent_in=para_indent_in,
        first_line_indent_in=first_line_indent_in,
    )

    return output_path, total_fixes, records, post_audit, warnings
