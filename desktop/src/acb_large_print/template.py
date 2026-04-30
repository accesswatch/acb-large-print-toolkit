"""Create ACB Large Print Word templates (.dotx) with all styles pre-configured."""

from __future__ import annotations

from dataclasses import replace
import os
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from . import constants as C


def _apply_font(font_obj, spec: C.FontDef) -> None:
    """Apply a FontDef to a python-docx font object."""
    font_obj.name = spec.name
    font_obj.size = Pt(spec.size_pt)
    font_obj.bold = spec.bold
    font_obj.italic = False  # ACB: never italic
    font_obj.all_caps = spec.all_caps
    font_obj.color.rgb = RGBColor(0, 0, 0)


_ALIGN_MAP = {
    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _resolve_profile_style_overrides(
    standards_profile: str,
) -> tuple[str | None, float | None]:
    """Return (font_family_override, line_spacing_override) by standards profile."""
    normalized = (standards_profile or C.StandardsProfile.ACB_2025.value).strip().lower()

    if normalized == C.StandardsProfile.APH_SUBMISSION.value:
        # APH profile favors APHont and recommends ~1.25 line spacing.
        return C.APH_ACCEPTED_FONT_FAMILIES[0], C.APH_LINE_SPACING_RECOMMENDED

    # ACB and combined strict use existing ACB defaults for template construction.
    return None, None


def _apply_paragraph_format(pf, spec: C.ParaDef) -> None:
    """Apply a ParaDef to a python-docx ParagraphFormat."""
    pf.alignment = _ALIGN_MAP[spec.alignment]
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = spec.line_spacing
    pf.space_before = Pt(spec.space_before_pt)
    pf.space_after = Pt(spec.space_after_pt)
    pf.widow_control = spec.widow_control
    pf.keep_with_next = spec.keep_with_next

    if spec.first_line_indent_in:
        pf.first_line_indent = Inches(spec.first_line_indent_in)
    if spec.left_indent_in:
        pf.left_indent = Inches(spec.left_indent_in)
    if spec.hanging_indent_in and spec.left_indent_in:
        pf.first_line_indent = Inches(-spec.hanging_indent_in)


def _configure_styles(
    doc: Document,
    *,
    list_indent_in: float = C.LIST_INDENT_IN,
    list_hanging_in: float = C.LIST_HANGING_IN,
    font_family_override: str | None = None,
    line_spacing_override: float | None = None,
    style_size_overrides: dict[str, float] | None = None,
) -> None:
    """Set all ACB styles on a Document."""
    styles_table = C.effective_styles(style_size_overrides)
    for style_name, style_def in styles_table.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            # Style doesn't exist -- skip gracefully
            continue

        font_spec = (
            replace(style_def.font, name=font_family_override)
            if font_family_override
            else style_def.font
        )
        para_spec = (
            replace(style_def.para, line_spacing=line_spacing_override)
            if line_spacing_override is not None
            else style_def.para
        )

        _apply_font(style.font, font_spec)
        _apply_paragraph_format(style.paragraph_format, para_spec)

        # Apply user-configurable list indent to list styles
        if style_name in ("List Bullet", "List Number"):
            pf = style.paragraph_format
            pf.left_indent = Inches(list_indent_in)
            if list_indent_in > 0 and list_hanging_in > 0:
                pf.first_line_indent = Inches(-list_hanging_in)
            elif list_indent_in == 0:
                pf.first_line_indent = None

    # Remove theme color from headings (Word defaults to blue)
    for heading in (
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Heading 4",
        "Heading 5",
        "Heading 6",
    ):
        try:
            style = doc.styles[heading]
            rPr = style.element.find(qn("w:rPr"))
            if rPr is not None:
                color_elem = rPr.find(qn("w:color"))
                if color_elem is not None:
                    color_elem.attrib.pop(qn("w:themeColor"), None)
                    color_elem.attrib.pop(qn("w:themeShade"), None)
                    color_elem.set(qn("w:val"), "000000")
        except KeyError:
            continue


def _configure_page_setup(doc: Document, bound: bool = False) -> None:
    """Set page dimensions, margins, and orientation."""
    for section in doc.sections:
        section.page_width = Inches(C.PAGE_WIDTH_IN)
        section.page_height = Inches(C.PAGE_HEIGHT_IN)
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = Inches(C.MARGIN_TOP_IN)
        section.bottom_margin = Inches(C.MARGIN_BOTTOM_IN)
        section.right_margin = Inches(C.MARGIN_RIGHT_IN)
        left = C.MARGIN_LEFT_IN + (C.BINDING_EXTRA_IN if bound else 0)
        section.left_margin = Inches(left)


def _disable_hyphenation(doc: Document) -> None:
    """Disable automatic hyphenation at the document level."""
    settings = doc.settings.element
    auto_hyph = settings.find(qn("w:autoHyphenation"))
    if auto_hyph is not None:
        settings.remove(auto_hyph)
    elem = OxmlElement("w:autoHyphenation")
    elem.set(qn("w:val"), "false")
    settings.append(elem)


def _add_page_numbers(doc: Document, *, font_family: str = C.FONT_FAMILY) -> None:
    """Add ACB-compliant page numbers to the footer (Arial 18pt bold, right-aligned)."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        # Clear existing footer content
        for para in footer.paragraphs:
            for run in para.runs:
                run.text = ""

        paragraph = (
            footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = paragraph.add_run()
        run.font.name = font_family
        run.font.size = Pt(C.FOOTER_SIZE_PT)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Insert PAGE field code
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run._r.append(instr)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)


def _set_document_language(doc: Document, lang: str = "en-US") -> None:
    """Set the default document language."""
    styles_elem = doc.styles.element
    doc_defaults = styles_elem.find(qn("w:docDefaults"))
    rpr_default = None
    if doc_defaults is not None:
        rpr_def_elem = doc_defaults.find(qn("w:rPrDefault"))
        if rpr_def_elem is not None:
            rpr_default = rpr_def_elem.find(qn("w:rPr"))
    if rpr_default is not None:
        lang_elem = rpr_default.find(qn("w:lang"))
        if lang_elem is None:
            lang_elem = OxmlElement("w:lang")
            rpr_default.append(lang_elem)
        lang_elem.set(qn("w:val"), lang)
        lang_elem.set(qn("w:eastAsia"), lang)
        lang_elem.set(qn("w:bidi"), lang)


def _add_sample_content(doc: Document, *, allowed_heading_levels: list[int] | None = None) -> None:
    """Add sample paragraphs that demonstrate each style.

    The sample heading ladder respects user-selected allowed heading levels.
    """
    levels = sorted(set(allowed_heading_levels or [1, 2, 3]))
    levels = [level for level in levels if 1 <= level <= 6]
    if not levels:
        levels = [1]

    top_level = levels[0]
    doc.add_heading("ACB Large Print Template", level=top_level)

    p = doc.add_paragraph(
        "This template is pre-configured with all styles required by the "
        "American Council of the Blind (ACB) Large Print Guidelines. "
        "Replace this text with your content."
    )
    p.style = doc.styles["Normal"]

    for level in levels:
        if level == top_level:
            continue
        doc.add_heading(f"Heading {level} Example", level=level)
        doc.add_paragraph(f"Body text under a level {level} heading.")

    # Bullet list
    doc.add_paragraph("First bullet item", style="List Bullet")
    doc.add_paragraph("Second bullet item", style="List Bullet")
    doc.add_paragraph("Third bullet item", style="List Bullet")

    # Numbered list
    doc.add_paragraph("First numbered item", style="List Number")
    doc.add_paragraph("Second numbered item", style="List Number")


def _save_as_dotx(doc: Document, output_path: Path) -> None:
    """Save a Document as a .dotx template by adjusting the content type."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        doc.save(tmp_path)
        with zipfile.ZipFile(tmp_path, "r") as zin:
            with zipfile.ZipFile(str(output_path), "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == "[Content_Types].xml":
                        data = data.replace(
                            b"application/vnd.openxmlformats-officedocument"
                            b".wordprocessingml.document.main+xml",
                            b"application/vnd.openxmlformats-officedocument"
                            b".wordprocessingml.template.main+xml",
                        )
                    zout.writestr(item, data)
    finally:
        os.unlink(tmp_path)


def create_template(
    output_path: str | Path,
    *,
    bound: bool = False,
    include_sample: bool = True,
    title: str = "",
    list_indent_in: float = C.LIST_INDENT_IN,
    list_hanging_in: float = C.LIST_HANGING_IN,
    standards_profile: str = C.StandardsProfile.ACB_2025.value,
    allowed_heading_levels: list[int] | None = None,
    style_size_overrides: dict[str, float] | None = None,
) -> Path:
    """Create a complete ACB Large Print .dotx template.

    Args:
        output_path: Where to save the template file.
        bound: If True, add binding margin to left side.
        include_sample: If True, add demonstration content.
        title: Document title for properties (WCAG 2.4.2).
        list_indent_in: Left indent for list styles in inches.
        list_hanging_in: Hanging indent for list styles in inches.
        standards_profile: Standards profile for style defaults.
        allowed_heading_levels: Heading levels to include in sample content.

    Returns:
        Path to the created template file.
    """
    output_path = Path(output_path)
    doc = Document()

    # Core properties
    doc.core_properties.title = title or "ACB Large Print Document"
    doc.core_properties.author = "ACB Large Print Tool"

    font_family_override, line_spacing_override = _resolve_profile_style_overrides(
        standards_profile
    )

    _configure_styles(
        doc,
        list_indent_in=list_indent_in,
        list_hanging_in=list_hanging_in,
        font_family_override=font_family_override,
        line_spacing_override=line_spacing_override,
        style_size_overrides=style_size_overrides,
    )
    _configure_page_setup(doc, bound=bound)
    _disable_hyphenation(doc)
    _add_page_numbers(doc, font_family=font_family_override or C.FONT_FAMILY)
    _set_document_language(doc)

    if include_sample:
        _add_sample_content(doc, allowed_heading_levels=allowed_heading_levels)

    output_path.parent.mkdir(parents=True, exist_ok=True)

    if output_path.suffix.lower() == ".dotx":
        _save_as_dotx(doc, output_path)
    else:
        doc.save(str(output_path))

    return output_path


def install_template(template_path: str | Path) -> Path:
    """Copy a template to Word's user Templates folder.

    Returns:
        The destination path where the template was installed.
    """
    template_path = Path(template_path)
    templates_dir = Path(os.environ.get("APPDATA", "")) / "Microsoft" / "Templates"
    templates_dir.mkdir(parents=True, exist_ok=True)
    dest = templates_dir / template_path.name
    shutil.copy2(str(template_path), str(dest))
    return dest


def install_startup_template(template_path: str | Path) -> Path:
    """Copy a template to Word's STARTUP folder so it loads automatically.

    Returns:
        The destination path where the template was installed.
    """
    template_path = Path(template_path)
    startup_dir = Path(os.environ.get("APPDATA", "")) / "Microsoft" / "Word" / "STARTUP"
    startup_dir.mkdir(parents=True, exist_ok=True)
    dest = startup_dir / template_path.name
    shutil.copy2(str(template_path), str(dest))
    return dest
